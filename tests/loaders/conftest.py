# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Agentic Brain Contributors
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Loaders test fixtures."""

import json
from pathlib import Path

import pytest


@pytest.fixture
def fixtures_dir():
    """Return path to fixtures directory."""
    return Path(__file__).parent / "fixtures"


@pytest.fixture
def sample_txt(fixtures_dir):
    """Return path to sample text file."""
    return fixtures_dir / "sample.txt"


@pytest.fixture
def sample_md(fixtures_dir):
    """Return path to sample markdown file."""
    return fixtures_dir / "sample.md"


@pytest.fixture
def sample_html(fixtures_dir):
    """Return path to sample HTML file."""
    return fixtures_dir / "sample.html"


@pytest.fixture
def sample_csv(fixtures_dir):
    """Return path to sample CSV file."""
    return fixtures_dir / "sample.csv"


@pytest.fixture
def sample_json_file(fixtures_dir):
    """Return path to sample JSON file."""
    return fixtures_dir / "sample.json"


@pytest.fixture
def sample_jsonl(fixtures_dir):
    """Return path to sample JSONL file."""
    return fixtures_dir / "sample.jsonl"


@pytest.fixture
def sample_docx(fixtures_dir, tmp_path):
    """Create a sample DOCX file."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Second paragraph.")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"

        output_path = tmp_path / "sample.docx"
        doc.save(output_path)
        return output_path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture
def sample_pdf(fixtures_dir, tmp_path):
    """Create a sample PDF file."""
    try:
        from pypdf import PdfWriter
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        # Create PDF with reportlab
        pdf_path = tmp_path / "sample.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a test PDF.")
        c.drawString(100, 710, "Line 3 of content.")
        c.save()
        return pdf_path
    except ImportError:
        pytest.skip("pypdf or reportlab not installed")


@pytest.fixture
def temp_directory(tmp_path):
    """Create a temporary directory with various file types."""
    # Create text files
    (tmp_path / "file1.txt").write_text("Content 1")
    (tmp_path / "file2.txt").write_text("Content 2")

    # Create markdown
    (tmp_path / "readme.md").write_text("# Title\n\nContent")

    # Create CSV
    (tmp_path / "data.csv").write_text("a,b,c\n1,2,3\n4,5,6")

    # Create JSON
    json_path = tmp_path / "data.json"
    json_path.write_text(json.dumps({"key": "value"}))

    # Create subdirectory
    sub_dir = tmp_path / "subdir"
    sub_dir.mkdir()
    (sub_dir / "nested.txt").write_text("Nested content")

    return tmp_path
