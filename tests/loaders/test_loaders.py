# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Agentic Brain Contributors
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Tests for document loaders."""

import json
from pathlib import Path

import pytest

from agentic_brain.loaders import (
    CSVLoader,
    CSVRowLoader,
    DirectoryLoader,
    Document,
    DocumentLoader,
    DocxLoader,
    HTMLLoader,
    JSONLinesLoader,
    JSONLinesRowLoader,
    JSONLoader,
    MarkdownHeadingLoader,
    MarkdownLoader,
    PDFLoader,
    PDFPageLoader,
    TextLoader,
)

# Optional dependencies
try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pypdf

    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from bs4 import BeautifulSoup

    HAS_BEAUTIFULSOUP = True
except ImportError:
    HAS_BEAUTIFULSOUP = False


class TestDocument:
    """Test Document class."""

    def test_document_creation(self):
        """Test creating a document."""
        doc = Document(content="Test content", source="test.txt")
        assert doc.content == "Test content"
        assert doc.source == "test.txt"
        assert isinstance(doc.metadata, dict)

    def test_document_with_metadata(self):
        """Test document with metadata."""
        metadata = {"key": "value", "count": 42}
        doc = Document(content="Content", source="test.txt", metadata=metadata)
        assert doc.metadata == metadata
        assert doc.metadata["key"] == "value"

    def test_document_to_dict(self):
        """Test converting document to dict."""
        doc = Document(content="Content", source="test.txt", metadata={"key": "val"})
        doc_dict = doc.to_dict()
        assert doc_dict["content"] == "Content"
        assert doc_dict["source"] == "test.txt"
        assert doc_dict["metadata"]["key"] == "val"
        assert "loaded_at" in doc_dict


class TestTextLoader:
    """Test TextLoader."""

    @pytest.mark.asyncio
    async def test_load_text_file(self, sample_txt):
        """Test loading a text file."""
        loader = TextLoader()
        docs = await loader.load(sample_txt)
        assert len(docs) == 1
        assert "sample text file" in docs[0].content
        assert docs[0].source == str(sample_txt)

    def test_supported_extensions(self):
        """Test supported extensions."""
        loader = TextLoader()
        exts = loader.supported_extensions()
        assert ".txt" in exts
        assert ".py" in exts

    @pytest.mark.asyncio
    async def test_load_nonexistent_file(self, tmp_path):
        """Test loading nonexistent file raises error."""
        loader = TextLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(tmp_path / "nonexistent.txt")

    def test_can_load(self, sample_txt):
        """Test can_load method."""
        loader = TextLoader()
        assert loader.can_load(sample_txt)
        assert loader.can_load("file.py")
        assert not loader.can_load("file.pdf")


class TestMarkdownLoader:
    """Test MarkdownLoader."""

    @pytest.mark.asyncio
    async def test_load_markdown(self, sample_md):
        """Test loading markdown file."""
        loader = MarkdownLoader()
        docs = await loader.load(sample_md)
        assert len(docs) == 1
        assert "Main Title" in docs[0].content
        assert docs[0].metadata["filename"] == "sample.md"

    def test_markdown_metadata(self, sample_md):
        """Test markdown metadata extraction."""
        loader = MarkdownLoader()
        docs = loader.load_sync(sample_md)
        assert "title" in docs[0].metadata
        assert docs[0].metadata["title"] == "Main Title"

    @pytest.mark.asyncio
    async def test_markdown_with_links(self, sample_md):
        """Test markdown link extraction."""
        loader = MarkdownLoader()
        docs = await loader.load(sample_md)
        assert "links" in docs[0].metadata

    def test_supported_extensions(self):
        """Test supported markdown extensions."""
        loader = MarkdownLoader()
        exts = loader.supported_extensions()
        assert ".md" in exts
        assert ".markdown" in exts


class TestMarkdownHeadingLoader:
    """Test MarkdownHeadingLoader."""

    def test_load_by_headings(self, sample_md):
        """Test loading markdown split by headings."""
        loader = MarkdownHeadingLoader()
        docs = loader.load_sync(sample_md)
        assert len(docs) >= 2
        assert any("Section One" in d.metadata.get("heading", "") for d in docs)

    def test_heading_metadata(self, sample_md):
        """Test heading metadata."""
        loader = MarkdownHeadingLoader()
        docs = loader.load_sync(sample_md)
        for doc in docs:
            if "heading_level" in doc.metadata:
                assert doc.metadata["heading_level"] >= 1


@pytest.mark.skipif(not HAS_BEAUTIFULSOUP, reason="beautifulsoup4 not installed")
class TestHTMLLoader:
    """Test HTMLLoader."""

    @pytest.mark.asyncio
    async def test_load_html(self, sample_html):
        """Test loading HTML file."""
        loader = HTMLLoader()
        docs = await loader.load(sample_html)
        assert len(docs) == 1
        assert "Test Heading" in docs[0].content
        assert "script" not in docs[0].content.lower()

    def test_html_metadata(self, sample_html):
        """Test HTML metadata extraction."""
        loader = HTMLLoader(extract_metadata=True)
        docs = loader.load_sync(sample_html)
        assert "title" in docs[0].metadata
        assert docs[0].metadata["title"] == "Test Document"

    def test_html_links_extraction(self, sample_html):
        """Test HTML link extraction."""
        loader = HTMLLoader(extract_links=True)
        docs = loader.load_sync(sample_html)
        assert "links" in docs[0].metadata

    def test_supported_extensions(self):
        """Test supported HTML extensions."""
        loader = HTMLLoader()
        exts = loader.supported_extensions()
        assert ".html" in exts
        assert ".htm" in exts


class TestCSVLoader:
    """Test CSVLoader."""

    @pytest.mark.asyncio
    async def test_load_csv(self, sample_csv):
        """Test loading CSV file."""
        loader = CSVLoader()
        docs = await loader.load(sample_csv)
        assert len(docs) == 1
        assert "Alice" in docs[0].content
        assert "name" in docs[0].content

    def test_csv_metadata(self, sample_csv):
        """Test CSV metadata."""
        loader = CSVLoader()
        docs = loader.load_sync(sample_csv)
        assert docs[0].metadata["row_count"] == 3
        assert docs[0].metadata["column_count"] == 3

    def test_supported_extensions(self):
        """Test supported CSV extensions."""
        loader = CSVLoader()
        assert ".csv" in loader.supported_extensions()


class TestCSVRowLoader:
    """Test CSVRowLoader."""

    def test_load_csv_rows(self, sample_csv):
        """Test loading CSV with one document per row."""
        loader = CSVRowLoader()
        docs = loader.load_sync(sample_csv)
        assert len(docs) == 3
        assert docs[0].metadata["row_number"] == 1
        assert "Alice" in docs[0].content

    def test_csv_row_metadata(self, sample_csv):
        """Test CSV row metadata."""
        loader = CSVRowLoader()
        docs = loader.load_sync(sample_csv)
        for i, doc in enumerate(docs):
            assert doc.metadata["row_number"] == i + 1
            assert doc.metadata["total_rows"] == 3


class TestJSONLoader:
    """Test JSONLoader."""

    @pytest.mark.asyncio
    async def test_load_json(self, sample_json_file):
        """Test loading JSON file."""
        loader = JSONLoader()
        docs = await loader.load(sample_json_file)
        assert len(docs) == 1
        assert "Test Document" in docs[0].content

    def test_json_with_filter(self, sample_json_file):
        """Test loading JSON with jq filter."""
        loader = JSONLoader(jq_filter="data")
        docs = loader.load_sync(sample_json_file)
        assert len(docs) == 1
        content = json.loads(docs[0].content)
        assert "items" in content

    def test_json_metadata(self, sample_json_file):
        """Test JSON metadata."""
        loader = JSONLoader()
        docs = loader.load_sync(sample_json_file)
        assert docs[0].metadata["filename"] == "sample.json"

    def test_supported_extensions(self):
        """Test supported JSON extensions."""
        loader = JSONLoader()
        assert ".json" in loader.supported_extensions()


class TestJSONLinesLoader:
    """Test JSONLinesLoader."""

    @pytest.mark.asyncio
    async def test_load_jsonl(self, sample_jsonl):
        """Test loading JSONL file."""
        loader = JSONLinesLoader()
        docs = await loader.load(sample_jsonl)
        assert len(docs) == 1
        content = json.loads(docs[0].content)
        assert len(content) == 3

    def test_jsonl_metadata(self, sample_jsonl):
        """Test JSONL metadata."""
        loader = JSONLinesLoader()
        docs = loader.load_sync(sample_jsonl)
        assert docs[0].metadata["line_count"] == 3

    def test_supported_extensions(self):
        """Test supported JSONL extensions."""
        loader = JSONLinesLoader()
        exts = loader.supported_extensions()
        assert ".jsonl" in exts
        assert ".ndjson" in exts


class TestJSONLinesRowLoader:
    """Test JSONLinesRowLoader."""

    def test_load_jsonl_rows(self, sample_jsonl):
        """Test loading JSONL with one document per line."""
        loader = JSONLinesRowLoader()
        docs = loader.load_sync(sample_jsonl)
        assert len(docs) == 3
        assert docs[0].metadata["line_number"] == 1

    def test_jsonl_row_metadata(self, sample_jsonl):
        """Test JSONL row metadata."""
        loader = JSONLinesRowLoader()
        docs = loader.load_sync(sample_jsonl)
        for i, doc in enumerate(docs):
            assert doc.metadata["line_number"] == i + 1


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestDocxLoader:
    """Test DocxLoader."""

    def test_load_docx(self, sample_docx):
        """Test loading DOCX file."""
        loader = DocxLoader()
        docs = loader.load_sync(sample_docx)
        assert len(docs) == 1
        assert "Test Document" in docs[0].content
        assert "Test Title" in docs[0].metadata.get("title", "")

    def test_docx_metadata(self, sample_docx):
        """Test DOCX metadata."""
        loader = DocxLoader()
        docs = loader.load_sync(sample_docx)
        assert "Test Title" in docs[0].metadata.get("title", "")
        assert docs[0].metadata["filename"] == "sample.docx"

    def test_docx_with_table(self, sample_docx):
        """Test DOCX with table."""
        loader = DocxLoader()
        docs = loader.load_sync(sample_docx)
        assert "[TABLE]" in docs[0].content


@pytest.mark.skipif(not HAS_PYPDF, reason="pypdf not installed")
class TestPDFLoader:
    """Test PDFLoader."""

    def test_load_pdf(self, sample_pdf):
        """Test loading PDF file."""
        loader = PDFLoader()
        docs = loader.load_sync(sample_pdf)
        assert len(docs) == 1
        assert "Test PDF" in docs[0].content

    def test_pdf_metadata(self, sample_pdf):
        """Test PDF metadata."""
        loader = PDFLoader()
        docs = loader.load_sync(sample_pdf)
        assert "num_pages" in docs[0].metadata
        assert docs[0].metadata["filename"] == "sample.pdf"

    def test_supported_extensions(self):
        """Test supported PDF extensions."""
        loader = PDFLoader()
        assert ".pdf" in loader.supported_extensions()


@pytest.mark.skipif(not HAS_PYPDF, reason="pypdf not installed")
class TestPDFPageLoader:
    """Test PDFPageLoader."""

    def test_load_pdf_pages(self, sample_pdf):
        """Test loading PDF with one document per page."""
        loader = PDFPageLoader()
        docs = loader.load_sync(sample_pdf)
        assert len(docs) >= 1
        for doc in docs:
            assert "page" in doc.metadata


class TestDirectoryLoader:
    """Test DirectoryLoader."""

    @pytest.mark.asyncio
    async def test_load_directory(self, temp_directory):
        """Test loading directory."""
        loader = DirectoryLoader()
        docs = await loader.load(temp_directory)
        assert len(docs) > 0

    def test_load_directory_sync(self, temp_directory):
        """Test synchronous directory loading."""
        loader = DirectoryLoader()
        docs = loader.load_sync(temp_directory)
        assert len(docs) > 0

    def test_directory_with_glob_pattern(self, temp_directory):
        """Test directory loading with glob pattern."""
        loader = DirectoryLoader(glob_pattern="*.txt")
        docs = loader.load_sync(temp_directory)
        assert all(".txt" in d.source or "txt" in d.source for d in docs)

    def test_directory_recursive(self, temp_directory):
        """Test recursive directory loading."""
        loader = DirectoryLoader(recursive=True)
        docs = loader.load_sync(temp_directory)
        assert len(docs) > 0

    def test_add_custom_loader(self, temp_directory):
        """Test adding custom loader."""
        loader = DirectoryLoader()

        class CustomLoader(DocumentLoader):
            async def load(self, source):
                return []

            def supported_extensions(self):
                return [".custom"]

            def can_load(self, source):
                return False

        loader.add_loader(".custom", CustomLoader())
        assert ".custom" in loader.loaders

    def test_directory_not_found(self):
        """Test loading nonexistent directory."""
        loader = DirectoryLoader()
        with pytest.raises(NotADirectoryError):
            loader.load_sync("/nonexistent/directory")


class TestDocumentLoaderABC:
    """Test DocumentLoader abstract base class."""

    def test_cannot_instantiate_abstract(self):
        """Test that abstract DocumentLoader cannot be instantiated."""
        with pytest.raises(TypeError):
            DocumentLoader()


class TestLoaderIntegration:
    """Integration tests for loaders."""

    def test_mixed_file_types(self, temp_directory):
        """Test loading mixed file types."""
        loader = DirectoryLoader()
        docs = loader.load_sync(temp_directory)
        sources = [d.source for d in docs]
        assert len(sources) > 0

    def test_error_handling(self, tmp_path):
        """Test error handling for invalid files."""
        invalid_json = tmp_path / "invalid.json"
        invalid_json.write_text("{ invalid json }")

        loader = JSONLoader()
        with pytest.raises(ValueError):
            loader.load_sync(invalid_json)

    def test_metadata_preservation(self, sample_txt):
        """Test that metadata is preserved across loaders."""
        loader = TextLoader()
        docs = loader.load_sync(sample_txt)
        assert len(docs) == 1
        assert "filename" in docs[0].metadata
        assert "size" in docs[0].metadata

    def test_large_csv_handling(self, tmp_path):
        """Test handling large CSV files."""
        csv_file = tmp_path / "large.csv"
        lines = ["col1,col2,col3"] + [f"row{i},val{i},data{i}" for i in range(100)]
        csv_file.write_text("\n".join(lines))

        loader = CSVLoader()
        docs = loader.load_sync(csv_file)
        assert docs[0].metadata["row_count"] == 100

    def test_empty_file_handling(self, tmp_path):
        """Test handling empty files."""
        empty_file = tmp_path / "empty.txt"
        empty_file.write_text("")

        loader = TextLoader()
        docs = loader.load_sync(empty_file)
        assert len(docs) == 1
        assert docs[0].content == ""
