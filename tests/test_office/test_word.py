"""Tests for Word-focused office conversions."""

from __future__ import annotations

import importlib.util
import sys
from types import SimpleNamespace

import pytest

DOCX_LIBRARY_AVAILABLE = importlib.util.find_spec("docx") is not None


@pytest.mark.parametrize("platform_name", ["Darwin", "Windows"])
def test_docx_to_pdf_uses_docx2pdf_when_available(
    make_converter, monkeypatch, platform_name, sample_office_files, tmp_path
) -> None:
    converter = make_converter(platform_name=platform_name, docx2pdf=True)
    calls: list[tuple[str, str]] = []
    monkeypatch.setitem(
        sys.modules,
        "docx2pdf",
        SimpleNamespace(convert=lambda src, dst: calls.append((src, dst))),
    )

    output = tmp_path / "out" / "sample.pdf"
    result = converter.docx_to_pdf(sample_office_files.docx, output)

    assert result == output
    assert calls == [(str(sample_office_files.docx), str(output))]


def test_docx_to_pdf_falls_back_to_libreoffice_after_docx2pdf_error(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(platform_name="Darwin", docx2pdf=True)

    def explode(*_args, **_kwargs):
        raise RuntimeError("docx2pdf boom")

    libreoffice_output = tmp_path / "libreoffice" / "sample.pdf"
    libreoffice_output.parent.mkdir(parents=True)
    libreoffice_output.write_bytes(b"%PDF-1.7")

    monkeypatch.setitem(sys.modules, "docx2pdf", SimpleNamespace(convert=explode))
    monkeypatch.setattr(
        converter,
        "_run_libreoffice_conversion",
        lambda input_path, output_format, output_dir: libreoffice_output,
    )

    output = tmp_path / "renamed" / "final.pdf"
    result = converter.docx_to_pdf(sample_office_files.docx, output)

    assert result == output
    assert output.exists()
    assert not libreoffice_output.exists()


def test_docx_to_pdf_uses_libreoffice_on_non_docx2pdf_platform(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(platform_name="Linux", docx2pdf=True)
    called: list[tuple[str, str, str]] = []

    def fake_run(input_path, output_format, output_dir):
        called.append((str(input_path), output_format, str(output_dir)))
        output = tmp_path / "out" / "sample.pdf"
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_bytes(b"%PDF")
        return output

    monkeypatch.setattr(converter, "_run_libreoffice_conversion", fake_run)

    output = tmp_path / "out" / "sample.pdf"
    assert converter.docx_to_pdf(sample_office_files.docx, output) == output
    assert called == [(str(sample_office_files.docx), "pdf", str(output.parent))]


def test_docx_to_txt_uses_python_docx(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(python_docx=True)
    fake_document = SimpleNamespace(
        paragraphs=[SimpleNamespace(text="First"), SimpleNamespace(text="Second")]
    )
    monkeypatch.setitem(
        sys.modules, "docx", SimpleNamespace(Document=lambda path: fake_document)
    )

    output = tmp_path / "text" / "sample.txt"
    result = converter.docx_to_txt(sample_office_files.docx, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "First\nSecond"


def test_docx_to_txt_falls_back_to_libreoffice_on_python_docx_error(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(python_docx=True)
    monkeypatch.setitem(
        sys.modules,
        "docx",
        SimpleNamespace(Document=lambda path: (_ for _ in ()).throw(ValueError("bad"))),
    )
    fallback = tmp_path / "fallback" / "sample.txt"
    fallback.parent.mkdir(parents=True)
    fallback.write_text("fallback text", encoding="utf-8")
    monkeypatch.setattr(
        converter,
        "_run_libreoffice_conversion",
        lambda input_path, output_format, output_dir: fallback,
    )

    output = tmp_path / "converted" / "sample.txt"
    output.parent.mkdir(parents=True, exist_ok=True)
    result = converter.docx_to_txt(sample_office_files.docx, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "fallback text"


def test_docx_to_html_uses_mammoth_and_records_messages(
    make_converter, monkeypatch, sample_office_files, tmp_path, caplog
) -> None:
    converter = make_converter(mammoth=True)
    caplog.set_level("DEBUG")

    class Result:
        value = "<h1>Hello</h1>"
        messages = ["warning-one"]

    monkeypatch.setitem(
        sys.modules,
        "mammoth",
        SimpleNamespace(convert_to_html=lambda file_obj: Result()),
    )

    output = tmp_path / "html" / "sample.html"
    result = converter.docx_to_html(sample_office_files.docx, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "<h1>Hello</h1>"


def test_docx_to_html_falls_back_to_libreoffice_when_mammoth_fails(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(mammoth=True)
    monkeypatch.setitem(
        sys.modules,
        "mammoth",
        SimpleNamespace(
            convert_to_html=lambda file_obj: (_ for _ in ()).throw(RuntimeError("nope"))
        ),
    )
    fallback = tmp_path / "fallback" / "sample.html"
    fallback.parent.mkdir(parents=True)
    fallback.write_text("<p>fallback</p>", encoding="utf-8")
    monkeypatch.setattr(
        converter,
        "_run_libreoffice_conversion",
        lambda input_path, output_format, output_dir: fallback,
    )

    output = tmp_path / "html" / "final.html"
    assert converter.docx_to_html(sample_office_files.docx, output) == output
    assert output.read_text(encoding="utf-8") == "<p>fallback</p>"


def test_docx_to_markdown_uses_mammoth(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(mammoth=True)

    class Result:
        value = "# Heading\n\nBody"

    monkeypatch.setitem(
        sys.modules,
        "mammoth",
        SimpleNamespace(convert_to_markdown=lambda file_obj: Result()),
    )

    output = tmp_path / "md" / "sample.md"
    result = converter.docx_to_markdown(sample_office_files.docx, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "# Heading\n\nBody"


def test_docx_to_markdown_falls_back_via_html_conversion(
    make_converter, monkeypatch, sample_office_files, tmp_path
) -> None:
    converter = make_converter(mammoth=True)
    monkeypatch.setitem(
        sys.modules,
        "mammoth",
        SimpleNamespace(
            convert_to_markdown=lambda file_obj: (_ for _ in ()).throw(
                RuntimeError("markdown failure")
            )
        ),
    )

    def fake_docx_to_html(_input_path, html_path):
        html_path.write_text("<h1>Title</h1><p>Body</p>", encoding="utf-8")
        return html_path

    monkeypatch.setattr(converter, "docx_to_html", fake_docx_to_html)

    output = tmp_path / "md" / "sample.md"
    result = converter.docx_to_markdown(sample_office_files.docx, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "TitleBody"
    assert not output.with_suffix(".html").exists()


@pytest.mark.parametrize(
    ("target_format", "method_name"),
    [
        ("pdf", "docx_to_pdf"),
        ("txt", "docx_to_txt"),
        ("html", "docx_to_html"),
        ("md", "docx_to_markdown"),
    ],
)
def test_word_converters_are_registered(
    make_converter, target_format: str, method_name: str
) -> None:
    converter = make_converter()
    method = converter.get_converter("docx", target_format)
    assert method is not None
    assert method.__name__ == method_name


@pytest.mark.parametrize("fixture_name", ["empty_docx", "corrupted_docx", "large_docx"])
def test_word_edge_cases_can_fall_back_to_libreoffice(
    make_converter, monkeypatch, sample_office_files, tmp_path, fixture_name: str
) -> None:
    converter = make_converter(python_docx=True)
    source = getattr(sample_office_files, fixture_name)
    monkeypatch.setitem(
        sys.modules,
        "docx",
        SimpleNamespace(Document=lambda path: (_ for _ in ()).throw(ValueError("bad"))),
    )

    def fake_conversion(input_path, output_format, output_dir):
        output = tmp_path / "out" / f"{input_path.stem}.txt"
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(f"converted:{input_path.name}", encoding="utf-8")
        return output

    monkeypatch.setattr(converter, "_run_libreoffice_conversion", fake_conversion)

    output = tmp_path / "final" / f"{source.stem}.txt"
    output.parent.mkdir(parents=True, exist_ok=True)
    result = converter.docx_to_txt(source, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == f"converted:{source.name}"


@pytest.mark.skipif(not DOCX_LIBRARY_AVAILABLE, reason="python-docx not installed")
def test_docx_to_txt_with_real_python_docx(make_converter, tmp_path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    source = tmp_path / "real.docx"
    doc.save(source)

    converter = make_converter(python_docx=True)
    output = tmp_path / "real.txt"
    result = converter.docx_to_txt(source, output)

    assert result == output
    assert output.read_text(encoding="utf-8") == "Hello\nWorld"
