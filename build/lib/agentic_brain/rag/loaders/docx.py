# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Joseph Webber <joseph.webber@me.com>
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Microsoft Word document loaders for RAG pipelines."""

import logging
from io import BytesIO
from pathlib import Path
from typing import Optional

from .base import BaseLoader, LoadedDocument

logger = logging.getLogger(__name__)

# Check for python-docx
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxLoader(BaseLoader):
    """Load and extract text from DOCX documents.

    Example:
        loader = DocxLoader()
        doc = loader.load_document("report.docx")
        docs = loader.load_folder("documents/")
    """

    def __init__(
        self,
        base_path: str = ".",
        include_headers_footers: bool = False,
        include_comments: bool = False,
    ):
        self.base_path = Path(base_path)
        self.include_headers_footers = include_headers_footers
        self.include_comments = include_comments

    @property
    def source_name(self) -> str:
        return "docx"

    def authenticate(self) -> bool:
        """No authentication needed for local files."""
        return True

    def _extract_text(self, docx_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed")
            return "[DOCX content - install python-docx]"

        try:
            document = docx.Document(BytesIO(docx_bytes))
            paragraphs = []

            # Main body
            for para in document.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single DOCX document."""
        try:
            path = Path(doc_id)
            if not path.is_absolute():
                path = self.base_path / path

            if not path.exists():
                logger.error(f"File not found: {path}")
                return None

            if path.suffix.lower() not in (".docx", ".doc"):
                logger.warning(f"Not a Word document: {path}")
                return None

            with open(path, "rb") as f:
                docx_bytes = f.read()

            content = self._extract_text(docx_bytes)

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=str(path),
                filename=path.name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=len(docx_bytes),
                metadata={"path": str(path)},
            )
        except Exception as e:
            logger.error(f"Failed to load DOCX {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all DOCX files from a folder."""
        docs = []
        path = Path(folder_path)
        if not path.is_absolute():
            path = self.base_path / path

        if not path.exists():
            logger.error(f"Folder not found: {path}")
            return docs

        pattern = "**/*.docx" if recursive else "*.docx"

        for docx_path in path.glob(pattern):
            doc = self.load_document(str(docx_path))
            if doc:
                docs.append(doc)

        logger.info(f"Loaded {len(docs)} DOCX files from {folder_path}")
        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search not supported for local files."""
        return []


# Alias for backward compatibility
WordLoader = DocxLoader


__all__ = ["DocxLoader", "WordLoader"]
