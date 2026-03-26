# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 Joseph Webber <joseph.webber@me.com>
#
# This file is part of Agentic Brain.
"""Cloud and On-Premise Document Loaders for RAG.

Load documents from various sources for ingestion into RAG pipelines.

Supported services:
- Google Drive (docs, sheets, PDFs, folders)
- Gmail (emails with attachments)
- iCloud Drive (documents, folders)
- Firebase Firestore (collections, real-time sync)
- Amazon S3 / MinIO (object storage, self-hosted)
- MongoDB (document database, on-premise)

Example:
    from agentic_brain.rag import (
        GoogleDriveLoader, GmailLoader, iCloudLoader,
        FirestoreLoader, S3Loader, MongoDBLoader
    )

    # Google Drive
    drive = GoogleDriveLoader(credentials_path="credentials.json")
    docs = drive.load_folder("My Project")

    # Gmail
    gmail = GmailLoader(credentials_path="credentials.json")
    emails = gmail.load_recent(days=7, query="from:boss@company.com")

    # iCloud
    icloud = iCloudLoader(apple_id="user@icloud.com", password="app-password")
    docs = icloud.load_folder("Documents/Work")

    # Firestore
    firestore = FirestoreLoader(service_account_path="firebase-adminsdk.json")
    docs = firestore.load_collection("articles", where=[("status", "==", "published")])

    # S3 / MinIO
    s3 = S3Loader(bucket="documents", endpoint_url="http://minio:9000")
    docs = s3.load_folder("reports/2024/")

    # MongoDB
    mongo = MongoDBLoader(uri="mongodb://localhost:27017", database="knowledge")
    docs = mongo.load_collection(filter={"category": "tech"})
"""

import base64
import json
import logging
import mimetypes
import os
import re
from abc import ABC, abstractmethod
from collections.abc import Generator
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """A document loaded from a cloud source."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    source: str = ""  # e.g., "google_drive", "gmail", "icloud"
    source_id: str = ""  # Original ID in the source system
    filename: str = ""
    mime_type: str = "text/plain"
    created_at: Optional[datetime] = None
    modified_at: Optional[datetime] = None
    size_bytes: int = 0

    def to_dict(self) -> dict[str, Any]:
        """Serialize document to dictionary."""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "source": self.source,
            "source_id": self.source_id,
            "filename": self.filename,
            "mime_type": self.mime_type,
            "created_at": self.created_at.isoformat() if self.created_at else None,
            "modified_at": self.modified_at.isoformat() if self.modified_at else None,
            "size_bytes": self.size_bytes,
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "LoadedDocument":
        """Create document from dictionary."""
        created_at = None
        modified_at = None
        if data.get("created_at"):
            created_at = datetime.fromisoformat(data["created_at"])
        if data.get("modified_at"):
            modified_at = datetime.fromisoformat(data["modified_at"])

        return cls(
            content=data.get("content", ""),
            metadata=data.get("metadata", {}),
            source=data.get("source", ""),
            source_id=data.get("source_id", ""),
            filename=data.get("filename", ""),
            mime_type=data.get("mime_type", "text/plain"),
            created_at=created_at,
            modified_at=modified_at,
            size_bytes=data.get("size_bytes", 0),
        )


class BaseLoader(ABC):
    """Abstract base class for document loaders."""

    @property
    @abstractmethod
    def source_name(self) -> str:
        """Name of the source (e.g., 'google_drive')."""
        pass

    @abstractmethod
    def authenticate(self) -> bool:
        """Authenticate with the service. Returns True if successful."""
        pass

    @abstractmethod
    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by ID."""
        pass

    @abstractmethod
    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a folder."""
        pass

    @abstractmethod
    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents."""
        pass

    def _extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes. Tries multiple methods."""
        # Try PyPDF2
        try:
            from io import BytesIO

            import PyPDF2

            reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            if text.strip():
                return text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed: {e}")

        # Try pdfplumber
        try:
            from io import BytesIO

            import pdfplumber

            with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            if text.strip():
                return text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {e}")

        logger.warning(
            "No PDF extraction library available. Install PyPDF2 or pdfplumber."
        )
        return "[PDF content - extraction library not available]"

    def _clean_html(self, html: str) -> str:
        """Convert HTML to plain text."""
        # Try BeautifulSoup
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html, "html.parser")
            # Remove script and style elements
            for element in soup(["script", "style", "head", "meta", "link"]):
                element.decompose()
            text = soup.get_text(separator="\n")
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            return "\n".join(line for line in lines if line)
        except ImportError:
            pass

        # Fallback: basic regex
        text = re.sub(
            r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE
        )
        text = re.sub(
            r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE
        )
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()


# ============================================================================
# GOOGLE DRIVE LOADER
# ============================================================================

# Check for Google API
try:
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google.oauth2.service_account import Credentials as ServiceAccountCredentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload

    GOOGLE_API_AVAILABLE = True
except ImportError:
    GOOGLE_API_AVAILABLE = False


class GoogleDriveLoader(BaseLoader):
    """Load documents from Google Drive.

    Supports:
    - Google Docs (exported as plain text)
    - Google Sheets (exported as CSV)
    - Google Slides (exported as plain text)
    - PDFs (text extraction)
    - Plain text files
    - Word documents (.docx)

    Authentication:
    - Service account (recommended for servers)
    - OAuth2 user credentials (for desktop apps)

    Example:
        # Using service account
        loader = GoogleDriveLoader(
            credentials_path="service-account.json",
            use_service_account=True
        )

        # Using OAuth2 (will open browser for auth)
        loader = GoogleDriveLoader(
            credentials_path="client_secrets.json",
            token_path="token.json"
        )

        # Load a folder
        docs = loader.load_folder("My Project Documents")

        # Search
        results = loader.search("quarterly report 2024")
    """

    SCOPES = [
        "https://www.googleapis.com/auth/drive.readonly",
        "https://www.googleapis.com/auth/drive.metadata.readonly",
    ]

    # Google Workspace MIME types and their export formats
    GOOGLE_MIME_TYPES = {
        "application/vnd.google-apps.document": ("text/plain", ".txt"),
        "application/vnd.google-apps.spreadsheet": ("text/csv", ".csv"),
        "application/vnd.google-apps.presentation": ("text/plain", ".txt"),
        "application/vnd.google-apps.drawing": ("image/png", ".png"),
    }

    # Supported file types for text extraction
    SUPPORTED_MIME_TYPES = {
        "text/plain",
        "text/csv",
        "text/markdown",
        "text/html",
        "application/pdf",
        "application/json",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }

    def __init__(
        self,
        credentials_path: Optional[str] = None,
        token_path: str = "gdrive_token.json",
        use_service_account: bool = False,
        max_file_size_mb: int = 50,
    ):
        """Initialize Google Drive loader.

        Args:
            credentials_path: Path to credentials JSON file
            token_path: Path to store OAuth2 token
            use_service_account: Use service account instead of OAuth2
            max_file_size_mb: Maximum file size to download
        """
        if not GOOGLE_API_AVAILABLE:
            raise ImportError(
                "Google API libraries not installed. Run: "
                "pip install google-auth google-auth-oauthlib google-auth-httplib2 google-api-python-client"
            )

        self.credentials_path = credentials_path
        self.token_path = token_path
        self.use_service_account = use_service_account
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self._service = None
        self._credentials = None

    @property
    def source_name(self) -> str:
        return "google_drive"

    def authenticate(self) -> bool:
        """Authenticate with Google Drive API."""
        try:
            if self.use_service_account:
                if not self.credentials_path:
                    raise ValueError("credentials_path required for service account")
                self._credentials = ServiceAccountCredentials.from_service_account_file(
                    self.credentials_path, scopes=self.SCOPES
                )
            else:
                # OAuth2 flow
                if os.path.exists(self.token_path):
                    self._credentials = Credentials.from_authorized_user_file(
                        self.token_path, self.SCOPES
                    )

                if not self._credentials or not self._credentials.valid:
                    if (
                        self._credentials
                        and self._credentials.expired
                        and self._credentials.refresh_token
                    ):
                        self._credentials.refresh(Request())
                    else:
                        if not self.credentials_path:
                            raise ValueError(
                                "credentials_path required for OAuth2 flow"
                            )
                        flow = InstalledAppFlow.from_client_secrets_file(
                            self.credentials_path, self.SCOPES
                        )
                        self._credentials = flow.run_local_server(port=0)

                    # Save token for future use
                    with open(self.token_path, "w") as token:
                        token.write(self._credentials.to_json())

            self._service = build("drive", "v3", credentials=self._credentials)
            logger.info("Google Drive authentication successful")
            return True

        except Exception as e:
            logger.error(f"Google Drive authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        """Ensure service is authenticated."""
        if not self._service and not self.authenticate():
            raise RuntimeError("Failed to authenticate with Google Drive")

    def _get_file_content(
        self, file_id: str, mime_type: str, name: str
    ) -> Optional[str]:
        """Download and extract text content from a file."""
        try:
            from io import BytesIO

            # Handle Google Workspace files (need export)
            if mime_type in self.GOOGLE_MIME_TYPES:
                export_mime, _ = self.GOOGLE_MIME_TYPES[mime_type]
                request = self._service.files().export_media(
                    fileId=file_id, mimeType=export_mime
                )
            else:
                request = self._service.files().get_media(fileId=file_id)

            buffer = BytesIO()
            downloader = MediaIoBaseDownload(buffer, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()

            content_bytes = buffer.getvalue()

            # Extract text based on type
            if mime_type == "application/pdf":
                return self._extract_text_from_pdf(content_bytes)
            elif mime_type in ("text/html", "application/xhtml+xml"):
                return self._clean_html(content_bytes.decode("utf-8", errors="replace"))
            elif mime_type.startswith("text/") or mime_type in ("application/json",):
                return content_bytes.decode("utf-8", errors="replace")
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # Word document
                try:
                    from io import BytesIO

                    import docx

                    doc = docx.Document(BytesIO(content_bytes))
                    return "\n".join(p.text for p in doc.paragraphs)
                except ImportError:
                    logger.warning(
                        "python-docx not installed for Word document extraction"
                    )
                    return f"[Word document: {name}]"
            elif mime_type in self.GOOGLE_MIME_TYPES:
                # Already exported to text
                return content_bytes.decode("utf-8", errors="replace")
            else:
                logger.debug(f"Unsupported mime type: {mime_type}")
                return None

        except Exception as e:
            logger.error(f"Failed to get content for {name}: {e}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by its Google Drive ID."""
        self._ensure_authenticated()

        try:
            # Get file metadata
            file = (
                self._service.files()
                .get(
                    fileId=doc_id,
                    fields="id, name, mimeType, size, createdTime, modifiedTime, parents",
                )
                .execute()
            )

            mime_type = file.get("mimeType", "")
            name = file.get("name", "")

            # Check if supported
            is_google_type = mime_type in self.GOOGLE_MIME_TYPES
            is_supported = mime_type in self.SUPPORTED_MIME_TYPES or is_google_type

            if not is_supported:
                logger.debug(f"Skipping unsupported file type: {name} ({mime_type})")
                return None

            # Check size
            size = int(file.get("size", 0))
            if size > self.max_file_size and not is_google_type:
                logger.warning(f"Skipping large file: {name} ({size} bytes)")
                return None

            # Get content
            content = self._get_file_content(doc_id, mime_type, name)
            if not content:
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=name,
                mime_type=mime_type,
                created_at=(
                    datetime.fromisoformat(file["createdTime"].replace("Z", "+00:00"))
                    if file.get("createdTime")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(file["modifiedTime"].replace("Z", "+00:00"))
                    if file.get("modifiedTime")
                    else None
                ),
                size_bytes=size,
                metadata={
                    "drive_id": doc_id,
                    "name": name,
                    "mime_type": mime_type,
                    "parents": file.get("parents", []),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load document {doc_id}: {e}")
            return None

    def _get_folder_id(self, folder_path: str) -> Optional[str]:
        """Get folder ID from path like 'Folder/Subfolder'."""
        if not folder_path or folder_path == "/":
            return "root"

        parts = [p for p in folder_path.split("/") if p]
        current_parent = "root"

        for part in parts:
            query = f"name='{part}' and '{current_parent}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
            results = (
                self._service.files()
                .list(q=query, fields="files(id, name)", pageSize=1)
                .execute()
            )

            files = results.get("files", [])
            if not files:
                logger.warning(f"Folder not found: {part}")
                return None

            current_parent = files[0]["id"]

        return current_parent

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a Google Drive folder."""
        self._ensure_authenticated()

        folder_id = self._get_folder_id(folder_path)
        if not folder_id:
            logger.error(f"Folder not found: {folder_path}")
            return []

        return self._load_folder_by_id(folder_id, recursive)

    def _load_folder_by_id(
        self, folder_id: str, recursive: bool
    ) -> list[LoadedDocument]:
        """Load all documents from a folder by ID."""
        documents = []
        page_token = None

        while True:
            query = f"'{folder_id}' in parents and trashed=false"
            results = (
                self._service.files()
                .list(
                    q=query,
                    fields="nextPageToken, files(id, name, mimeType, size, createdTime, modifiedTime)",
                    pageSize=100,
                    pageToken=page_token,
                )
                .execute()
            )

            for file in results.get("files", []):
                if file["mimeType"] == "application/vnd.google-apps.folder":
                    if recursive:
                        documents.extend(self._load_folder_by_id(file["id"], recursive))
                else:
                    doc = self.load_document(file["id"])
                    if doc:
                        documents.append(doc)

            page_token = results.get("nextPageToken")
            if not page_token:
                break

        logger.info(f"Loaded {len(documents)} documents from folder")
        return documents

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents in Google Drive."""
        self._ensure_authenticated()

        documents = []
        page_token = None

        # Build search query - escape single quotes
        escaped_query = query.replace("'", "\\'")
        search_query = f"fullText contains '{escaped_query}' and trashed=false"

        while len(documents) < max_results:
            results = (
                self._service.files()
                .list(
                    q=search_query,
                    fields="nextPageToken, files(id, name, mimeType, size, createdTime, modifiedTime)",
                    pageSize=min(100, max_results - len(documents)),
                    pageToken=page_token,
                )
                .execute()
            )

            for file in results.get("files", []):
                if file["mimeType"] != "application/vnd.google-apps.folder":
                    doc = self.load_document(file["id"])
                    if doc:
                        documents.append(doc)

            page_token = results.get("nextPageToken")
            if not page_token:
                break

        logger.info(f"Found {len(documents)} documents matching '{query}'")
        return documents

    def list_folders(self, parent_path: str = "") -> list[dict[str, str]]:
        """List subfolders in a path."""
        self._ensure_authenticated()

        parent_id = self._get_folder_id(parent_path) if parent_path else "root"
        if not parent_id:
            return []

        folders = []
        query = f"'{parent_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"

        results = (
            self._service.files()
            .list(q=query, fields="files(id, name)", pageSize=100)
            .execute()
        )

        for f in results.get("files", []):
            folders.append({"id": f["id"], "name": f["name"]})

        return folders


# ============================================================================
# GMAIL LOADER
# ============================================================================


class GmailLoader(BaseLoader):
    """Load emails from Gmail as documents.

    Supports:
    - Email body (plain text or HTML converted to text)
    - Attachments (text, PDF, Word docs)
    - Search with Gmail query syntax
    - Label-based filtering

    Example:
        loader = GmailLoader(credentials_path="client_secrets.json")

        # Load recent emails
        emails = loader.load_recent(days=7)

        # Search with Gmail syntax
        emails = loader.search("from:boss@company.com has:attachment")

        # Load by label
        emails = loader.load_by_label("IMPORTANT")
    """

    SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]

    def __init__(
        self,
        credentials_path: Optional[str] = None,
        token_path: str = "gmail_token.json",
        include_attachments: bool = True,
        max_attachment_size_mb: int = 25,
    ):
        """Initialize Gmail loader.

        Args:
            credentials_path: Path to OAuth2 credentials JSON
            token_path: Path to store OAuth2 token
            include_attachments: Whether to extract attachments
            max_attachment_size_mb: Maximum attachment size to download
        """
        if not GOOGLE_API_AVAILABLE:
            raise ImportError(
                "Google API libraries not installed. Run: "
                "pip install google-auth google-auth-oauthlib google-api-python-client"
            )

        self.credentials_path = credentials_path
        self.token_path = token_path
        self.include_attachments = include_attachments
        self.max_attachment_size = max_attachment_size_mb * 1024 * 1024
        self._service = None
        self._credentials = None

    @property
    def source_name(self) -> str:
        return "gmail"

    def authenticate(self) -> bool:
        """Authenticate with Gmail API."""
        try:
            if os.path.exists(self.token_path):
                self._credentials = Credentials.from_authorized_user_file(
                    self.token_path, self.SCOPES
                )

            if not self._credentials or not self._credentials.valid:
                if (
                    self._credentials
                    and self._credentials.expired
                    and self._credentials.refresh_token
                ):
                    self._credentials.refresh(Request())
                else:
                    if not self.credentials_path:
                        raise ValueError("credentials_path required for OAuth2 flow")
                    flow = InstalledAppFlow.from_client_secrets_file(
                        self.credentials_path, self.SCOPES
                    )
                    self._credentials = flow.run_local_server(port=0)

                with open(self.token_path, "w") as token:
                    token.write(self._credentials.to_json())

            self._service = build("gmail", "v1", credentials=self._credentials)
            logger.info("Gmail authentication successful")
            return True

        except Exception as e:
            logger.error(f"Gmail authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        """Ensure service is authenticated."""
        if not self._service and not self.authenticate():
            raise RuntimeError("Failed to authenticate with Gmail")

    def _get_message_body(self, payload: dict) -> str:
        """Extract message body from payload."""
        body = ""

        if "body" in payload and payload["body"].get("data"):
            body = base64.urlsafe_b64decode(payload["body"]["data"]).decode(
                "utf-8", errors="replace"
            )

        if "parts" in payload:
            for part in payload["parts"]:
                mime_type = part.get("mimeType", "")

                if mime_type == "text/plain":
                    if part.get("body", {}).get("data"):
                        body = base64.urlsafe_b64decode(part["body"]["data"]).decode(
                            "utf-8", errors="replace"
                        )
                        break
                elif mime_type == "text/html":
                    if part.get("body", {}).get("data"):
                        html = base64.urlsafe_b64decode(part["body"]["data"]).decode(
                            "utf-8", errors="replace"
                        )
                        body = self._clean_html(html)
                elif mime_type.startswith("multipart/"):
                    body = self._get_message_body(part)
                    if body:
                        break

        return body

    def _get_attachments(self, message_id: str, payload: dict) -> list[LoadedDocument]:
        """Extract attachments from email."""
        attachments = []

        def process_parts(parts: list[dict]):
            for part in parts:
                filename = part.get("filename", "")
                if filename and part.get("body", {}).get("attachmentId"):
                    mime_type = part.get("mimeType", "")
                    size = int(part.get("body", {}).get("size", 0))

                    if size > self.max_attachment_size:
                        logger.debug(f"Skipping large attachment: {filename}")
                        continue

                    try:
                        attachment = (
                            self._service.users()
                            .messages()
                            .attachments()
                            .get(
                                userId="me",
                                messageId=message_id,
                                id=part["body"]["attachmentId"],
                            )
                            .execute()
                        )

                        data = base64.urlsafe_b64decode(attachment["data"])

                        # Extract text based on type
                        content = None
                        if mime_type == "application/pdf":
                            content = self._extract_text_from_pdf(data)
                        elif mime_type.startswith("text/"):
                            content = data.decode("utf-8", errors="replace")

                        if content:
                            attachments.append(
                                LoadedDocument(
                                    content=content,
                                    source="gmail_attachment",
                                    source_id=f"{message_id}_{part['body']['attachmentId']}",
                                    filename=filename,
                                    mime_type=mime_type,
                                    size_bytes=size,
                                    metadata={
                                        "message_id": message_id,
                                        "attachment_filename": filename,
                                    },
                                )
                            )
                    except Exception as e:
                        logger.error(f"Failed to get attachment {filename}: {e}")

                if "parts" in part:
                    process_parts(part["parts"])

        if "parts" in payload:
            process_parts(payload["parts"])

        return attachments

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single email by message ID."""
        self._ensure_authenticated()

        try:
            message = (
                self._service.users()
                .messages()
                .get(userId="me", id=doc_id, format="full")
                .execute()
            )

            # Extract headers
            headers = {
                h["name"].lower(): h["value"]
                for h in message.get("payload", {}).get("headers", [])
            }

            subject = headers.get("subject", "(No Subject)")
            from_addr = headers.get("from", "")
            to_addr = headers.get("to", "")
            date_str = headers.get("date", "")

            # Parse date
            created_at = None
            if date_str:
                try:
                    from email.utils import parsedate_to_datetime

                    created_at = parsedate_to_datetime(date_str)
                except Exception:
                    pass

            # Get body
            body = self._get_message_body(message.get("payload", {}))

            # Format content
            content = f"Subject: {subject}\nFrom: {from_addr}\nTo: {to_addr}\nDate: {date_str}\n\n{body}"

            doc = LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=f"{subject}.eml",
                mime_type="message/rfc822",
                created_at=created_at,
                metadata={
                    "message_id": doc_id,
                    "subject": subject,
                    "from": from_addr,
                    "to": to_addr,
                    "labels": message.get("labelIds", []),
                    "thread_id": message.get("threadId", ""),
                },
            )

            return doc

        except Exception as e:
            logger.error(f"Failed to load email {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load emails by label (folder). Use label name like 'INBOX' or 'IMPORTANT'."""
        return self.load_by_label(folder_path)

    def load_by_label(self, label: str, max_results: int = 100) -> list[LoadedDocument]:
        """Load emails with a specific label."""
        self._ensure_authenticated()
        return self._search_messages(f"label:{label}", max_results)

    def load_recent(
        self, days: int = 7, max_results: int = 100
    ) -> list[LoadedDocument]:
        """Load recent emails from the last N days."""
        self._ensure_authenticated()
        return self._search_messages(f"newer_than:{days}d", max_results)

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search emails using Gmail query syntax."""
        self._ensure_authenticated()
        return self._search_messages(query, max_results)

    def _search_messages(self, query: str, max_results: int) -> list[LoadedDocument]:
        """Internal search implementation."""
        documents = []
        page_token = None

        while len(documents) < max_results:
            results = (
                self._service.users()
                .messages()
                .list(
                    userId="me",
                    q=query,
                    maxResults=min(100, max_results - len(documents)),
                    pageToken=page_token,
                )
                .execute()
            )

            for msg in results.get("messages", []):
                doc = self.load_document(msg["id"])
                if doc:
                    documents.append(doc)

                    # Get attachments if enabled
                    if self.include_attachments:
                        full_msg = (
                            self._service.users()
                            .messages()
                            .get(userId="me", id=msg["id"], format="full")
                            .execute()
                        )
                        attachments = self._get_attachments(
                            msg["id"], full_msg.get("payload", {})
                        )
                        documents.extend(attachments)

            page_token = results.get("nextPageToken")
            if not page_token:
                break

        logger.info(f"Loaded {len(documents)} emails matching '{query}'")
        return documents

    def list_labels(self) -> list[dict[str, str]]:
        """List all Gmail labels."""
        self._ensure_authenticated()

        results = self._service.users().labels().list(userId="me").execute()
        return [{"id": l["id"], "name": l["name"]} for l in results.get("labels", [])]


# ============================================================================
# ICLOUD LOADER
# ============================================================================

# Check for pyicloud
try:
    from pyicloud import PyiCloudService
    from pyicloud.exceptions import PyiCloudFailedLoginException

    PYICLOUD_AVAILABLE = True
except ImportError:
    PYICLOUD_AVAILABLE = False


class iCloudLoader(BaseLoader):
    """Load documents from iCloud Drive.

    Supports:
    - Text files (.txt, .md, .json, .csv)
    - PDFs (with text extraction)
    - Word documents (.docx)
    - Folder navigation

    Note: Requires 2FA. On first use, you'll need to verify via SMS/device.

    Example:
        loader = iCloudLoader(
            apple_id="your@icloud.com",
            password="your-app-specific-password"  # Use app-specific password!
        )

        # Load from a folder
        docs = loader.load_folder("Documents/Work")

        # Search
        results = loader.search("report")
    """

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".json",
        ".csv",
        ".pdf",
        ".docx",
        ".rtf",
        ".html",
    }

    def __init__(
        self,
        apple_id: Optional[str] = None,
        password: Optional[str] = None,
        cookie_directory: str = ".icloud",
        max_file_size_mb: int = 50,
    ):
        """Initialize iCloud loader.

        Args:
            apple_id: Apple ID email address
            password: Password or app-specific password
            cookie_directory: Directory to store session cookies
            max_file_size_mb: Maximum file size to download
        """
        if not PYICLOUD_AVAILABLE:
            raise ImportError("pyicloud not installed. Run: pip install pyicloud")

        self.apple_id = apple_id or os.environ.get("ICLOUD_APPLE_ID")
        self.password = password or os.environ.get("ICLOUD_PASSWORD")
        self.cookie_directory = cookie_directory
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self._api = None

    @property
    def source_name(self) -> str:
        return "icloud"

    def authenticate(self) -> bool:
        """Authenticate with iCloud."""
        try:
            if not self.apple_id or not self.password:
                raise ValueError("apple_id and password required")

            self._api = PyiCloudService(
                self.apple_id, self.password, cookie_directory=self.cookie_directory
            )

            if self._api.requires_2fa:
                logger.warning("Two-factor authentication required")
                print("Two-factor authentication required.")
                code = input("Enter the code sent to your device: ")
                if not self._api.validate_2fa_code(code):
                    logger.error("Invalid 2FA code")
                    return False

                # Trust the session
                if not self._api.is_trusted_session:
                    print("Session not trusted. Requesting trust...")
                    self._api.trust_session()

            logger.info("iCloud authentication successful")
            return True

        except PyiCloudFailedLoginException as e:
            logger.error(f"iCloud login failed: {e}")
            return False
        except Exception as e:
            logger.error(f"iCloud authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        """Ensure API is authenticated."""
        if not self._api and not self.authenticate():
            raise RuntimeError("Failed to authenticate with iCloud")

    def _get_folder(self, path: str):
        """Navigate to a folder by path."""
        if not path or path == "/":
            return self._api.drive.root

        current = self._api.drive.root
        parts = [p for p in path.split("/") if p]

        for part in parts:
            found = False
            for item in current.dir():
                if item == part:
                    current = current[item]
                    found = True
                    break

            if not found:
                raise FileNotFoundError(f"Folder not found: {part}")

        return current

    def _download_file(self, item) -> Optional[bytes]:
        """Download file content."""
        try:
            response = item.open(stream=True)
            return response.raw.read()
        except Exception as e:
            logger.error(f"Failed to download {item.name}: {e}")
            return None

    def _extract_content(
        self, data: bytes, filename: str, mime_type: str
    ) -> Optional[str]:
        """Extract text content from file data."""
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._extract_text_from_pdf(data)
        elif ext in (".txt", ".md", ".csv", ".json"):
            return data.decode("utf-8", errors="replace")
        elif ext == ".html":
            return self._clean_html(data.decode("utf-8", errors="replace"))
        elif ext == ".docx":
            try:
                from io import BytesIO

                import docx

                doc = docx.Document(BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                logger.warning("python-docx not installed")
                return None

        return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a document by path."""
        self._ensure_authenticated()

        try:
            # Navigate to the file
            parts = doc_id.rsplit("/", 1)
            if len(parts) == 2:
                folder_path, filename = parts
                folder = self._get_folder(folder_path)
                item = folder[filename]
            else:
                item = self._api.drive.root[doc_id]

            # Check if it's a file
            if item.type != "file":
                return None

            # Check extension
            ext = Path(item.name).suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                logger.debug(f"Unsupported extension: {item.name}")
                return None

            # Check size
            size = item.size or 0
            if size > self.max_file_size:
                logger.warning(f"File too large: {item.name} ({size} bytes)")
                return None

            # Download and extract
            data = self._download_file(item)
            if not data:
                return None

            mime_type = mimetypes.guess_type(item.name)[0] or "application/octet-stream"
            content = self._extract_content(data, item.name, mime_type)

            if not content:
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=item.name,
                mime_type=mime_type,
                modified_at=(
                    item.date_modified if hasattr(item, "date_modified") else None
                ),
                size_bytes=size,
                metadata={"path": doc_id, "name": item.name, "type": item.type},
            )

        except Exception as e:
            logger.error(f"Failed to load iCloud document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from an iCloud Drive folder."""
        self._ensure_authenticated()

        try:
            folder = self._get_folder(folder_path)
            return self._load_folder_items(folder, folder_path, recursive)
        except FileNotFoundError as e:
            logger.error(str(e))
            return []

    def _load_folder_items(
        self, folder, path: str, recursive: bool
    ) -> list[LoadedDocument]:
        """Recursively load items from folder."""
        documents = []

        for name in folder.dir():
            item = folder[name]
            item_path = f"{path}/{name}" if path else name

            if item.type == "folder" and recursive:
                documents.extend(self._load_folder_items(item, item_path, recursive))
            elif item.type == "file":
                doc = self.load_document(item_path)
                if doc:
                    documents.append(doc)

        logger.info(f"Loaded {len(documents)} documents from {path}")
        return documents

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents (simple name-based search)."""
        self._ensure_authenticated()

        # iCloud doesn't have a search API, so we do a simple name match
        query_lower = query.lower()
        documents = []

        def search_folder(folder, path: str):
            for name in folder.dir():
                if len(documents) >= max_results:
                    return

                item = folder[name]
                item_path = f"{path}/{name}" if path else name

                if item.type == "folder":
                    search_folder(item, item_path)
                elif query_lower in name.lower():
                    doc = self.load_document(item_path)
                    if doc:
                        documents.append(doc)

        search_folder(self._api.drive.root, "")
        logger.info(f"Found {len(documents)} documents matching '{query}'")
        return documents

    def list_folders(self, path: str = "") -> list[dict[str, str]]:
        """List subfolders in a path."""
        self._ensure_authenticated()

        try:
            folder = self._get_folder(path)
            folders = []

            for name in folder.dir():
                item = folder[name]
                if item.type == "folder":
                    folders.append(
                        {"name": name, "path": f"{path}/{name}" if path else name}
                    )

            return folders
        except FileNotFoundError:
            return []


# ============================================================================
# FIRESTORE DOCUMENT LOADER
# ============================================================================

# Check for Firebase
try:
    import firebase_admin
    from firebase_admin import credentials as firebase_credentials
    from firebase_admin import firestore

    FIREBASE_AVAILABLE = True
except ImportError:
    FIREBASE_AVAILABLE = False


class FirestoreLoader(BaseLoader):
    """Load documents from Firebase Firestore.

    Firestore is a document database that stores JSON-like documents
    in collections. This loader pulls documents into the RAG pipeline.

    Authentication options:
        1. Service account JSON file (for server deployments)
        2. Default credentials (for GCP environments)
        3. Existing Firebase app (for apps already using Firebase)

    Example:
        # With service account
        loader = FirestoreLoader(
            service_account_path="firebase-adminsdk.json",
            project_id="my-project"
        )

        # Load all docs from a collection
        docs = loader.load_collection("documents")

        # Load with query filter
        docs = loader.load_collection(
            "documents",
            where=[("status", "==", "published"), ("category", "==", "tech")]
        )

        # Real-time listener (returns generator)
        for doc in loader.watch_collection("documents"):
            print(f"New/updated doc: {doc.filename}")
    """

    def __init__(
        self,
        service_account_path: Optional[str] = None,
        project_id: Optional[str] = None,
        app_name: str = "[DEFAULT]",
        content_field: str = "content",
        title_field: str = "title",
        metadata_fields: Optional[list[str]] = None,
    ):
        """Initialize Firestore loader.

        Args:
            service_account_path: Path to service account JSON
            project_id: Firebase project ID (optional if in service account)
            app_name: Firebase app name (for multiple apps)
            content_field: Field name containing document content
            title_field: Field name containing document title
            metadata_fields: Additional fields to include in metadata
        """
        if not FIREBASE_AVAILABLE:
            raise ImportError(
                "Firebase Admin SDK not available. Install with: "
                "pip install firebase-admin"
            )

        self._service_account_path = service_account_path
        self._project_id = project_id
        self._app_name = app_name
        self._content_field = content_field
        self._title_field = title_field
        self._metadata_fields = metadata_fields or []
        self._db: Optional[Any] = None
        self._app: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "firestore"

    def authenticate(self) -> bool:
        """Initialize Firebase Admin and get Firestore client."""
        if self._authenticated and self._db is not None:
            return True

        try:
            # Check if app already exists
            try:
                self._app = firebase_admin.get_app(self._app_name)
                logger.info(f"Using existing Firebase app: {self._app_name}")
            except ValueError:
                # App doesn't exist, create it
                if self._service_account_path:
                    cred = firebase_credentials.Certificate(self._service_account_path)
                    options = (
                        {"projectId": self._project_id} if self._project_id else {}
                    )
                    self._app = firebase_admin.initialize_app(
                        cred,
                        options,
                        name=self._app_name if self._app_name != "[DEFAULT]" else None,
                    )
                    logger.info("Firebase authenticated with service account")
                else:
                    # Use default credentials (GCP environment)
                    self._app = firebase_admin.initialize_app(
                        options=(
                            {"projectId": self._project_id}
                            if self._project_id
                            else None
                        ),
                        name=self._app_name if self._app_name != "[DEFAULT]" else None,
                    )
                    logger.info("Firebase authenticated with default credentials")

            self._db = firestore.client(self._app)
            self._authenticated = True
            return True

        except Exception as e:
            logger.error(f"Firebase authentication failed: {e}")
            return False

    def _ensure_authenticated(self) -> None:
        """Ensure we're authenticated before operations."""
        if not self._authenticated and not self.authenticate():
            raise RuntimeError("Firestore authentication required")

    def _doc_to_loaded_document(
        self, doc_snapshot: Any, collection: str
    ) -> Optional[LoadedDocument]:
        """Convert Firestore document snapshot to LoadedDocument."""
        try:
            data = doc_snapshot.to_dict()
            if not data:
                return None

            # Get content
            content = data.get(self._content_field, "")
            if not content:
                # Try common fallbacks
                for fallback in ["text", "body", "description", "data"]:
                    if fallback in data:
                        content = str(data[fallback])
                        break

            if not content:
                # Use all fields as content
                content = json.dumps(data, indent=2, default=str)

            # Get title/filename
            title = data.get(self._title_field, "")
            if not title:
                for fallback in ["name", "filename", "subject", "heading"]:
                    if fallback in data:
                        title = str(data[fallback])
                        break
            if not title:
                title = doc_snapshot.id

            # Build metadata
            metadata = {
                "collection": collection,
                "document_id": doc_snapshot.id,
            }

            # Add specified metadata fields
            for field in self._metadata_fields:
                if field in data:
                    metadata[field] = data[field]

            # Add all non-content fields to metadata
            for key, value in data.items():
                if key not in [self._content_field, self._title_field]:
                    if key not in metadata:
                        try:
                            # Ensure JSON-serializable
                            json.dumps(value, default=str)
                            metadata[key] = value
                        except (TypeError, ValueError):
                            metadata[key] = str(value)

            # Handle timestamps
            created_at = None
            modified_at = None

            if hasattr(doc_snapshot, "create_time") and doc_snapshot.create_time:
                created_at = doc_snapshot.create_time
            if hasattr(doc_snapshot, "update_time") and doc_snapshot.update_time:
                modified_at = doc_snapshot.update_time

            # Also check for timestamp fields in data
            for ts_field in ["created_at", "createdAt", "timestamp", "created"]:
                if ts_field in data and created_at is None:
                    ts = data[ts_field]
                    if hasattr(ts, "isoformat"):
                        created_at = ts
                    break

            for ts_field in ["updated_at", "updatedAt", "modified", "modified_at"]:
                if ts_field in data and modified_at is None:
                    ts = data[ts_field]
                    if hasattr(ts, "isoformat"):
                        modified_at = ts
                    break

            return LoadedDocument(
                content=str(content),
                metadata=metadata,
                source="firestore",
                source_id=f"{collection}/{doc_snapshot.id}",
                filename=title,
                mime_type="application/json",
                created_at=created_at,
                modified_at=modified_at,
                size_bytes=len(str(content).encode("utf-8")),
            )

        except Exception as e:
            logger.error(f"Failed to convert Firestore document: {e}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by ID.

        Args:
            doc_id: Document path as "collection/document_id"
        """
        self._ensure_authenticated()

        try:
            # Parse collection/doc_id format
            if "/" not in doc_id:
                raise ValueError(
                    f"Invalid doc_id format: {doc_id}. Use 'collection/document_id'"
                )

            parts = doc_id.split("/")
            collection = "/".join(parts[:-1])
            document_id = parts[-1]

            doc_ref = self._db.collection(collection).document(document_id)
            doc_snapshot = doc_ref.get()

            if not doc_snapshot.exists:
                logger.warning(f"Document not found: {doc_id}")
                return None

            return self._doc_to_loaded_document(doc_snapshot, collection)

        except Exception as e:
            logger.error(f"Failed to load Firestore document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a collection.

        In Firestore, 'folder' = collection.

        Args:
            folder_path: Collection name (e.g., "documents" or "users/uid/posts")
            recursive: If True, also load subcollections
        """
        return self.load_collection(folder_path, recursive=recursive)

    def load_collection(
        self,
        collection: str,
        where: Optional[list[tuple]] = None,
        order_by: Optional[str] = None,
        limit: Optional[int] = None,
        recursive: bool = False,
    ) -> list[LoadedDocument]:
        """Load documents from a Firestore collection.

        Args:
            collection: Collection path (e.g., "documents")
            where: List of filter tuples: [(field, op, value), ...]
                   Operators: ==, !=, <, <=, >, >=, in, not-in, array-contains
            order_by: Field to order by (prefix with - for descending)
            limit: Maximum documents to return
            recursive: Also load subcollections

        Example:
            docs = loader.load_collection(
                "articles",
                where=[("status", "==", "published"), ("views", ">", 100)],
                order_by="-created_at",
                limit=50
            )
        """
        self._ensure_authenticated()

        documents = []

        try:
            # Build query
            query = self._db.collection(collection)

            if where:
                for field, op, value in where:
                    query = query.where(field, op, value)

            if order_by:
                if order_by.startswith("-"):
                    query = query.order_by(
                        order_by[1:], direction=firestore.Query.DESCENDING
                    )
                else:
                    query = query.order_by(order_by)

            if limit:
                query = query.limit(limit)

            # Execute query
            for doc_snapshot in query.stream():
                loaded_doc = self._doc_to_loaded_document(doc_snapshot, collection)
                if loaded_doc:
                    documents.append(loaded_doc)

                # Handle subcollections if recursive
                if recursive:
                    subcollections = doc_snapshot.reference.collections()
                    for subcoll in subcollections:
                        subcoll_path = f"{collection}/{doc_snapshot.id}/{subcoll.id}"
                        subdocs = self.load_collection(subcoll_path, recursive=True)
                        documents.extend(subdocs)

            logger.info(f"Loaded {len(documents)} documents from {collection}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load collection {collection}: {e}")
            return []

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents containing text.

        Note: Firestore doesn't support full-text search natively.
        This searches across all collections by title/content fields.
        For production, consider using Algolia or Elasticsearch.

        Args:
            query: Search text
            max_results: Maximum results to return
        """
        self._ensure_authenticated()

        documents = []
        query_lower = query.lower()

        try:
            # Get all collections
            collections = self._db.collections()

            for collection in collections:
                for doc_snapshot in collection.stream():
                    data = doc_snapshot.to_dict()
                    if not data:
                        continue

                    # Check if query matches any text field
                    searchable = []
                    searchable.append(str(data.get(self._content_field, "")))
                    searchable.append(str(data.get(self._title_field, "")))
                    searchable.append(str(data.get("description", "")))
                    searchable.append(doc_snapshot.id)

                    combined = " ".join(searchable).lower()

                    if query_lower in combined:
                        loaded_doc = self._doc_to_loaded_document(
                            doc_snapshot, collection.id
                        )
                        if loaded_doc:
                            documents.append(loaded_doc)
                            if len(documents) >= max_results:
                                break

                if len(documents) >= max_results:
                    break

            logger.info(f"Found {len(documents)} documents matching '{query}'")
            return documents

        except Exception as e:
            logger.error(f"Search failed: {e}")
            return []

    def watch_collection(
        self, collection: str, where: Optional[list[tuple]] = None
    ) -> Generator[LoadedDocument, None, None]:
        """Watch a collection for real-time updates.

        Yields LoadedDocument for each new or updated document.
        This is a blocking generator - runs until cancelled.

        Args:
            collection: Collection to watch
            where: Optional filters

        Example:
            for doc in loader.watch_collection("messages"):
                pipeline.add_document(doc.content, metadata=doc.metadata)
        """
        self._ensure_authenticated()

        import queue
        import threading

        doc_queue: queue.Queue = queue.Queue()
        stop_event = threading.Event()

        def on_snapshot(doc_snapshot, changes, read_time):
            for change in changes:
                if change.type.name in ("ADDED", "MODIFIED"):
                    doc_queue.put(change.document)

        # Build query
        query = self._db.collection(collection)
        if where:
            for field, op, value in where:
                query = query.where(field, op, value)

        # Start listener
        unsubscribe = query.on_snapshot(on_snapshot)

        try:
            while not stop_event.is_set():
                try:
                    doc_snapshot = doc_queue.get(timeout=1.0)
                    loaded_doc = self._doc_to_loaded_document(doc_snapshot, collection)
                    if loaded_doc:
                        yield loaded_doc
                except queue.Empty:
                    continue
        finally:
            unsubscribe()

    def save_document(
        self,
        collection: str,
        content: str,
        title: str = "",
        metadata: Optional[dict[str, Any]] = None,
        document_id: Optional[str] = None,
    ) -> Optional[str]:
        """Save a document to Firestore.

        Useful for storing RAG-processed documents back to Firestore.

        Args:
            collection: Target collection
            content: Document content
            title: Document title
            metadata: Additional fields
            document_id: Specific ID (auto-generated if not provided)

        Returns:
            Document ID if successful, None otherwise
        """
        self._ensure_authenticated()

        try:
            data = {
                self._content_field: content,
                self._title_field: title or "Untitled",
                "created_at": firestore.SERVER_TIMESTAMP,
                "updated_at": firestore.SERVER_TIMESTAMP,
            }

            if metadata:
                data.update(metadata)

            if document_id:
                self._db.collection(collection).document(document_id).set(data)
                return document_id
            else:
                doc_ref = self._db.collection(collection).add(data)
                return doc_ref[1].id

        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return None

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from Firestore.

        Args:
            doc_id: Document path as "collection/document_id"
        """
        self._ensure_authenticated()

        try:
            if "/" not in doc_id:
                raise ValueError(f"Invalid doc_id format: {doc_id}")

            parts = doc_id.split("/")
            collection = "/".join(parts[:-1])
            document_id = parts[-1]

            self._db.collection(collection).document(document_id).delete()
            logger.info(f"Deleted document: {doc_id}")
            return True

        except Exception as e:
            logger.error(f"Failed to delete document {doc_id}: {e}")
            return False

    def list_collections(self) -> list[str]:
        """List all root collections in Firestore."""
        self._ensure_authenticated()

        try:
            return [coll.id for coll in self._db.collections()]
        except Exception as e:
            logger.error(f"Failed to list collections: {e}")
            return []


# ============================================================================
# S3 / MINIO DOCUMENT LOADER
# ============================================================================

# Check for boto3 (AWS SDK)
try:
    import boto3
    from botocore.config import Config as BotoConfig
    from botocore.exceptions import ClientError, NoCredentialsError

    BOTO3_AVAILABLE = True
except ImportError:
    BOTO3_AVAILABLE = False


class S3Loader(BaseLoader):
    """Load documents from Amazon S3 or MinIO (S3-compatible).

    Works with:
    - Amazon S3 (cloud)
    - MinIO (self-hosted, S3-compatible)
    - Any S3-compatible storage (Wasabi, DigitalOcean Spaces, etc.)

    Authentication options:
        1. AWS credentials file (~/.aws/credentials)
        2. Environment variables (AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY)
        3. Explicit credentials in constructor
        4. IAM role (for EC2/Lambda)

    Example:
        # AWS S3
        loader = S3Loader(bucket="my-documents")
        docs = loader.load_folder("reports/2024/")

        # MinIO (self-hosted)
        loader = S3Loader(
            bucket="documents",
            endpoint_url="http://minio.local:9000",
            access_key="minioadmin",
            secret_key="minioadmin"
        )
        docs = loader.load_folder("project-docs/")
    """

    # Supported text file extensions
    TEXT_EXTENSIONS = {
        ".txt",
        ".md",
        ".markdown",
        ".rst",
        ".csv",
        ".json",
        ".xml",
        ".yaml",
        ".yml",
        ".html",
        ".htm",
        ".log",
        ".ini",
        ".cfg",
        ".py",
        ".js",
        ".ts",
        ".java",
        ".cpp",
        ".c",
        ".h",
        ".go",
        ".rs",
        ".rb",
        ".php",
        ".sql",
        ".sh",
        ".bash",
        ".zsh",
    }

    def __init__(
        self,
        bucket: str,
        endpoint_url: Optional[str] = None,
        access_key: Optional[str] = None,
        secret_key: Optional[str] = None,
        region: str = "us-east-1",
        prefix: str = "",
        include_metadata: bool = True,
    ):
        """Initialize S3/MinIO loader.

        Args:
            bucket: S3 bucket name
            endpoint_url: Custom endpoint for MinIO/S3-compatible (e.g., "http://minio:9000")
            access_key: AWS access key ID (or MinIO access key)
            secret_key: AWS secret access key (or MinIO secret key)
            region: AWS region (default: us-east-1)
            prefix: Default prefix/folder to use
            include_metadata: Include S3 object metadata in documents
        """
        if not BOTO3_AVAILABLE:
            raise ImportError("boto3 not available. Install with: pip install boto3")

        self._bucket = bucket
        self._endpoint_url = endpoint_url
        self._access_key = access_key
        self._secret_key = secret_key
        self._region = region
        self._prefix = prefix.strip("/")
        self._include_metadata = include_metadata
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "s3"

    def authenticate(self) -> bool:
        """Initialize S3 client."""
        if self._authenticated and self._client is not None:
            return True

        try:
            # Build client kwargs
            client_kwargs = {
                "service_name": "s3",
                "region_name": self._region,
            }

            if self._endpoint_url:
                client_kwargs["endpoint_url"] = self._endpoint_url
                # For MinIO, often need path-style access
                client_kwargs["config"] = BotoConfig(
                    signature_version="s3v4", s3={"addressing_style": "path"}
                )

            if self._access_key and self._secret_key:
                client_kwargs["aws_access_key_id"] = self._access_key
                client_kwargs["aws_secret_access_key"] = self._secret_key

            self._client = boto3.client(**client_kwargs)

            # Test connection by checking if bucket exists
            self._client.head_bucket(Bucket=self._bucket)

            self._authenticated = True
            endpoint = self._endpoint_url or "AWS S3"
            logger.info(f"S3 authenticated: {endpoint}, bucket: {self._bucket}")
            return True

        except NoCredentialsError:
            logger.error("S3 authentication failed: No credentials found")
            return False
        except ClientError as e:
            error_code = e.response.get("Error", {}).get("Code", "Unknown")
            logger.error(f"S3 authentication failed: {error_code}")
            return False
        except Exception as e:
            logger.error(f"S3 authentication failed: {e}")
            return False

    def _ensure_authenticated(self) -> None:
        """Ensure we're authenticated before operations."""
        if not self._authenticated and not self.authenticate():
            raise RuntimeError("S3 authentication required")

    def _get_object_content(self, key: str) -> Optional[tuple]:
        """Get object content and metadata.

        Returns:
            Tuple of (content_str, metadata_dict) or None
        """
        try:
            response = self._client.get_object(Bucket=self._bucket, Key=key)

            content_type = response.get("ContentType", "application/octet-stream")
            body = response["Body"].read()

            # Handle different content types
            if content_type.startswith("text/") or self._is_text_file(key):
                # Text file - decode
                try:
                    content = body.decode("utf-8")
                except UnicodeDecodeError:
                    content = body.decode("latin-1")
            elif content_type == "application/json":
                content = body.decode("utf-8")
            elif content_type == "application/pdf":
                content = self._extract_text_from_pdf(body)
            else:
                # Binary file - base64 encode or skip
                ext = Path(key).suffix.lower()
                if ext in self.TEXT_EXTENSIONS:
                    try:
                        content = body.decode("utf-8")
                    except UnicodeDecodeError:
                        logger.debug(f"Skipping binary file: {key}")
                        return None
                else:
                    logger.debug(f"Skipping non-text file: {key}")
                    return None

            metadata = {
                "content_type": content_type,
                "content_length": response.get("ContentLength", 0),
                "etag": response.get("ETag", "").strip('"'),
                "last_modified": response.get("LastModified"),
            }

            # Include S3 user metadata
            if self._include_metadata and "Metadata" in response:
                metadata["s3_metadata"] = response["Metadata"]

            return content, metadata

        except ClientError as e:
            logger.error(f"Failed to get object {key}: {e}")
            return None

    def _is_text_file(self, key: str) -> bool:
        """Check if file is likely a text file based on extension."""
        ext = Path(key).suffix.lower()
        return ext in self.TEXT_EXTENSIONS

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by S3 key.

        Args:
            doc_id: S3 object key (e.g., "reports/2024/q1.pdf")
        """
        self._ensure_authenticated()

        try:
            result = self._get_object_content(doc_id)
            if result is None:
                return None

            content, metadata = result

            # Get object info for timestamps
            head = self._client.head_object(Bucket=self._bucket, Key=doc_id)

            return LoadedDocument(
                content=content,
                metadata={"bucket": self._bucket, "key": doc_id, **metadata},
                source="s3",
                source_id=f"s3://{self._bucket}/{doc_id}",
                filename=Path(doc_id).name,
                mime_type=metadata.get("content_type", "text/plain"),
                created_at=None,  # S3 doesn't track creation time
                modified_at=head.get("LastModified"),
                size_bytes=metadata.get("content_length", 0),
            )

        except Exception as e:
            logger.error(f"Failed to load S3 document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from an S3 prefix (folder).

        Args:
            folder_path: S3 prefix (e.g., "reports/2024/")
            recursive: Include nested prefixes
        """
        self._ensure_authenticated()

        documents = []
        prefix = folder_path.strip("/")
        if prefix:
            prefix = f"{prefix}/"

        try:
            paginator = self._client.get_paginator("list_objects_v2")

            for page in paginator.paginate(Bucket=self._bucket, Prefix=prefix):
                for obj in page.get("Contents", []):
                    key = obj["Key"]

                    # Skip "folder" markers
                    if key.endswith("/"):
                        continue

                    # Skip non-recursive nested items
                    if not recursive:
                        remaining = key[len(prefix) :]
                        if "/" in remaining:
                            continue

                    doc = self.load_document(key)
                    if doc:
                        documents.append(doc)

            logger.info(
                f"Loaded {len(documents)} documents from s3://{self._bucket}/{prefix}"
            )
            return documents

        except Exception as e:
            logger.error(f"Failed to load S3 folder {folder_path}: {e}")
            return []

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents by filename pattern.

        Note: S3 doesn't support content search. This searches by key name.
        For full-text search, use Elasticsearch or similar.

        Args:
            query: Search pattern (matched against object keys)
            max_results: Maximum results to return
        """
        self._ensure_authenticated()

        documents = []
        query_lower = query.lower()
        prefix = self._prefix + "/" if self._prefix else ""

        try:
            paginator = self._client.get_paginator("list_objects_v2")

            for page in paginator.paginate(Bucket=self._bucket, Prefix=prefix):
                for obj in page.get("Contents", []):
                    key = obj["Key"]

                    if query_lower in key.lower():
                        doc = self.load_document(key)
                        if doc:
                            documents.append(doc)
                            if len(documents) >= max_results:
                                break

                if len(documents) >= max_results:
                    break

            logger.info(f"Found {len(documents)} documents matching '{query}'")
            return documents

        except Exception as e:
            logger.error(f"S3 search failed: {e}")
            return []

    def list_folders(self, prefix: str = "") -> list[dict[str, str]]:
        """List 'folders' (common prefixes) in a prefix.

        Args:
            prefix: Parent prefix to list
        """
        self._ensure_authenticated()

        folders = []
        prefix = prefix.strip("/")
        if prefix:
            prefix = f"{prefix}/"

        try:
            response = self._client.list_objects_v2(
                Bucket=self._bucket, Prefix=prefix, Delimiter="/"
            )

            for cp in response.get("CommonPrefixes", []):
                folder_path = cp["Prefix"].rstrip("/")
                folder_name = folder_path.split("/")[-1]
                folders.append({"name": folder_name, "path": folder_path})

            return folders

        except Exception as e:
            logger.error(f"Failed to list S3 folders: {e}")
            return []

    def upload_document(
        self,
        key: str,
        content: str,
        content_type: str = "text/plain",
        metadata: Optional[dict[str, str]] = None,
    ) -> bool:
        """Upload a document to S3.

        Args:
            key: S3 object key
            content: Document content
            content_type: MIME type
            metadata: Optional S3 user metadata
        """
        self._ensure_authenticated()

        try:
            put_kwargs = {
                "Bucket": self._bucket,
                "Key": key,
                "Body": content.encode("utf-8"),
                "ContentType": content_type,
            }

            if metadata:
                put_kwargs["Metadata"] = metadata

            self._client.put_object(**put_kwargs)
            logger.info(f"Uploaded document to s3://{self._bucket}/{key}")
            return True

        except Exception as e:
            logger.error(f"Failed to upload to S3: {e}")
            return False

    def delete_document(self, key: str) -> bool:
        """Delete a document from S3."""
        self._ensure_authenticated()

        try:
            self._client.delete_object(Bucket=self._bucket, Key=key)
            logger.info(f"Deleted s3://{self._bucket}/{key}")
            return True
        except Exception as e:
            logger.error(f"Failed to delete S3 object: {e}")
            return False


# ============================================================================
# MONGODB DOCUMENT LOADER
# ============================================================================

# Check for pymongo
try:
    from pymongo import MongoClient
    from pymongo.errors import ConnectionFailure, OperationFailure

    PYMONGO_AVAILABLE = True
except ImportError:
    PYMONGO_AVAILABLE = False


class MongoDBLoader(BaseLoader):
    """Load documents from MongoDB.

    MongoDB is a document database storing JSON-like documents.
    Perfect for RAG as documents are already structured.

    Authentication options:
        1. Connection URI (mongodb://user:pass@host:port/db)
        2. Separate host/port/credentials

    Example:
        # With connection URI
        loader = MongoDBLoader(
            uri="mongodb://localhost:27017",
            database="knowledge_base",
            collection="documents"
        )
        docs = loader.load_collection()

        # With authentication
        loader = MongoDBLoader(
            host="mongodb.company.com",
            port=27017,
            database="docs",
            collection="articles",
            username="reader",
            password="secret",
            auth_source="admin"
        )
        docs = loader.load_collection(filter={"status": "published"})
    """

    def __init__(
        self,
        uri: Optional[str] = None,
        host: str = "localhost",
        port: int = 27017,
        database: str = "test",
        collection: str = "documents",
        username: Optional[str] = None,
        password: Optional[str] = None,
        auth_source: str = "admin",
        content_field: str = "content",
        title_field: str = "title",
        metadata_fields: Optional[list[str]] = None,
    ):
        """Initialize MongoDB loader.

        Args:
            uri: MongoDB connection URI (overrides host/port/auth)
            host: MongoDB host
            port: MongoDB port
            database: Database name
            collection: Default collection name
            username: Auth username
            password: Auth password
            auth_source: Auth database
            content_field: Field containing document content
            title_field: Field containing document title
            metadata_fields: Additional fields to include in metadata
        """
        if not PYMONGO_AVAILABLE:
            raise ImportError(
                "pymongo not available. Install with: pip install pymongo"
            )

        self._uri = uri
        self._host = host
        self._port = port
        self._database_name = database
        self._collection_name = collection
        self._username = username
        self._password = password
        self._auth_source = auth_source
        self._content_field = content_field
        self._title_field = title_field
        self._metadata_fields = metadata_fields or []

        self._client: Optional[Any] = None
        self._db: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "mongodb"

    def authenticate(self) -> bool:
        """Connect to MongoDB."""
        if self._authenticated and self._client is not None:
            return True

        try:
            if self._uri:
                self._client = MongoClient(self._uri)
            else:
                # Build connection kwargs
                conn_kwargs = {
                    "host": self._host,
                    "port": self._port,
                }

                if self._username and self._password:
                    conn_kwargs["username"] = self._username
                    conn_kwargs["password"] = self._password
                    conn_kwargs["authSource"] = self._auth_source

                self._client = MongoClient(**conn_kwargs)

            # Test connection
            self._client.admin.command("ping")

            self._db = self._client[self._database_name]
            self._authenticated = True

            logger.info(f"MongoDB connected: {self._database_name}")
            return True

        except ConnectionFailure as e:
            logger.error(f"MongoDB connection failed: {e}")
            return False
        except Exception as e:
            logger.error(f"MongoDB authentication failed: {e}")
            return False

    def _ensure_authenticated(self) -> None:
        """Ensure we're connected before operations."""
        if not self._authenticated and not self.authenticate():
            raise RuntimeError("MongoDB connection required")

    def _doc_to_loaded_document(
        self, doc: dict[str, Any], collection_name: str
    ) -> Optional[LoadedDocument]:
        """Convert MongoDB document to LoadedDocument."""
        try:
            # Get content
            content = doc.get(self._content_field, "")
            if not content:
                # Try fallbacks
                for fallback in ["text", "body", "description", "data"]:
                    if fallback in doc:
                        content = str(doc[fallback])
                        break

            if not content:
                # Use entire document as JSON
                # Remove _id for cleaner output (it's in metadata)
                doc_copy = {k: v for k, v in doc.items() if k != "_id"}
                content = json.dumps(doc_copy, indent=2, default=str)

            # Get title
            title = doc.get(self._title_field, "")
            if not title:
                for fallback in ["name", "filename", "subject", "heading"]:
                    if fallback in doc:
                        title = str(doc[fallback])
                        break
            if not title:
                title = str(doc.get("_id", "Untitled"))

            # Build metadata
            doc_id = doc.get("_id")
            metadata = {
                "database": self._database_name,
                "collection": collection_name,
                "document_id": str(doc_id) if doc_id else None,
            }

            # Add specified metadata fields
            for field in self._metadata_fields:
                if field in doc and field not in [self._content_field, "_id"]:
                    try:
                        json.dumps(doc[field], default=str)
                        metadata[field] = doc[field]
                    except (TypeError, ValueError):
                        metadata[field] = str(doc[field])

            # Add all non-content fields
            for key, value in doc.items():
                if key not in [self._content_field, self._title_field, "_id"]:
                    if key not in metadata:
                        try:
                            json.dumps(value, default=str)
                            metadata[key] = value
                        except (TypeError, ValueError):
                            metadata[key] = str(value)

            # Handle timestamps
            created_at = None
            modified_at = None

            for ts_field in ["created_at", "createdAt", "timestamp", "created"]:
                if ts_field in doc:
                    ts = doc[ts_field]
                    if isinstance(ts, datetime):
                        created_at = ts
                    break

            for ts_field in ["updated_at", "updatedAt", "modified", "modified_at"]:
                if ts_field in doc:
                    ts = doc[ts_field]
                    if isinstance(ts, datetime):
                        modified_at = ts
                    break

            return LoadedDocument(
                content=str(content),
                metadata=metadata,
                source="mongodb",
                source_id=f"{self._database_name}/{collection_name}/{doc_id}",
                filename=title,
                mime_type="application/json",
                created_at=created_at,
                modified_at=modified_at,
                size_bytes=len(str(content).encode("utf-8")),
            )

        except Exception as e:
            logger.error(f"Failed to convert MongoDB document: {e}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by ID.

        Args:
            doc_id: Document _id (string or ObjectId format)
        """
        self._ensure_authenticated()

        try:
            from bson import ObjectId

            collection = self._db[self._collection_name]

            # Try as ObjectId first, then as string
            try:
                doc = collection.find_one({"_id": ObjectId(doc_id)})
            except Exception:
                doc = collection.find_one({"_id": doc_id})

            if not doc:
                logger.warning(f"Document not found: {doc_id}")
                return None

            return self._doc_to_loaded_document(doc, self._collection_name)

        except Exception as e:
            logger.error(f"Failed to load MongoDB document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a collection.

        In MongoDB, 'folder' = collection.

        Args:
            folder_path: Collection name
            recursive: Ignored (MongoDB doesn't have nested collections)
        """
        return self.load_collection(collection_name=folder_path)

    def load_collection(
        self,
        collection_name: Optional[str] = None,
        filter: Optional[dict[str, Any]] = None,
        projection: Optional[dict[str, int]] = None,
        sort: Optional[list[tuple]] = None,
        limit: int = 0,
        skip: int = 0,
    ) -> list[LoadedDocument]:
        """Load documents from a MongoDB collection.

        Args:
            collection_name: Collection to query (uses default if None)
            filter: MongoDB query filter (e.g., {"status": "published"})
            projection: Fields to include/exclude
            sort: Sort order (e.g., [("created_at", -1)])
            limit: Maximum documents (0 = no limit)
            skip: Documents to skip (for pagination)

        Example:
            docs = loader.load_collection(
                filter={"category": "tech", "views": {"$gt": 100}},
                sort=[("created_at", -1)],
                limit=50
            )
        """
        self._ensure_authenticated()

        collection_name = collection_name or self._collection_name
        documents = []

        try:
            collection = self._db[collection_name]

            cursor = collection.find(filter=filter or {}, projection=projection)

            if sort:
                cursor = cursor.sort(sort)
            if skip:
                cursor = cursor.skip(skip)
            if limit:
                cursor = cursor.limit(limit)

            for doc in cursor:
                loaded_doc = self._doc_to_loaded_document(doc, collection_name)
                if loaded_doc:
                    documents.append(loaded_doc)

            logger.info(f"Loaded {len(documents)} documents from {collection_name}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load collection {collection_name}: {e}")
            return []

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents using MongoDB text search.

        Requires a text index on the collection.
        Falls back to regex search if no text index.

        Args:
            query: Search text
            max_results: Maximum results
        """
        self._ensure_authenticated()

        documents = []
        collection = self._db[self._collection_name]

        try:
            # Try text search first (requires text index)
            try:
                cursor = (
                    collection.find(
                        {"$text": {"$search": query}}, {"score": {"$meta": "textScore"}}
                    )
                    .sort([("score", {"$meta": "textScore"})])
                    .limit(max_results)
                )

                for doc in cursor:
                    loaded_doc = self._doc_to_loaded_document(
                        doc, self._collection_name
                    )
                    if loaded_doc:
                        documents.append(loaded_doc)

                if documents:
                    logger.info(f"Found {len(documents)} documents via text search")
                    return documents

            except OperationFailure:
                # No text index, fall back to regex
                pass

            # Fallback: regex search on content and title fields
            regex_filter = {
                "$or": [
                    {self._content_field: {"$regex": query, "$options": "i"}},
                    {self._title_field: {"$regex": query, "$options": "i"}},
                ]
            }

            cursor = collection.find(regex_filter).limit(max_results)

            for doc in cursor:
                loaded_doc = self._doc_to_loaded_document(doc, self._collection_name)
                if loaded_doc:
                    documents.append(loaded_doc)

            logger.info(f"Found {len(documents)} documents via regex search")
            return documents

        except Exception as e:
            logger.error(f"MongoDB search failed: {e}")
            return []

    def save_document(
        self,
        content: str,
        title: str = "",
        metadata: Optional[dict[str, Any]] = None,
        collection_name: Optional[str] = None,
        document_id: Optional[str] = None,
    ) -> Optional[str]:
        """Save a document to MongoDB.

        Args:
            content: Document content
            title: Document title
            metadata: Additional fields
            collection_name: Target collection (uses default if None)
            document_id: Specific _id (auto-generated if None)

        Returns:
            Document ID if successful
        """
        self._ensure_authenticated()

        collection_name = collection_name or self._collection_name

        try:
            doc = {
                self._content_field: content,
                self._title_field: title or "Untitled",
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }

            if metadata:
                doc.update(metadata)

            collection = self._db[collection_name]

            if document_id:
                from bson import ObjectId

                try:
                    doc["_id"] = ObjectId(document_id)
                except Exception:
                    doc["_id"] = document_id
                collection.replace_one({"_id": doc["_id"]}, doc, upsert=True)
                return str(doc["_id"])
            else:
                result = collection.insert_one(doc)
                return str(result.inserted_id)

        except Exception as e:
            logger.error(f"Failed to save MongoDB document: {e}")
            return None

    def delete_document(
        self, doc_id: str, collection_name: Optional[str] = None
    ) -> bool:
        """Delete a document from MongoDB."""
        self._ensure_authenticated()

        collection_name = collection_name or self._collection_name

        try:
            from bson import ObjectId

            collection = self._db[collection_name]

            # Try as ObjectId first
            try:
                result = collection.delete_one({"_id": ObjectId(doc_id)})
            except Exception:
                result = collection.delete_one({"_id": doc_id})

            if result.deleted_count > 0:
                logger.info(f"Deleted MongoDB document: {doc_id}")
                return True
            else:
                logger.warning(f"Document not found: {doc_id}")
                return False

        except Exception as e:
            logger.error(f"Failed to delete MongoDB document: {e}")
            return False

    def list_collections(self) -> list[str]:
        """List all collections in the database."""
        self._ensure_authenticated()

        try:
            return self._db.list_collection_names()
        except Exception as e:
            logger.error(f"Failed to list collections: {e}")
            return []

    def create_text_index(
        self, fields: Optional[list[str]] = None, collection_name: Optional[str] = None
    ) -> bool:
        """Create a text index for full-text search.

        Args:
            fields: Fields to index (defaults to content and title)
            collection_name: Target collection
        """
        self._ensure_authenticated()

        collection_name = collection_name or self._collection_name
        fields = fields or [self._content_field, self._title_field]

        try:
            collection = self._db[collection_name]

            # Build index specification
            index_spec = [(field, "text") for field in fields]

            collection.create_index(index_spec, name="text_search_index")
            logger.info(f"Created text index on {collection_name}: {fields}")
            return True

        except Exception as e:
            logger.error(f"Failed to create text index: {e}")
            return False

    def aggregate(
        self, pipeline: list[dict[str, Any]], collection_name: Optional[str] = None
    ) -> list[LoadedDocument]:
        """Run aggregation pipeline and return as documents.

        Args:
            pipeline: MongoDB aggregation pipeline
            collection_name: Target collection

        Example:
            docs = loader.aggregate([
                {"$match": {"status": "published"}},
                {"$group": {"_id": "$category", "count": {"$sum": 1}}},
                {"$sort": {"count": -1}}
            ])
        """
        self._ensure_authenticated()

        collection_name = collection_name or self._collection_name
        documents = []

        try:
            collection = self._db[collection_name]

            for doc in collection.aggregate(pipeline):
                loaded_doc = self._doc_to_loaded_document(doc, collection_name)
                if loaded_doc:
                    documents.append(loaded_doc)

            return documents

        except Exception as e:
            logger.error(f"Aggregation failed: {e}")
            return []


# ============================================================================
# GITHUB DOCUMENT LOADER
# ============================================================================

# Check for PyGithub
try:
    from github import Github, GithubException
    from github.ContentFile import ContentFile
    from github.Repository import Repository

    PYGITHUB_AVAILABLE = True
except ImportError:
    PYGITHUB_AVAILABLE = False


class GitHubLoader(BaseLoader):
    """Load documents from GitHub repositories.

    Loads code, documentation, issues, PRs, and wikis for RAG.
    Perfect for building code-aware assistants.

    Authentication options:
        1. Personal Access Token (recommended)
        2. GitHub App installation token
        3. No auth (public repos only, rate limited)

    Example:
        # Load repository code and docs
        loader = GitHubLoader(token="ghp_xxxxx")
        docs = loader.load_repository("owner/repo")

        # Load issues for context
        issues = loader.load_issues("owner/repo", state="open")

        # Load specific folder
        code = loader.load_folder("owner/repo", path="src/")

        # Search across repos
        results = loader.search_code("def authenticate", language="python")
    """

    # File extensions to include as text
    CODE_EXTENSIONS = {
        ".py",
        ".js",
        ".ts",
        ".jsx",
        ".tsx",
        ".java",
        ".cpp",
        ".c",
        ".h",
        ".go",
        ".rs",
        ".rb",
        ".php",
        ".cs",
        ".swift",
        ".kt",
        ".scala",
        ".sh",
        ".bash",
        ".zsh",
        ".ps1",
        ".sql",
        ".r",
        ".m",
        ".mm",
    }

    DOC_EXTENSIONS = {
        ".md",
        ".markdown",
        ".rst",
        ".txt",
        ".adoc",
        ".asciidoc",
        ".html",
        ".htm",
        ".xml",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".ini",
        ".cfg",
        ".conf",
        ".env",
        ".gitignore",
        ".dockerignore",
    }

    def __init__(
        self,
        token: Optional[str] = None,
        base_url: Optional[str] = None,
        include_code: bool = True,
        include_docs: bool = True,
        max_file_size: int = 1_000_000,  # 1MB
    ):
        """Initialize GitHub loader.

        Args:
            token: GitHub personal access token or app token
            base_url: GitHub Enterprise base URL (e.g., "https://github.company.com/api/v3")
            include_code: Include source code files
            include_docs: Include documentation files
            max_file_size: Maximum file size to load (bytes)
        """
        if not PYGITHUB_AVAILABLE:
            raise ImportError(
                "PyGithub not available. Install with: pip install PyGithub"
            )

        self._token = token or os.environ.get("GITHUB_TOKEN")
        self._base_url = base_url
        self._include_code = include_code
        self._include_docs = include_docs
        self._max_file_size = max_file_size
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "github"

    def authenticate(self) -> bool:
        """Initialize GitHub client."""
        if self._authenticated and self._client is not None:
            return True

        try:
            if self._base_url:
                # GitHub Enterprise
                self._client = Github(
                    login_or_token=self._token, base_url=self._base_url
                )
            else:
                # GitHub.com
                self._client = Github(login_or_token=self._token)

            # Test connection
            user = self._client.get_user()
            logger.info(
                f"GitHub authenticated as: {user.login if self._token else 'anonymous'}"
            )
            self._authenticated = True
            return True

        except GithubException as e:
            logger.error(f"GitHub authentication failed: {e}")
            return False
        except Exception as e:
            logger.error(f"GitHub connection failed: {e}")
            return False

    def _ensure_authenticated(self) -> None:
        """Ensure we're authenticated."""
        if not self._authenticated and not self.authenticate():
            raise RuntimeError("GitHub authentication required")

    def _should_include_file(self, path: str, size: int) -> bool:
        """Check if file should be included."""
        if size > self._max_file_size:
            return False

        ext = Path(path).suffix.lower()

        if self._include_code and ext in self.CODE_EXTENSIONS:
            return True
        if self._include_docs and ext in self.DOC_EXTENSIONS:
            return True

        # Include common important files
        filename = Path(path).name.lower()
        return filename in {
            "readme",
            "license",
            "changelog",
            "contributing",
            "makefile",
            "dockerfile",
        }

    def _content_to_document(
        self, content: "ContentFile", repo_name: str
    ) -> Optional[LoadedDocument]:
        """Convert GitHub ContentFile to LoadedDocument."""
        try:
            if content.type != "file":
                return None

            if not self._should_include_file(content.path, content.size):
                return None

            # Get decoded content
            try:
                text_content = content.decoded_content.decode("utf-8")
            except UnicodeDecodeError:
                logger.debug(f"Skipping binary file: {content.path}")
                return None

            ext = Path(content.path).suffix.lower()
            mime_type = "text/plain"
            if ext in {".json"}:
                mime_type = "application/json"
            elif ext in {".md", ".markdown"}:
                mime_type = "text/markdown"
            elif ext in {".html", ".htm"}:
                mime_type = "text/html"
            elif ext in self.CODE_EXTENSIONS:
                mime_type = f"text/x-{ext[1:]}"

            return LoadedDocument(
                content=text_content,
                metadata={
                    "repository": repo_name,
                    "path": content.path,
                    "sha": content.sha,
                    "url": content.html_url,
                    "type": "code" if ext in self.CODE_EXTENSIONS else "documentation",
                },
                source="github",
                source_id=f"github:{repo_name}:{content.path}",
                filename=content.name,
                mime_type=mime_type,
                size_bytes=content.size,
            )

        except Exception as e:
            logger.error(f"Failed to convert GitHub content: {e}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single file from a repository.

        Args:
            doc_id: Format "owner/repo:path/to/file.py"
        """
        self._ensure_authenticated()

        try:
            if ":" not in doc_id:
                raise ValueError(
                    f"Invalid doc_id format: {doc_id}. Use 'owner/repo:path/to/file'"
                )

            repo_name, file_path = doc_id.split(":", 1)
            repo = self._client.get_repo(repo_name)
            content = repo.get_contents(file_path)

            if isinstance(content, list):
                logger.warning(f"Path is a directory: {file_path}")
                return None

            return self._content_to_document(content, repo_name)

        except GithubException as e:
            logger.error(f"GitHub API error: {e}")
            return None
        except Exception as e:
            logger.error(f"Failed to load GitHub document: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load files from a repository folder.

        Args:
            folder_path: Format "owner/repo" or "owner/repo:path/to/folder"
            recursive: Include subdirectories
        """
        self._ensure_authenticated()

        documents = []

        try:
            if ":" in folder_path:
                repo_name, path = folder_path.split(":", 1)
            else:
                repo_name = folder_path
                path = ""

            repo = self._client.get_repo(repo_name)

            def process_contents(contents, current_path=""):
                for content in contents:
                    if content.type == "file":
                        doc = self._content_to_document(content, repo_name)
                        if doc:
                            documents.append(doc)
                    elif content.type == "dir" and recursive:
                        try:
                            subcontents = repo.get_contents(content.path)
                            if isinstance(subcontents, list):
                                process_contents(subcontents, content.path)
                        except GithubException:
                            pass  # Skip inaccessible directories

            contents = repo.get_contents(path)
            if isinstance(contents, list):
                process_contents(contents, path)
            else:
                doc = self._content_to_document(contents, repo_name)
                if doc:
                    documents.append(doc)

            logger.info(f"Loaded {len(documents)} files from {repo_name}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load GitHub folder: {e}")
            return []

    def load_repository(
        self,
        repo_name: str,
        branch: Optional[str] = None,
        paths: Optional[list[str]] = None,
    ) -> list[LoadedDocument]:
        """Load entire repository or specific paths.

        Args:
            repo_name: "owner/repo"
            branch: Branch name (default: default branch)
            paths: Specific paths to load (loads all if None)
        """
        self._ensure_authenticated()

        documents = []

        try:
            repo = self._client.get_repo(repo_name)

            if paths:
                for path in paths:
                    docs = self.load_folder(f"{repo_name}:{path}")
                    documents.extend(docs)
            else:
                documents = self.load_folder(repo_name)

            # Also load README if not already included
            try:
                readme = repo.get_readme()
                readme_doc = self._content_to_document(readme, repo_name)
                if readme_doc and readme_doc.source_id not in [
                    d.source_id for d in documents
                ]:
                    documents.insert(0, readme_doc)
            except GithubException:
                pass

            return documents

        except Exception as e:
            logger.error(f"Failed to load repository: {e}")
            return []

    def load_issues(
        self,
        repo_name: str,
        state: str = "open",
        labels: Optional[list[str]] = None,
        limit: int = 100,
    ) -> list[LoadedDocument]:
        """Load issues from a repository.

        Args:
            repo_name: "owner/repo"
            state: "open", "closed", or "all"
            labels: Filter by labels
            limit: Maximum issues to load
        """
        self._ensure_authenticated()

        documents = []

        try:
            repo = self._client.get_repo(repo_name)

            issues = repo.get_issues(state=state, labels=labels or [])

            count = 0
            for issue in issues:
                if count >= limit:
                    break

                # Skip pull requests (they show up in issues API)
                if issue.pull_request:
                    continue

                # Build content from issue body and comments
                content_parts = [
                    f"# {issue.title}",
                    f"\n**State:** {issue.state}",
                    f"**Author:** {issue.user.login}",
                    f"**Created:** {issue.created_at}",
                ]

                if issue.labels:
                    content_parts.append(
                        f"**Labels:** {', '.join(l.name for l in issue.labels)}"
                    )

                if issue.body:
                    content_parts.append(f"\n## Description\n{issue.body}")

                # Include top comments
                comments = list(issue.get_comments()[:10])
                if comments:
                    content_parts.append("\n## Comments")
                    for comment in comments:
                        content_parts.append(
                            f"\n**{comment.user.login}** ({comment.created_at}):\n{comment.body}"
                        )

                doc = LoadedDocument(
                    content="\n".join(content_parts),
                    metadata={
                        "repository": repo_name,
                        "issue_number": issue.number,
                        "state": issue.state,
                        "author": issue.user.login,
                        "labels": [l.name for l in issue.labels],
                        "url": issue.html_url,
                        "comments_count": issue.comments,
                    },
                    source="github",
                    source_id=f"github:{repo_name}:issue:{issue.number}",
                    filename=f"Issue #{issue.number}: {issue.title}",
                    mime_type="text/markdown",
                    created_at=issue.created_at,
                    modified_at=issue.updated_at,
                )
                documents.append(doc)
                count += 1

            logger.info(f"Loaded {len(documents)} issues from {repo_name}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load issues: {e}")
            return []

    def load_pull_requests(
        self, repo_name: str, state: str = "open", limit: int = 50
    ) -> list[LoadedDocument]:
        """Load pull requests from a repository.

        Args:
            repo_name: "owner/repo"
            state: "open", "closed", or "all"
            limit: Maximum PRs to load
        """
        self._ensure_authenticated()

        documents = []

        try:
            repo = self._client.get_repo(repo_name)
            prs = repo.get_pulls(state=state)

            count = 0
            for pr in prs:
                if count >= limit:
                    break

                content_parts = [
                    f"# PR #{pr.number}: {pr.title}",
                    f"\n**State:** {pr.state}",
                    f"**Author:** {pr.user.login}",
                    f"**Branch:** {pr.head.ref} → {pr.base.ref}",
                    f"**Created:** {pr.created_at}",
                ]

                if pr.body:
                    content_parts.append(f"\n## Description\n{pr.body}")

                # Include changed files list
                try:
                    files = list(pr.get_files()[:20])
                    if files:
                        content_parts.append("\n## Changed Files")
                        for f in files:
                            content_parts.append(
                                f"- {f.filename} (+{f.additions}/-{f.deletions})"
                            )
                except GithubException:
                    pass

                doc = LoadedDocument(
                    content="\n".join(content_parts),
                    metadata={
                        "repository": repo_name,
                        "pr_number": pr.number,
                        "state": pr.state,
                        "author": pr.user.login,
                        "head_branch": pr.head.ref,
                        "base_branch": pr.base.ref,
                        "url": pr.html_url,
                        "mergeable": pr.mergeable,
                    },
                    source="github",
                    source_id=f"github:{repo_name}:pr:{pr.number}",
                    filename=f"PR #{pr.number}: {pr.title}",
                    mime_type="text/markdown",
                    created_at=pr.created_at,
                    modified_at=pr.updated_at,
                )
                documents.append(doc)
                count += 1

            logger.info(f"Loaded {len(documents)} PRs from {repo_name}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load PRs: {e}")
            return []

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for code across GitHub.

        Args:
            query: Search query (GitHub code search syntax)
            max_results: Maximum results
        """
        self._ensure_authenticated()

        documents = []

        try:
            results = self._client.search_code(query)

            count = 0
            for result in results:
                if count >= max_results:
                    break

                try:
                    doc = self._content_to_document(result, result.repository.full_name)
                    if doc:
                        documents.append(doc)
                        count += 1
                except Exception:
                    continue

            logger.info(f"Found {len(documents)} results for '{query}'")
            return documents

        except Exception as e:
            logger.error(f"GitHub search failed: {e}")
            return []

    def search_repositories(
        self, query: str, max_results: int = 10
    ) -> list[dict[str, Any]]:
        """Search for repositories.

        Returns repository info, not LoadedDocuments.
        Use load_repository() to load content.
        """
        self._ensure_authenticated()

        repos = []

        try:
            results = self._client.search_repositories(query)

            count = 0
            for repo in results:
                if count >= max_results:
                    break

                repos.append(
                    {
                        "full_name": repo.full_name,
                        "description": repo.description,
                        "stars": repo.stargazers_count,
                        "language": repo.language,
                        "url": repo.html_url,
                        "default_branch": repo.default_branch,
                    }
                )
                count += 1

            return repos

        except Exception as e:
            logger.error(f"Repository search failed: {e}")
            return []


# ============================================================================
# MICROSOFT 365 DOCUMENT LOADER
# ============================================================================

# Check for Microsoft Graph SDK
try:
    import requests
    from msal import ConfidentialClientApplication, PublicClientApplication

    MSAL_AVAILABLE = True
except ImportError:
    MSAL_AVAILABLE = False


class Microsoft365Loader(BaseLoader):
    """Load documents from Microsoft 365 (OneDrive, Outlook, SharePoint).

    Uses Microsoft Graph API to access all Microsoft 365 services:
    - OneDrive: Personal and business file storage
    - Outlook: Emails and attachments
    - SharePoint: Team sites and document libraries
    - OneNote: Notebooks and pages

    Authentication options:
        1. Client credentials (app-only, for automation)
        2. Device code flow (interactive, for users)
        3. Existing access token

    Example:
        # App-only authentication (automation)
        loader = Microsoft365Loader(
            client_id="your-app-id",
            client_secret="your-secret",
            tenant_id="your-tenant-id"
        )

        # Load OneDrive files
        docs = loader.load_onedrive_folder("Documents/Projects")

        # Load recent emails
        emails = loader.load_outlook_emails(days=7, query="from:boss@company.com")

        # Load SharePoint site
        sharepoint = loader.load_sharepoint_library("sites/Engineering", "Shared Documents")
    """

    GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"

    # Scopes for different services
    SCOPES_DELEGATED = [
        "Files.Read.All",
        "Mail.Read",
        "Sites.Read.All",
        "Notes.Read.All",
        "User.Read",
    ]

    SCOPES_APPLICATION = ["https://graph.microsoft.com/.default"]

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        tenant_id: Optional[str] = None,
        access_token: Optional[str] = None,
        user_id: Optional[str] = None,
        use_device_flow: bool = False,
    ):
        """Initialize Microsoft 365 loader.

        Args:
            client_id: Azure AD application (client) ID
            client_secret: Client secret (for app-only auth)
            tenant_id: Azure AD tenant ID
            access_token: Pre-acquired access token
            user_id: User ID for delegated access (default: "me")
            use_device_flow: Use device code flow for interactive auth
        """
        if not MSAL_AVAILABLE:
            raise ImportError(
                "MSAL not available. Install with: pip install msal requests"
            )

        self._client_id = client_id or os.environ.get("MICROSOFT_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get("MICROSOFT_CLIENT_SECRET")
        self._tenant_id = tenant_id or os.environ.get("MICROSOFT_TENANT_ID")
        self._access_token = access_token
        self._user_id = user_id or "me"
        self._use_device_flow = use_device_flow

        self._app: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "microsoft365"

    def authenticate(self) -> bool:
        """Authenticate with Microsoft Graph."""
        if self._authenticated and self._access_token:
            return True

        try:
            if self._access_token:
                # Use provided token
                self._authenticated = True
                return True

            if not self._client_id or not self._tenant_id:
                raise ValueError("client_id and tenant_id are required")

            authority = f"https://login.microsoftonline.com/{self._tenant_id}"

            if self._client_secret:
                # App-only authentication
                self._app = ConfidentialClientApplication(
                    self._client_id,
                    authority=authority,
                    client_credential=self._client_secret,
                )

                result = self._app.acquire_token_for_client(
                    scopes=self.SCOPES_APPLICATION
                )

            elif self._use_device_flow:
                # Device code flow (interactive)
                self._app = PublicClientApplication(
                    self._client_id, authority=authority
                )

                flow = self._app.initiate_device_flow(scopes=self.SCOPES_DELEGATED)

                if "user_code" not in flow:
                    raise ValueError(
                        f"Failed to create device flow: {flow.get('error_description')}"
                    )

                print(f"\nTo authenticate, visit: {flow['verification_uri']}")
                print(f"Enter code: {flow['user_code']}\n")

                result = self._app.acquire_token_by_device_flow(flow)
            else:
                raise ValueError(
                    "Either client_secret or use_device_flow must be provided"
                )

            if "access_token" in result:
                self._access_token = result["access_token"]
                self._authenticated = True
                logger.info("Microsoft 365 authenticated successfully")
                return True
            else:
                error = result.get(
                    "error_description", result.get("error", "Unknown error")
                )
                logger.error(f"Microsoft 365 authentication failed: {error}")
                return False

        except Exception as e:
            logger.error(f"Microsoft 365 authentication failed: {e}")
            return False

    def _ensure_authenticated(self) -> None:
        """Ensure we're authenticated."""
        if not self._authenticated and not self.authenticate():
            raise RuntimeError("Microsoft 365 authentication required")

    def _graph_request(
        self,
        endpoint: str,
        method: str = "GET",
        params: Optional[dict] = None,
        data: Optional[dict] = None,
    ) -> Optional[dict]:
        """Make a request to Microsoft Graph API."""
        url = f"{self.GRAPH_BASE_URL}/{endpoint.lstrip('/')}"

        headers = {
            "Authorization": f"Bearer {self._access_token}",
            "Content-Type": "application/json",
        }

        try:
            response = requests.request(
                method, url, headers=headers, params=params, json=data, timeout=30
            )
            response.raise_for_status()
            return response.json() if response.content else {}

        except requests.exceptions.HTTPError as e:
            logger.error(
                f"Graph API error: {e.response.status_code} - {e.response.text}"
            )
            return None
        except Exception as e:
            logger.error(f"Graph API request failed: {e}")
            return None

    def _download_file_content(self, download_url: str) -> Optional[bytes]:
        """Download file content from OneDrive/SharePoint."""
        try:
            response = requests.get(
                download_url,
                headers={"Authorization": f"Bearer {self._access_token}"},
                timeout=60,
            )
            response.raise_for_status()
            return response.content
        except Exception as e:
            logger.error(f"Failed to download file: {e}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by ID.

        Args:
            doc_id: Format "onedrive:item_id" or "sharepoint:site_id:item_id"
        """
        self._ensure_authenticated()

        try:
            parts = doc_id.split(":")

            if parts[0] == "onedrive":
                item_id = parts[1]
                endpoint = f"/me/drive/items/{item_id}"
            elif parts[0] == "sharepoint":
                site_id, item_id = parts[1], parts[2]
                endpoint = f"/sites/{site_id}/drive/items/{item_id}"
            else:
                raise ValueError(f"Invalid doc_id format: {doc_id}")

            item = self._graph_request(endpoint)
            if not item:
                return None

            # Get download URL
            download_url = item.get("@microsoft.graph.downloadUrl")
            if not download_url:
                # Try to get download URL
                download_item = self._graph_request(f"{endpoint}/content", method="GET")
                if not download_item:
                    return None

            content = self._download_file_content(download_url)
            if not content:
                return None

            # Decode content based on mime type
            mime_type = item.get("file", {}).get("mimeType", "application/octet-stream")

            if mime_type.startswith("text/") or mime_type == "application/json":
                text_content = content.decode("utf-8")
            elif mime_type == "application/pdf":
                text_content = self._extract_text_from_pdf(content)
            else:
                logger.debug(f"Skipping non-text file: {item.get('name')}")
                return None

            return LoadedDocument(
                content=text_content,
                metadata={
                    "item_id": item.get("id"),
                    "web_url": item.get("webUrl"),
                    "parent_path": item.get("parentReference", {}).get("path", ""),
                    "created_by": item.get("createdBy", {})
                    .get("user", {})
                    .get("displayName"),
                    "modified_by": item.get("lastModifiedBy", {})
                    .get("user", {})
                    .get("displayName"),
                },
                source="microsoft365",
                source_id=doc_id,
                filename=item.get("name", "Untitled"),
                mime_type=mime_type,
                created_at=(
                    datetime.fromisoformat(
                        item["createdDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("createdDateTime")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(
                        item["lastModifiedDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("lastModifiedDateTime")
                    else None
                ),
                size_bytes=item.get("size", 0),
            )

        except Exception as e:
            logger.error(f"Failed to load document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load files from OneDrive folder.

        Args:
            folder_path: OneDrive folder path (e.g., "Documents/Projects")
            recursive: Include subfolders
        """
        return self.load_onedrive_folder(folder_path, recursive)

    def load_onedrive_folder(
        self,
        folder_path: str = "",
        recursive: bool = True,
        file_types: Optional[list[str]] = None,
    ) -> list[LoadedDocument]:
        """Load files from OneDrive.

        Args:
            folder_path: Folder path (empty for root)
            recursive: Include subfolders
            file_types: File extensions to include (e.g., [".docx", ".pdf"])
        """
        self._ensure_authenticated()

        documents = []

        try:
            # Build endpoint
            if folder_path:
                endpoint = f"/{self._user_id}/drive/root:/{folder_path}:/children"
            else:
                endpoint = f"/{self._user_id}/drive/root/children"

            def process_items(ep: str):
                result = self._graph_request(ep)
                if not result:
                    return

                for item in result.get("value", []):
                    if "folder" in item:
                        # Directory
                        if recursive:
                            child_endpoint = (
                                f"/{self._user_id}/drive/items/{item['id']}/children"
                            )
                            process_items(child_endpoint)
                    elif "file" in item:
                        # File
                        if file_types:
                            ext = Path(item["name"]).suffix.lower()
                            if ext not in file_types:
                                continue

                        doc_id = f"onedrive:{item['id']}"
                        doc = self.load_document(doc_id)
                        if doc:
                            documents.append(doc)

                # Handle pagination
                next_link = result.get("@odata.nextLink")
                if next_link:
                    process_items(next_link.replace(self.GRAPH_BASE_URL, ""))

            process_items(endpoint)
            logger.info(f"Loaded {len(documents)} files from OneDrive")
            return documents

        except Exception as e:
            logger.error(f"Failed to load OneDrive folder: {e}")
            return []

    def load_outlook_emails(
        self,
        days: int = 7,
        query: Optional[str] = None,
        folder: str = "inbox",
        limit: int = 50,
        include_attachments: bool = True,
    ) -> list[LoadedDocument]:
        """Load emails from Outlook.

        Args:
            days: Load emails from last N days
            query: OData filter query (e.g., "from/emailAddress/address eq 'boss@company.com'")
            folder: Mail folder (inbox, sentitems, drafts, etc.)
            limit: Maximum emails to load
            include_attachments: Include attachment content
        """
        self._ensure_authenticated()

        documents = []

        try:
            # Build endpoint with filters
            endpoint = f"/{self._user_id}/mailFolders/{folder}/messages"

            params = {
                "$top": limit,
                "$orderby": "receivedDateTime desc",
                "$select": "id,subject,bodyPreview,body,from,toRecipients,receivedDateTime,hasAttachments",
            }

            # Add date filter
            from_date = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
            filter_parts = [f"receivedDateTime ge {from_date}"]

            if query:
                filter_parts.append(query)

            params["$filter"] = " and ".join(filter_parts)

            result = self._graph_request(endpoint, params=params)
            if not result:
                return []

            for email in result.get("value", []):
                # Build email content
                content_parts = [
                    f"# {email.get('subject', 'No Subject')}",
                    f"\n**From:** {email.get('from', {}).get('emailAddress', {}).get('address', 'Unknown')}",
                    f"**To:** {', '.join(r.get('emailAddress', {}).get('address', '') for r in email.get('toRecipients', []))}",
                    f"**Date:** {email.get('receivedDateTime', '')}",
                    "\n---\n",
                    email.get("body", {}).get("content", email.get("bodyPreview", "")),
                ]

                content = "\n".join(content_parts)

                # Clean HTML if needed
                if email.get("body", {}).get("contentType") == "html":
                    content = self._clean_html(content)

                metadata = {
                    "email_id": email.get("id"),
                    "from": email.get("from", {})
                    .get("emailAddress", {})
                    .get("address"),
                    "to": [
                        r.get("emailAddress", {}).get("address")
                        for r in email.get("toRecipients", [])
                    ],
                    "has_attachments": email.get("hasAttachments", False),
                    "folder": folder,
                }

                # Handle attachments
                if include_attachments and email.get("hasAttachments"):
                    attachments_result = self._graph_request(
                        f"/{self._user_id}/messages/{email['id']}/attachments"
                    )
                    if attachments_result:
                        for att in attachments_result.get("value", []):
                            if (
                                att.get("@odata.type")
                                == "#microsoft.graph.fileAttachment"
                            ):
                                att_content = base64.b64decode(
                                    att.get("contentBytes", "")
                                )
                                mime = att.get("contentType", "")

                                if (
                                    mime.startswith("text/")
                                    or mime == "application/json"
                                ):
                                    try:
                                        content += (
                                            f"\n\n--- Attachment: {att['name']} ---\n"
                                        )
                                        content += att_content.decode("utf-8")
                                    except UnicodeDecodeError:
                                        pass
                                elif mime == "application/pdf":
                                    pdf_text = self._extract_text_from_pdf(att_content)
                                    content += (
                                        f"\n\n--- Attachment: {att['name']} ---\n"
                                    )
                                    content += pdf_text

                doc = LoadedDocument(
                    content=content,
                    metadata=metadata,
                    source="microsoft365",
                    source_id=f"outlook:{email.get('id')}",
                    filename=email.get("subject", "No Subject"),
                    mime_type="text/plain",
                    created_at=(
                        datetime.fromisoformat(
                            email["receivedDateTime"].replace("Z", "+00:00")
                        )
                        if email.get("receivedDateTime")
                        else None
                    ),
                )
                documents.append(doc)

            logger.info(f"Loaded {len(documents)} emails from Outlook")
            return documents

        except Exception as e:
            logger.error(f"Failed to load Outlook emails: {e}")
            return []

    def load_sharepoint_library(
        self,
        site_path: str,
        library_name: str = "Shared Documents",
        folder_path: str = "",
        recursive: bool = True,
    ) -> list[LoadedDocument]:
        """Load documents from SharePoint document library.

        Args:
            site_path: Site path (e.g., "sites/Engineering" or site ID)
            library_name: Document library name
            folder_path: Subfolder path within library
            recursive: Include subfolders
        """
        self._ensure_authenticated()

        documents = []

        try:
            # Get site
            site_result = self._graph_request(f"/sites/{site_path}")
            if not site_result:
                logger.error(f"SharePoint site not found: {site_path}")
                return []

            site_id = site_result.get("id")

            # Get drive (document library)
            drives_result = self._graph_request(f"/sites/{site_id}/drives")
            if not drives_result:
                return []

            drive_id = None
            for drive in drives_result.get("value", []):
                if drive.get("name") == library_name:
                    drive_id = drive.get("id")
                    break

            if not drive_id:
                logger.error(f"Document library not found: {library_name}")
                return []

            # Build endpoint
            if folder_path:
                endpoint = f"/drives/{drive_id}/root:/{folder_path}:/children"
            else:
                endpoint = f"/drives/{drive_id}/root/children"

            def process_items(ep: str):
                result = self._graph_request(ep)
                if not result:
                    return

                for item in result.get("value", []):
                    if "folder" in item:
                        if recursive:
                            child_endpoint = (
                                f"/drives/{drive_id}/items/{item['id']}/children"
                            )
                            process_items(child_endpoint)
                    elif "file" in item:
                        doc_id = f"sharepoint:{site_id}:{item['id']}"
                        doc = self.load_document(doc_id)
                        if doc:
                            documents.append(doc)

                next_link = result.get("@odata.nextLink")
                if next_link:
                    process_items(next_link.replace(self.GRAPH_BASE_URL, ""))

            process_items(endpoint)
            logger.info(f"Loaded {len(documents)} files from SharePoint")
            return documents

        except Exception as e:
            logger.error(f"Failed to load SharePoint library: {e}")
            return []

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search across OneDrive and SharePoint.

        Args:
            query: Search query
            max_results: Maximum results
        """
        self._ensure_authenticated()

        documents = []

        try:
            # Use Microsoft Search API
            search_request = {
                "requests": [
                    {
                        "entityTypes": ["driveItem"],
                        "query": {"queryString": query},
                        "from": 0,
                        "size": max_results,
                    }
                ]
            }

            result = self._graph_request(
                "/search/query", method="POST", data=search_request
            )

            if not result:
                return []

            for hit_container in result.get("value", []):
                for hit in hit_container.get("hitsContainers", [{}])[0].get("hits", []):
                    resource = hit.get("resource", {})

                    # Determine source
                    if "parentReference" in resource:
                        drive_id = resource["parentReference"].get("driveId", "")
                        item_id = resource.get("id", "")

                        if drive_id and item_id:
                            doc_id = f"onedrive:{item_id}"
                            doc = self.load_document(doc_id)
                            if doc:
                                documents.append(doc)

            logger.info(f"Found {len(documents)} results for '{query}'")
            return documents

        except Exception as e:
            logger.error(f"Microsoft 365 search failed: {e}")
            return []

    def list_sharepoint_sites(self) -> list[dict[str, str]]:
        """List accessible SharePoint sites."""
        self._ensure_authenticated()

        try:
            result = self._graph_request("/sites?search=*")
            if not result:
                return []

            return [
                {
                    "id": site.get("id"),
                    "name": site.get("displayName"),
                    "url": site.get("webUrl"),
                }
                for site in result.get("value", [])
            ]

        except Exception as e:
            logger.error(f"Failed to list SharePoint sites: {e}")
            return []


# ============================================================================
# NOTION LOADER
# ============================================================================

# Check for notion-client availability
try:
    from notion_client import Client as NotionClient

    NOTION_AVAILABLE = True
except ImportError:
    NOTION_AVAILABLE = False


class NotionLoader(BaseLoader):
    """Document loader for Notion workspaces.

    Loads pages, databases, and blocks from Notion using the official API.

    Features:
    - Load individual pages with full content
    - Load database entries as documents
    - Search across workspace
    - Extract rich text, code blocks, tables
    - Preserve page hierarchy in metadata

    Requirements:
        pip install notion-client

    Example:
        loader = NotionLoader(token="secret_xxx")
        loader.authenticate()

        # Load a specific page
        doc = loader.load_document("page_id_here")

        # Load all pages from a database
        docs = loader.load_database("database_id_here")

        # Search workspace
        results = loader.search("project roadmap")
    """

    def __init__(self, token: str, **kwargs):
        """Initialize Notion loader.

        Args:
            token: Notion integration token (starts with 'secret_')
        """
        if not NOTION_AVAILABLE:
            raise ImportError(
                "notion-client is required for NotionLoader. "
                "Install with: pip install notion-client"
            )

        self._token = token
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Notion"

    def authenticate(self) -> bool:
        """Authenticate with Notion API."""
        try:
            self._client = NotionClient(auth=self._token)
            # Test connection by listing users
            self._client.users.me()
            self._authenticated = True
            logger.info("Notion authentication successful")
            return True
        except Exception as e:
            logger.error(f"Notion authentication failed: {e}")
            return False

    def _extract_rich_text(self, rich_text_array: list[dict]) -> str:
        """Extract plain text from Notion rich text array."""
        if not rich_text_array:
            return ""
        return "".join(item.get("plain_text", "") for item in rich_text_array)

    def _extract_block_content(self, block: dict) -> str:
        """Extract text content from a Notion block."""
        block_type = block.get("type", "")
        block_data = block.get(block_type, {})

        # Handle different block types
        if block_type in [
            "paragraph",
            "heading_1",
            "heading_2",
            "heading_3",
            "bulleted_list_item",
            "numbered_list_item",
            "quote",
            "callout",
            "toggle",
        ]:
            return self._extract_rich_text(block_data.get("rich_text", []))

        elif block_type == "code":
            code = self._extract_rich_text(block_data.get("rich_text", []))
            language = block_data.get("language", "")
            return f"```{language}\n{code}\n```"

        elif block_type == "to_do":
            text = self._extract_rich_text(block_data.get("rich_text", []))
            checked = "☑" if block_data.get("checked") else "☐"
            return f"{checked} {text}"

        elif block_type == "divider":
            return "---"

        elif block_type == "table_row":
            cells = block_data.get("cells", [])
            return " | ".join(self._extract_rich_text(cell) for cell in cells)

        elif block_type == "child_page":
            return f"[Page: {block_data.get('title', 'Untitled')}]"

        elif block_type == "child_database":
            return f"[Database: {block_data.get('title', 'Untitled')}]"

        elif block_type == "embed" or block_type == "bookmark":
            return f"[Link: {block_data.get('url', '')}]"

        elif block_type == "image":
            img_data = block_data.get("file", {}) or block_data.get("external", {})
            return f"[Image: {img_data.get('url', '')}]"

        return ""

    def _get_page_blocks(self, page_id: str, max_depth: int = 3) -> list[str]:
        """Recursively get all blocks from a page."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        content_parts = []

        try:
            response = self._client.blocks.children.list(block_id=page_id)
            blocks = response.get("results", [])

            for block in blocks:
                text = self._extract_block_content(block)
                if text:
                    # Add indentation for nested blocks
                    content_parts.append(text)

                # Recursively get children if has_children and within depth
                if block.get("has_children") and max_depth > 0:
                    child_content = self._get_page_blocks(block["id"], max_depth - 1)
                    content_parts.extend(f"  {line}" for line in child_content)

        except Exception as e:
            logger.warning(f"Error getting blocks for {page_id}: {e}")

        return content_parts

    def _get_page_title(self, page: dict) -> str:
        """Extract title from a page object."""
        properties = page.get("properties", {})

        # Try common title property names
        for prop_name in ["title", "Title", "Name", "name"]:
            if prop_name in properties:
                prop = properties[prop_name]
                if prop.get("type") == "title":
                    return self._extract_rich_text(prop.get("title", []))

        # Fallback: try first title-type property
        for prop in properties.values():
            if prop.get("type") == "title":
                return self._extract_rich_text(prop.get("title", []))

        return "Untitled"

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Notion page.

        Args:
            doc_id: Notion page ID (with or without dashes)

        Returns:
            LoadedDocument or None if not found
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            # Get page metadata
            page = self._client.pages.retrieve(page_id=doc_id)
            title = self._get_page_title(page)

            # Get page content
            content_parts = self._get_page_blocks(doc_id)
            content = "\n".join(content_parts)

            # Build metadata
            metadata = {
                "notion_id": page.get("id"),
                "url": page.get("url"),
                "created_time": page.get("created_time"),
                "last_edited_time": page.get("last_edited_time"),
                "created_by": page.get("created_by", {}).get("id"),
                "parent_type": page.get("parent", {}).get("type"),
            }

            return LoadedDocument(
                content=content,
                doc_id=doc_id,
                title=title,
                source="notion",
                metadata=metadata,
                created_at=page.get("created_time"),
                modified_at=page.get("last_edited_time"),
            )

        except Exception as e:
            logger.error(f"Failed to load Notion page {doc_id}: {e}")
            return None

    def load_database(
        self,
        database_id: str,
        filter_dict: Optional[dict] = None,
        sorts: Optional[list[dict]] = None,
        max_results: int = 100,
    ) -> list[LoadedDocument]:
        """Load entries from a Notion database.

        Args:
            database_id: Notion database ID
            filter_dict: Optional Notion filter object
            sorts: Optional list of sort objects
            max_results: Maximum entries to load

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            query_params = {"database_id": database_id}
            if filter_dict:
                query_params["filter"] = filter_dict
            if sorts:
                query_params["sorts"] = sorts

            # Paginate through results
            has_more = True
            next_cursor = None

            while has_more and len(docs) < max_results:
                if next_cursor:
                    query_params["start_cursor"] = next_cursor

                response = self._client.databases.query(**query_params)

                for page in response.get("results", []):
                    if len(docs) >= max_results:
                        break

                    doc = self.load_document(page["id"])
                    if doc:
                        docs.append(doc)

                has_more = response.get("has_more", False)
                next_cursor = response.get("next_cursor")

        except Exception as e:
            logger.error(f"Failed to load Notion database {database_id}: {e}")

        return docs

    def load_folder(self, folder_id: str) -> list[LoadedDocument]:
        """Load all pages from a database (alias for load_database).

        Args:
            folder_id: Notion database ID

        Returns:
            List of LoadedDocument
        """
        return self.load_database(folder_id)

    def search(
        self, query: str, max_results: int = 10, filter_type: Optional[str] = None
    ) -> list[LoadedDocument]:
        """Search Notion workspace.

        Args:
            query: Search query
            max_results: Maximum results to return
            filter_type: 'page' or 'database' to filter results

        Returns:
            List of matching LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            search_params = {"query": query}

            if filter_type:
                search_params["filter"] = {"property": "object", "value": filter_type}

            response = self._client.search(**search_params)

            for result in response.get("results", [])[:max_results]:
                if result.get("object") == "page":
                    doc = self.load_document(result["id"])
                    if doc:
                        docs.append(doc)

        except Exception as e:
            logger.error(f"Notion search failed: {e}")

        return docs

    def list_databases(self) -> list[dict[str, str]]:
        """List all databases in the workspace.

        Returns:
            List of dicts with 'id' and 'title'
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        databases = []

        try:
            response = self._client.search(
                filter={"property": "object", "value": "database"}
            )

            for db in response.get("results", []):
                title_prop = db.get("title", [])
                title = (
                    self._extract_rich_text(title_prop) if title_prop else "Untitled"
                )
                databases.append(
                    {
                        "id": db["id"],
                        "title": title,
                        "url": db.get("url", ""),
                    }
                )

        except Exception as e:
            logger.error(f"Failed to list Notion databases: {e}")

        return databases


# ============================================================================
# CONFLUENCE LOADER
# ============================================================================

# Check for atlassian-python-api availability
try:
    from atlassian import Confluence

    CONFLUENCE_AVAILABLE = True
except ImportError:
    CONFLUENCE_AVAILABLE = False


class ConfluenceLoader(BaseLoader):
    """Document loader for Atlassian Confluence.

    Loads pages, spaces, and attachments from Confluence Cloud or Server.

    Features:
    - Load individual pages with content
    - Load all pages from a space
    - Search using CQL (Confluence Query Language)
    - Extract attachments
    - Preserve page hierarchy

    Requirements:
        pip install atlassian-python-api

    Example:
        # Cloud
        loader = ConfluenceLoader(
            url="https://company.atlassian.net/wiki",
            username="user@company.com",
            api_token="xxx"
        )

        # Server
        loader = ConfluenceLoader(
            url="https://confluence.company.com",
            username="user",
            password="xxx"
        )

        loader.authenticate()

        # Load a page
        doc = loader.load_document("123456")

        # Load entire space
        docs = loader.load_space("PROJ")

        # Search
        results = loader.search("project AND roadmap")
    """

    def __init__(
        self,
        url: str,
        username: str,
        api_token: Optional[str] = None,
        password: Optional[str] = None,
        cloud: bool = True,
        **kwargs,
    ):
        """Initialize Confluence loader.

        Args:
            url: Confluence URL (e.g., https://company.atlassian.net/wiki)
            username: Username or email
            api_token: API token (for Cloud)
            password: Password (for Server)
            cloud: Whether this is Confluence Cloud (default True)
        """
        if not CONFLUENCE_AVAILABLE:
            raise ImportError(
                "atlassian-python-api is required for ConfluenceLoader. "
                "Install with: pip install atlassian-python-api"
            )

        self._url = url.rstrip("/")
        self._username = username
        self._api_token = api_token
        self._password = password
        self._cloud = cloud
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Confluence"

    def authenticate(self) -> bool:
        """Authenticate with Confluence."""
        try:
            if self._cloud:
                self._client = Confluence(
                    url=self._url,
                    username=self._username,
                    password=self._api_token,
                    cloud=True,
                )
            else:
                self._client = Confluence(
                    url=self._url, username=self._username, password=self._password
                )

            # Test connection
            self._client.get_all_spaces(limit=1)
            self._authenticated = True
            logger.info("Confluence authentication successful")
            return True

        except Exception as e:
            logger.error(f"Confluence authentication failed: {e}")
            return False

    def _clean_html(self, html_content: str) -> str:
        """Convert HTML to plain text."""
        if not html_content:
            return ""

        try:
            from html.parser import HTMLParser
            from io import StringIO

            class MLStripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.reset()
                    self.strict = False
                    self.convert_charrefs = True
                    self.text = StringIO()

                def handle_data(self, d):
                    self.text.write(d)

                def get_data(self):
                    return self.text.getvalue()

            s = MLStripper()
            s.feed(html_content)
            return s.get_data()

        except Exception:
            # Fallback: simple tag removal
            import re

            return re.sub(r"<[^>]+>", "", html_content)

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Confluence page.

        Args:
            doc_id: Page ID

        Returns:
            LoadedDocument or None if not found
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            page = self._client.get_page_by_id(
                doc_id, expand="body.storage,version,space,ancestors"
            )

            title = page.get("title", "Untitled")
            html_content = page.get("body", {}).get("storage", {}).get("value", "")
            content = self._clean_html(html_content)

            # Build hierarchy from ancestors
            ancestors = page.get("ancestors", [])
            hierarchy = [a.get("title", "") for a in ancestors]
            hierarchy.append(title)

            metadata = {
                "confluence_id": page.get("id"),
                "space_key": page.get("space", {}).get("key"),
                "space_name": page.get("space", {}).get("name"),
                "version": page.get("version", {}).get("number"),
                "url": f"{self._url}/pages/viewpage.action?pageId={doc_id}",
                "hierarchy": " > ".join(hierarchy),
                "ancestors": [a.get("id") for a in ancestors],
            }

            version_info = page.get("version", {})

            return LoadedDocument(
                content=content,
                doc_id=doc_id,
                title=title,
                source="confluence",
                metadata=metadata,
                created_at=(
                    version_info.get("when")
                    if version_info.get("number") == 1
                    else None
                ),
                modified_at=version_info.get("when"),
            )

        except Exception as e:
            logger.error(f"Failed to load Confluence page {doc_id}: {e}")
            return None

    def load_space(
        self, space_key: str, max_results: int = 100, content_type: str = "page"
    ) -> list[LoadedDocument]:
        """Load all pages from a Confluence space.

        Args:
            space_key: Space key (e.g., "PROJ")
            max_results: Maximum pages to load
            content_type: 'page' or 'blogpost'

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Get all pages in space
            start = 0
            limit = min(50, max_results)

            while len(docs) < max_results:
                pages = self._client.get_all_pages_from_space(
                    space_key,
                    start=start,
                    limit=limit,
                    content_type=content_type,
                    expand="body.storage,version",
                )

                if not pages:
                    break

                for page in pages:
                    if len(docs) >= max_results:
                        break

                    doc = self.load_document(page["id"])
                    if doc:
                        docs.append(doc)

                if len(pages) < limit:
                    break

                start += limit

        except Exception as e:
            logger.error(f"Failed to load Confluence space {space_key}: {e}")

        return docs

    def load_folder(self, folder_id: str) -> list[LoadedDocument]:
        """Load pages from a space (alias for load_space).

        Args:
            folder_id: Space key

        Returns:
            List of LoadedDocument
        """
        return self.load_space(folder_id)

    def search(
        self, query: str, max_results: int = 10, space_key: Optional[str] = None
    ) -> list[LoadedDocument]:
        """Search Confluence using CQL.

        Args:
            query: Search query (CQL supported)
            max_results: Maximum results
            space_key: Optional space to limit search

        Returns:
            List of matching LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Build CQL query
            cql = f'text ~ "{query}"'
            if space_key:
                cql = f'{cql} AND space = "{space_key}"'

            results = self._client.cql(cql, limit=max_results)

            for result in results.get("results", []):
                content = result.get("content", {})
                if content.get("type") == "page":
                    doc = self.load_document(content["id"])
                    if doc:
                        docs.append(doc)

        except Exception as e:
            logger.error(f"Confluence search failed: {e}")

        return docs

    def list_spaces(self) -> list[dict[str, str]]:
        """List all accessible spaces.

        Returns:
            List of dicts with 'key', 'name', 'type'
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        spaces = []

        try:
            result = self._client.get_all_spaces(limit=500)

            for space in result.get("results", []):
                spaces.append(
                    {
                        "key": space.get("key"),
                        "name": space.get("name"),
                        "type": space.get("type"),
                    }
                )

        except Exception as e:
            logger.error(f"Failed to list Confluence spaces: {e}")

        return spaces

    def get_page_attachments(
        self, page_id: str, load_content: bool = False
    ) -> list[dict]:
        """Get attachments from a page.

        Args:
            page_id: Page ID
            load_content: Whether to download attachment content

        Returns:
            List of attachment info dicts
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        attachments = []

        try:
            result = self._client.get_attachments_from_content(page_id)

            for att in result.get("results", []):
                info = {
                    "id": att.get("id"),
                    "title": att.get("title"),
                    "filename": att.get("title"),
                    "media_type": att.get("metadata", {}).get("mediaType"),
                    "size": att.get("extensions", {}).get("fileSize"),
                    "download_url": f"{self._url}{att.get('_links', {}).get('download', '')}",
                }

                if load_content and info["media_type"] and "text" in info["media_type"]:
                    try:
                        content = self._client.download_attachment(att["id"])
                        info["content"] = content.decode("utf-8", errors="ignore")
                    except Exception:
                        pass

                attachments.append(info)

        except Exception as e:
            logger.error(f"Failed to get attachments for {page_id}: {e}")

        return attachments


# ============================================================================
# SLACK LOADER
# ============================================================================

# Check for slack-sdk availability
try:
    from slack_sdk import WebClient as SlackClient
    from slack_sdk.errors import SlackApiError

    SLACK_AVAILABLE = True
except ImportError:
    SLACK_AVAILABLE = False


class SlackLoader(BaseLoader):
    """Document loader for Slack workspaces.

    Loads messages, threads, and files from Slack channels.

    Features:
    - Load messages from channels
    - Load thread replies
    - Search messages across workspace
    - Extract file content
    - User mention resolution

    Requirements:
        pip install slack-sdk

    Scopes required:
        - channels:history (public channels)
        - groups:history (private channels)
        - im:history (DMs)
        - mpim:history (group DMs)
        - search:read (search)
        - files:read (files)
        - users:read (user info)

    Example:
        loader = SlackLoader(token="xoxb-xxx")
        loader.authenticate()

        # Load messages from a channel
        docs = loader.load_channel("C1234567890", days=7)

        # Search messages
        results = loader.search("project update")

        # Load a thread
        thread = loader.load_thread("C123", "1234567890.123456")
    """

    def __init__(self, token: str, **kwargs):
        """Initialize Slack loader.

        Args:
            token: Slack Bot token (xoxb-...) or User token (xoxp-...)
        """
        if not SLACK_AVAILABLE:
            raise ImportError(
                "slack-sdk is required for SlackLoader. "
                "Install with: pip install slack-sdk"
            )

        self._token = token
        self._client: Optional[Any] = None
        self._authenticated = False
        self._user_cache: dict[str, str] = {}

    @property
    def source_name(self) -> str:
        return "Slack"

    def authenticate(self) -> bool:
        """Authenticate with Slack API."""
        try:
            self._client = SlackClient(token=self._token)

            # Test connection
            response = self._client.auth_test()
            if response["ok"]:
                self._authenticated = True
                logger.info(
                    f"Slack authentication successful (team: {response.get('team')})"
                )
                return True
            return False

        except Exception as e:
            logger.error(f"Slack authentication failed: {e}")
            return False

    def _get_user_name(self, user_id: str) -> str:
        """Get user display name from ID (cached)."""
        if user_id in self._user_cache:
            return self._user_cache[user_id]

        try:
            response = self._client.users_info(user=user_id)
            if response["ok"]:
                user = response["user"]
                name = user.get("real_name") or user.get("name", user_id)
                self._user_cache[user_id] = name
                return name
        except Exception:
            pass

        return user_id

    def _resolve_mentions(self, text: str) -> str:
        """Replace user mentions with names."""
        import re

        def replace_mention(match):
            user_id = match.group(1)
            return f"@{self._get_user_name(user_id)}"

        # Replace <@U123> patterns
        return re.sub(r"<@(U[A-Z0-9]+)>", replace_mention, text)

    def _format_message(self, msg: dict) -> str:
        """Format a Slack message to text."""
        text = msg.get("text", "")
        text = self._resolve_mentions(text)

        # Handle attachments
        attachments = msg.get("attachments", [])
        for att in attachments:
            if att.get("text"):
                text += f"\n[Attachment: {att.get('text')}]"
            elif att.get("fallback"):
                text += f"\n[Attachment: {att.get('fallback')}]"

        # Handle files
        files = msg.get("files", [])
        for f in files:
            text += f"\n[File: {f.get('name', 'unnamed')}]"

        return text

    def load_channel(
        self,
        channel_id: str,
        days: int = 7,
        max_messages: int = 1000,
        include_threads: bool = True,
    ) -> list[LoadedDocument]:
        """Load messages from a Slack channel.

        Args:
            channel_id: Channel ID (C..., G..., D...)
            days: Number of days to look back
            max_messages: Maximum messages to load
            include_threads: Whether to include thread replies

        Returns:
            List of LoadedDocument (one per message or thread)
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            from datetime import datetime, timedelta, timezone

            oldest = (datetime.now() - timedelta(days=days)).timestamp()

            # Paginate through messages
            cursor = None
            messages = []

            while len(messages) < max_messages:
                kwargs = {
                    "channel": channel_id,
                    "oldest": str(oldest),
                    "limit": min(200, max_messages - len(messages)),
                }
                if cursor:
                    kwargs["cursor"] = cursor

                response = self._client.conversations_history(**kwargs)

                if not response["ok"]:
                    break

                messages.extend(response.get("messages", []))

                cursor = response.get("response_metadata", {}).get("next_cursor")
                if not cursor:
                    break

            # Process messages
            for msg in messages:
                user = self._get_user_name(msg.get("user", "unknown"))
                text = self._format_message(msg)
                ts = msg.get("ts", "")

                # Load thread replies if present
                thread_content = ""
                if include_threads and msg.get("reply_count", 0) > 0:
                    thread_docs = self._load_thread_replies(channel_id, ts)
                    if thread_docs:
                        thread_content = "\n\nReplies:\n" + "\n".join(
                            f"  {d.metadata.get('user', 'unknown')}: {d.content}"
                            for d in thread_docs
                        )

                content = f"{user}: {text}{thread_content}"

                # Convert timestamp to datetime
                try:
                    dt = datetime.fromtimestamp(float(ts))
                    timestamp = dt.isoformat()
                except Exception:
                    timestamp = ts

                docs.append(
                    LoadedDocument(
                        content=content,
                        doc_id=f"{channel_id}:{ts}",
                        title=f"Slack message from {user}",
                        source="slack",
                        metadata={
                            "channel_id": channel_id,
                            "thread_ts": ts,
                            "user": user,
                            "user_id": msg.get("user"),
                            "reply_count": msg.get("reply_count", 0),
                        },
                        created_at=timestamp,
                    )
                )

        except Exception as e:
            logger.error(f"Failed to load Slack channel {channel_id}: {e}")

        return docs

    def _load_thread_replies(
        self, channel_id: str, thread_ts: str
    ) -> list[LoadedDocument]:
        """Load replies in a thread."""
        docs = []

        try:
            response = self._client.conversations_replies(
                channel=channel_id, ts=thread_ts, limit=100
            )

            if response["ok"]:
                # Skip first message (it's the parent)
                for msg in response.get("messages", [])[1:]:
                    user = self._get_user_name(msg.get("user", "unknown"))
                    text = self._format_message(msg)

                    docs.append(
                        LoadedDocument(
                            content=text,
                            doc_id=f"{channel_id}:{msg.get('ts')}",
                            title=f"Reply from {user}",
                            source="slack",
                            metadata={
                                "channel_id": channel_id,
                                "thread_ts": thread_ts,
                                "user": user,
                            },
                        )
                    )

        except Exception as e:
            logger.warning(f"Failed to load thread {thread_ts}: {e}")

        return docs

    def load_thread(self, channel_id: str, thread_ts: str) -> list[LoadedDocument]:
        """Load a complete thread.

        Args:
            channel_id: Channel ID
            thread_ts: Thread timestamp

        Returns:
            List of LoadedDocument for all messages in thread
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.conversations_replies(
                channel=channel_id, ts=thread_ts, limit=1000
            )

            if response["ok"]:
                for msg in response.get("messages", []):
                    user = self._get_user_name(msg.get("user", "unknown"))
                    text = self._format_message(msg)

                    docs.append(
                        LoadedDocument(
                            content=text,
                            doc_id=f"{channel_id}:{msg.get('ts')}",
                            title=f"Message from {user}",
                            source="slack",
                            metadata={
                                "channel_id": channel_id,
                                "thread_ts": thread_ts,
                                "user": user,
                                "is_parent": msg.get("ts") == thread_ts,
                            },
                        )
                    )

        except Exception as e:
            logger.error(f"Failed to load thread: {e}")

        return docs

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single message.

        Args:
            doc_id: Message ID in format "channel_id:timestamp"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            channel_id, ts = doc_id.split(":", 1)

            response = self._client.conversations_history(
                channel=channel_id, oldest=ts, latest=ts, inclusive=True, limit=1
            )

            if response["ok"] and response.get("messages"):
                msg = response["messages"][0]
                user = self._get_user_name(msg.get("user", "unknown"))
                text = self._format_message(msg)

                return LoadedDocument(
                    content=text,
                    doc_id=doc_id,
                    title=f"Message from {user}",
                    source="slack",
                    metadata={
                        "channel_id": channel_id,
                        "thread_ts": ts,
                        "user": user,
                    },
                )

        except Exception as e:
            logger.error(f"Failed to load Slack message {doc_id}: {e}")

        return None

    def load_folder(self, folder_id: str) -> list[LoadedDocument]:
        """Load messages from a channel (alias for load_channel).

        Args:
            folder_id: Channel ID

        Returns:
            List of LoadedDocument
        """
        return self.load_channel(folder_id)

    def search(
        self, query: str, max_results: int = 20, sort: str = "timestamp"
    ) -> list[LoadedDocument]:
        """Search Slack messages.

        Args:
            query: Search query (supports Slack search modifiers)
            max_results: Maximum results
            sort: 'timestamp' or 'score'

        Returns:
            List of matching LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.search_messages(
                query=query, count=max_results, sort=sort
            )

            if response["ok"]:
                for match in response.get("messages", {}).get("matches", []):
                    user = self._get_user_name(match.get("user", "unknown"))
                    text = self._resolve_mentions(match.get("text", ""))
                    channel = match.get("channel", {})

                    docs.append(
                        LoadedDocument(
                            content=text,
                            doc_id=f"{channel.get('id')}:{match.get('ts')}",
                            title=f"Message from {user} in #{channel.get('name', 'unknown')}",
                            source="slack",
                            metadata={
                                "channel_id": channel.get("id"),
                                "channel_name": channel.get("name"),
                                "user": user,
                                "permalink": match.get("permalink"),
                            },
                        )
                    )

        except Exception as e:
            logger.error(f"Slack search failed: {e}")

        return docs

    def list_channels(
        self, types: str = "public_channel,private_channel"
    ) -> list[dict[str, str]]:
        """List accessible channels.

        Args:
            types: Channel types to list

        Returns:
            List of dicts with 'id', 'name', 'is_private'
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        channels = []

        try:
            cursor = None

            while True:
                kwargs = {"types": types, "limit": 200}
                if cursor:
                    kwargs["cursor"] = cursor

                response = self._client.conversations_list(**kwargs)

                if response["ok"]:
                    for ch in response.get("channels", []):
                        channels.append(
                            {
                                "id": ch.get("id"),
                                "name": ch.get("name"),
                                "is_private": ch.get("is_private", False),
                                "is_member": ch.get("is_member", False),
                            }
                        )

                cursor = response.get("response_metadata", {}).get("next_cursor")
                if not cursor:
                    break

        except Exception as e:
            logger.error(f"Failed to list Slack channels: {e}")

        return channels


# ============================================================================
# DROPBOX LOADER
# ============================================================================

# Check for dropbox availability
try:
    import dropbox
    from dropbox.exceptions import AuthError as DropboxAuthError

    DROPBOX_AVAILABLE = True
except ImportError:
    DROPBOX_AVAILABLE = False


class DropboxLoader(BaseLoader):
    """Document loader for Dropbox files and folders.

    Load documents from Dropbox personal and business accounts.

    Features:
    - Load files from folders
    - Search files by name or content
    - Support for shared folders
    - Automatic text extraction

    Requirements:
        pip install dropbox

    Environment variables:
        DROPBOX_ACCESS_TOKEN: OAuth2 access token
        DROPBOX_REFRESH_TOKEN: Refresh token for long-lived access
        DROPBOX_APP_KEY: App key (required for refresh)
        DROPBOX_APP_SECRET: App secret (required for refresh)

    Example:
        loader = DropboxLoader(access_token="xxx")
        loader.authenticate()
        docs = loader.load_folder("/Work/Projects")
        results = loader.search("quarterly report")
    """

    def __init__(
        self,
        access_token: Optional[str] = None,
        refresh_token: Optional[str] = None,
        app_key: Optional[str] = None,
        app_secret: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Dropbox loader.

        Args:
            access_token: OAuth2 access token
            refresh_token: Refresh token for auto-renewal
            app_key: App key for refresh
            app_secret: App secret for refresh
        """
        if not DROPBOX_AVAILABLE:
            raise ImportError(
                "dropbox is required for DropboxLoader. "
                "Install with: pip install dropbox"
            )

        self._access_token = access_token or os.environ.get("DROPBOX_ACCESS_TOKEN")
        self._refresh_token = refresh_token or os.environ.get("DROPBOX_REFRESH_TOKEN")
        self._app_key = app_key or os.environ.get("DROPBOX_APP_KEY")
        self._app_secret = app_secret or os.environ.get("DROPBOX_APP_SECRET")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Dropbox"

    def authenticate(self) -> bool:
        """Authenticate with Dropbox API."""
        try:
            if self._refresh_token and self._app_key and self._app_secret:
                self._client = dropbox.Dropbox(
                    oauth2_refresh_token=self._refresh_token,
                    app_key=self._app_key,
                    app_secret=self._app_secret,
                )
            elif self._access_token:
                self._client = dropbox.Dropbox(self._access_token)
            else:
                raise ValueError("No access token or refresh token provided")

            # Test connection
            account = self._client.users_get_current_account()
            self._authenticated = True
            logger.info(f"Dropbox authentication successful (user: {account.email})")
            return True

        except Exception as e:
            logger.error(f"Dropbox authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single file by path.

        Args:
            doc_id: File path in Dropbox (e.g., "/folder/file.pdf")

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            metadata, response = self._client.files_download(doc_id)
            content_bytes = response.content

            # Extract text based on mime type
            content = self._extract_content(content_bytes, metadata.name)

            return LoadedDocument(
                content=content,
                source="dropbox",
                source_id=metadata.id,
                filename=metadata.name,
                mime_type=mimetypes.guess_type(metadata.name)[0]
                or "application/octet-stream",
                created_at=None,
                modified_at=(
                    metadata.server_modified
                    if hasattr(metadata, "server_modified")
                    else None
                ),
                size_bytes=metadata.size if hasattr(metadata, "size") else 0,
                metadata={"path": doc_id, "rev": getattr(metadata, "rev", None)},
            )

        except Exception as e:
            logger.error(f"Failed to load Dropbox file {doc_id}: {e}")
            return None

    def _extract_content(self, content_bytes: bytes, filename: str) -> str:
        """Extract text content from file bytes."""
        mime_type = mimetypes.guess_type(filename)[0] or ""

        if mime_type == "application/pdf":
            return self._extract_text_from_pdf(content_bytes)
        elif mime_type.startswith("text/") or filename.endswith(
            (".txt", ".md", ".json", ".yaml", ".yml", ".csv")
        ):
            return content_bytes.decode("utf-8", errors="ignore")
        elif mime_type in ("application/json",):
            return content_bytes.decode("utf-8", errors="ignore")
        else:
            return f"[Binary file: {filename}]"

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a Dropbox folder.

        Args:
            folder_path: Folder path (e.g., "/Work/Projects")
            recursive: Include subfolders

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Normalize path
            if not folder_path.startswith("/"):
                folder_path = "/" + folder_path

            result = self._client.files_list_folder(folder_path, recursive=recursive)

            while True:
                for entry in result.entries:
                    if (
                        hasattr(entry, "is_downloadable")
                        or entry.__class__.__name__ == "FileMetadata"
                    ):
                        doc = self.load_document(entry.path_display)
                        if doc:
                            docs.append(doc)

                if not result.has_more:
                    break
                result = self._client.files_list_folder_continue(result.cursor)

        except Exception as e:
            logger.error(f"Failed to load Dropbox folder {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for files in Dropbox.

        Args:
            query: Search query
            max_results: Maximum results to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            result = self._client.files_search_v2(query)

            count = 0
            for match in result.matches:
                if count >= max_results:
                    break

                metadata = match.metadata.get_metadata()
                if hasattr(metadata, "path_display"):
                    doc = self.load_document(metadata.path_display)
                    if doc:
                        docs.append(doc)
                        count += 1

        except Exception as e:
            logger.error(f"Dropbox search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading (runs sync in executor)."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# BOX LOADER
# ============================================================================

# Check for boxsdk availability
try:
    from boxsdk import Client as BoxClient
    from boxsdk import OAuth2 as BoxOAuth2

    BOX_AVAILABLE = True
except ImportError:
    BOX_AVAILABLE = False


class BoxLoader(BaseLoader):
    """Document loader for Box.com enterprise file sharing.

    Load documents from Box enterprise and personal accounts.

    Features:
    - Load files from folders
    - Search across all content
    - Support for shared folders and collaborations
    - Automatic text extraction

    Requirements:
        pip install boxsdk

    Environment variables:
        BOX_CLIENT_ID: OAuth2 client ID
        BOX_CLIENT_SECRET: OAuth2 client secret
        BOX_ACCESS_TOKEN: Access token
        BOX_DEVELOPER_TOKEN: Developer token (for testing)

    Example:
        loader = BoxLoader(developer_token="xxx")
        loader.authenticate()
        docs = loader.load_folder("0")  # Root folder
        results = loader.search("contract")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        access_token: Optional[str] = None,
        developer_token: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Box loader.

        Args:
            client_id: OAuth2 client ID
            client_secret: OAuth2 client secret
            access_token: OAuth2 access token
            developer_token: Developer token (short-lived, for testing)
        """
        if not BOX_AVAILABLE:
            raise ImportError(
                "boxsdk is required for BoxLoader. " "Install with: pip install boxsdk"
            )

        self._client_id = client_id or os.environ.get("BOX_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get("BOX_CLIENT_SECRET")
        self._access_token = access_token or os.environ.get("BOX_ACCESS_TOKEN")
        self._developer_token = developer_token or os.environ.get("BOX_DEVELOPER_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Box"

    def authenticate(self) -> bool:
        """Authenticate with Box API."""
        try:
            if self._developer_token:
                oauth2 = BoxOAuth2(
                    client_id=self._client_id or "",
                    client_secret=self._client_secret or "",
                    access_token=self._developer_token,
                )
            elif self._access_token:
                oauth2 = BoxOAuth2(
                    client_id=self._client_id or "",
                    client_secret=self._client_secret or "",
                    access_token=self._access_token,
                )
            else:
                raise ValueError("No access token or developer token provided")

            self._client = BoxClient(oauth2)

            # Test connection
            user = self._client.user().get()
            self._authenticated = True
            logger.info(f"Box authentication successful (user: {user.login})")
            return True

        except Exception as e:
            logger.error(f"Box authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single file by ID.

        Args:
            doc_id: Box file ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            file_info = self._client.file(doc_id).get()
            content_bytes = self._client.file(doc_id).content()

            content = self._extract_content(content_bytes, file_info.name)

            return LoadedDocument(
                content=content,
                source="box",
                source_id=doc_id,
                filename=file_info.name,
                mime_type=mimetypes.guess_type(file_info.name)[0]
                or "application/octet-stream",
                created_at=(
                    datetime.fromisoformat(file_info.created_at.replace("Z", "+00:00"))
                    if file_info.created_at
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(file_info.modified_at.replace("Z", "+00:00"))
                    if file_info.modified_at
                    else None
                ),
                size_bytes=file_info.size or 0,
                metadata={
                    "parent_id": file_info.parent.id if file_info.parent else None
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Box file {doc_id}: {e}")
            return None

    def _extract_content(self, content_bytes: bytes, filename: str) -> str:
        """Extract text content from file bytes."""
        mime_type = mimetypes.guess_type(filename)[0] or ""

        if mime_type == "application/pdf":
            return self._extract_text_from_pdf(content_bytes)
        elif mime_type.startswith("text/") or filename.endswith(
            (".txt", ".md", ".json", ".yaml", ".yml", ".csv")
        ):
            return content_bytes.decode("utf-8", errors="ignore")
        else:
            return f"[Binary file: {filename}]"

    def load_folder(
        self, folder_id: str = "0", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a Box folder.

        Args:
            folder_id: Folder ID ("0" for root)
            recursive: Include subfolders

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            items = self._client.folder(folder_id).get_items()

            for item in items:
                if item.type == "file":
                    doc = self.load_document(item.id)
                    if doc:
                        docs.append(doc)
                elif item.type == "folder" and recursive:
                    docs.extend(self.load_folder(item.id, recursive=True))

        except Exception as e:
            logger.error(f"Failed to load Box folder {folder_id}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for files in Box.

        Args:
            query: Search query
            max_results: Maximum results to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            results = self._client.search().query(query, limit=max_results)

            for item in results:
                if item.type == "file":
                    doc = self.load_document(item.id)
                    if doc:
                        docs.append(doc)

        except Exception as e:
            logger.error(f"Box search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_id: str = "0", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_folder, folder_id, recursive)


# ============================================================================
# ONEDRIVE LOADER
# ============================================================================

# Uses MSAL (already imported for Microsoft365Loader)


class OneDriveLoader(BaseLoader):
    """Document loader for Microsoft OneDrive.

    Load documents from OneDrive personal and business accounts.

    Features:
    - Load files from folders
    - Search files
    - Support for shared items
    - Automatic text extraction

    Requirements:
        pip install msal requests

    Environment variables:
        ONEDRIVE_CLIENT_ID: Azure AD app client ID
        ONEDRIVE_CLIENT_SECRET: Azure AD app client secret
        ONEDRIVE_TENANT_ID: Azure AD tenant ID (or "consumers" for personal)

    Example:
        loader = OneDriveLoader(client_id="xxx", client_secret="yyy")
        loader.authenticate()
        docs = loader.load_folder("/Documents")
        results = loader.search("report")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        tenant_id: Optional[str] = None,
        **kwargs,
    ):
        """Initialize OneDrive loader.

        Args:
            client_id: Azure AD application client ID
            client_secret: Azure AD application client secret
            tenant_id: Azure AD tenant ID ("consumers" for personal accounts)
        """
        if not MSAL_AVAILABLE:
            raise ImportError(
                "msal is required for OneDriveLoader. " "Install with: pip install msal"
            )

        self._client_id = client_id or os.environ.get("ONEDRIVE_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get("ONEDRIVE_CLIENT_SECRET")
        self._tenant_id = tenant_id or os.environ.get("ONEDRIVE_TENANT_ID", "consumers")
        self._access_token: Optional[str] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "OneDrive"

    def authenticate(self) -> bool:
        """Authenticate with Microsoft Graph API."""
        try:
            import msal
            import requests

            authority = f"https://login.microsoftonline.com/{self._tenant_id}"
            app = msal.ConfidentialClientApplication(
                self._client_id,
                authority=authority,
                client_credential=self._client_secret,
            )

            scopes = ["https://graph.microsoft.com/.default"]
            result = app.acquire_token_for_client(scopes=scopes)

            if "access_token" in result:
                self._access_token = result["access_token"]
                self._authenticated = True
                logger.info("OneDrive authentication successful")
                return True
            else:
                logger.error(f"OneDrive auth failed: {result.get('error_description')}")
                return False

        except Exception as e:
            logger.error(f"OneDrive authentication failed: {e}")
            return False

    def _make_request(self, endpoint: str) -> Optional[dict]:
        """Make authenticated request to Graph API."""
        import requests

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0{endpoint}",
            headers=headers,
        )

        if response.status_code == 200:
            return response.json()
        else:
            logger.error(f"Graph API error: {response.status_code} - {response.text}")
            return None

    def _download_file(self, item_id: str) -> Optional[bytes]:
        """Download file content."""
        import requests

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0/me/drive/items/{item_id}/content",
            headers=headers,
        )

        if response.status_code == 200:
            return response.content
        return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single file by ID.

        Args:
            doc_id: OneDrive item ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            item = self._make_request(f"/me/drive/items/{doc_id}")
            if not item:
                return None

            content_bytes = self._download_file(doc_id)
            if not content_bytes:
                return None

            content = self._extract_content(content_bytes, item.get("name", ""))

            return LoadedDocument(
                content=content,
                source="onedrive",
                source_id=doc_id,
                filename=item.get("name", ""),
                mime_type=item.get("file", {}).get(
                    "mimeType", "application/octet-stream"
                ),
                created_at=(
                    datetime.fromisoformat(
                        item["createdDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("createdDateTime")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(
                        item["lastModifiedDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("lastModifiedDateTime")
                    else None
                ),
                size_bytes=item.get("size", 0),
                metadata={"webUrl": item.get("webUrl")},
            )

        except Exception as e:
            logger.error(f"Failed to load OneDrive file {doc_id}: {e}")
            return None

    def _extract_content(self, content_bytes: bytes, filename: str) -> str:
        """Extract text content from file bytes."""
        mime_type = mimetypes.guess_type(filename)[0] or ""

        if mime_type == "application/pdf":
            return self._extract_text_from_pdf(content_bytes)
        elif mime_type.startswith("text/") or filename.endswith(
            (".txt", ".md", ".json", ".yaml", ".yml", ".csv")
        ):
            return content_bytes.decode("utf-8", errors="ignore")
        else:
            return f"[Binary file: {filename}]"

    def load_folder(
        self, folder_path: str = "/", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a OneDrive folder.

        Args:
            folder_path: Folder path (e.g., "/Documents/Work")
            recursive: Include subfolders

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Get folder contents
            if folder_path == "/" or folder_path == "":
                endpoint = "/me/drive/root/children"
            else:
                endpoint = f"/me/drive/root:/{folder_path.strip('/')}:/children"

            result = self._make_request(endpoint)
            if not result:
                return docs

            for item in result.get("value", []):
                if "file" in item:
                    doc = self.load_document(item["id"])
                    if doc:
                        docs.append(doc)
                elif "folder" in item and recursive:
                    subfolder_path = f"{folder_path.rstrip('/')}/{item['name']}"
                    docs.extend(self.load_folder(subfolder_path, recursive=True))

        except Exception as e:
            logger.error(f"Failed to load OneDrive folder {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for files in OneDrive.

        Args:
            query: Search query
            max_results: Maximum results to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            result = self._make_request(
                f"/me/drive/root/search(q='{query}')?$top={max_results}"
            )
            if not result:
                return docs

            for item in result.get("value", []):
                if "file" in item:
                    doc = self.load_document(item["id"])
                    if doc:
                        docs.append(doc)

        except Exception as e:
            logger.error(f"OneDrive search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str = "/", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# SHAREPOINT LOADER
# ============================================================================


class SharePointLoader(BaseLoader):
    """Document loader for Microsoft SharePoint.

    Load documents from SharePoint sites, libraries, and lists.

    Features:
    - Load files from document libraries
    - Load list items
    - Search across sites
    - Support for multiple sites

    Requirements:
        pip install msal requests

    Environment variables:
        SHAREPOINT_CLIENT_ID: Azure AD app client ID
        SHAREPOINT_CLIENT_SECRET: Azure AD app client secret
        SHAREPOINT_TENANT_ID: Azure AD tenant ID
        SHAREPOINT_SITE_URL: SharePoint site URL

    Example:
        loader = SharePointLoader(
            client_id="xxx",
            client_secret="yyy",
            tenant_id="zzz",
            site_url="https://company.sharepoint.com/sites/team"
        )
        loader.authenticate()
        docs = loader.load_library("Documents")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        tenant_id: Optional[str] = None,
        site_url: Optional[str] = None,
        **kwargs,
    ):
        """Initialize SharePoint loader.

        Args:
            client_id: Azure AD application client ID
            client_secret: Azure AD application client secret
            tenant_id: Azure AD tenant ID
            site_url: SharePoint site URL
        """
        if not MSAL_AVAILABLE:
            raise ImportError(
                "msal is required for SharePointLoader. "
                "Install with: pip install msal"
            )

        self._client_id = client_id or os.environ.get("SHAREPOINT_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get(
            "SHAREPOINT_CLIENT_SECRET"
        )
        self._tenant_id = tenant_id or os.environ.get("SHAREPOINT_TENANT_ID")
        self._site_url = site_url or os.environ.get("SHAREPOINT_SITE_URL")
        self._access_token: Optional[str] = None
        self._site_id: Optional[str] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "SharePoint"

    def authenticate(self) -> bool:
        """Authenticate with Microsoft Graph API."""
        try:
            import msal
            import requests

            authority = f"https://login.microsoftonline.com/{self._tenant_id}"
            app = msal.ConfidentialClientApplication(
                self._client_id,
                authority=authority,
                client_credential=self._client_secret,
            )

            scopes = ["https://graph.microsoft.com/.default"]
            result = app.acquire_token_for_client(scopes=scopes)

            if "access_token" in result:
                self._access_token = result["access_token"]

                # Get site ID from URL
                if self._site_url:
                    self._site_id = self._get_site_id()

                self._authenticated = True
                logger.info("SharePoint authentication successful")
                return True
            else:
                logger.error(
                    f"SharePoint auth failed: {result.get('error_description')}"
                )
                return False

        except Exception as e:
            logger.error(f"SharePoint authentication failed: {e}")
            return False

    def _get_site_id(self) -> Optional[str]:
        """Get SharePoint site ID from URL."""
        import requests
        from urllib.parse import urlparse

        parsed = urlparse(self._site_url)
        host = parsed.netloc
        path = parsed.path

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0/sites/{host}:{path}",
            headers=headers,
        )

        if response.status_code == 200:
            return response.json().get("id")
        return None

    def _make_request(self, endpoint: str) -> Optional[dict]:
        """Make authenticated request to Graph API."""
        import requests

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0{endpoint}",
            headers=headers,
        )

        if response.status_code == 200:
            return response.json()
        else:
            logger.error(f"Graph API error: {response.status_code}")
            return None

    def _download_file(self, drive_id: str, item_id: str) -> Optional[bytes]:
        """Download file content."""
        import requests

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/content",
            headers=headers,
        )

        if response.status_code == 200:
            return response.content
        return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single file by ID.

        Args:
            doc_id: Item ID in format "drive_id:item_id"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            drive_id, item_id = doc_id.split(":", 1)

            item = self._make_request(f"/drives/{drive_id}/items/{item_id}")
            if not item:
                return None

            content_bytes = self._download_file(drive_id, item_id)
            if not content_bytes:
                return None

            content = self._extract_content(content_bytes, item.get("name", ""))

            return LoadedDocument(
                content=content,
                source="sharepoint",
                source_id=doc_id,
                filename=item.get("name", ""),
                mime_type=item.get("file", {}).get(
                    "mimeType", "application/octet-stream"
                ),
                created_at=(
                    datetime.fromisoformat(
                        item["createdDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("createdDateTime")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(
                        item["lastModifiedDateTime"].replace("Z", "+00:00")
                    )
                    if item.get("lastModifiedDateTime")
                    else None
                ),
                size_bytes=item.get("size", 0),
                metadata={"webUrl": item.get("webUrl"), "site_id": self._site_id},
            )

        except Exception as e:
            logger.error(f"Failed to load SharePoint file {doc_id}: {e}")
            return None

    def _extract_content(self, content_bytes: bytes, filename: str) -> str:
        """Extract text content from file bytes."""
        mime_type = mimetypes.guess_type(filename)[0] or ""

        if mime_type == "application/pdf":
            return self._extract_text_from_pdf(content_bytes)
        elif mime_type.startswith("text/") or filename.endswith(
            (".txt", ".md", ".json", ".yaml", ".yml", ".csv")
        ):
            return content_bytes.decode("utf-8", errors="ignore")
        else:
            return f"[Binary file: {filename}]"

    def load_folder(
        self, folder_path: str = "/", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a SharePoint library/folder.

        Args:
            folder_path: Library or folder path
            recursive: Include subfolders

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        return self.load_library(folder_path, recursive=recursive)

    def load_library(
        self, library_name: str = "Documents", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a SharePoint document library.

        Args:
            library_name: Library name (e.g., "Documents", "Shared Documents")
            recursive: Include subfolders

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Get drives (document libraries)
            drives = self._make_request(f"/sites/{self._site_id}/drives")
            if not drives:
                return docs

            # Find the library
            drive_id = None
            for drive in drives.get("value", []):
                if drive.get("name", "").lower() == library_name.lower():
                    drive_id = drive["id"]
                    break

            if not drive_id:
                logger.warning(f"Library '{library_name}' not found")
                return docs

            # Get items
            docs = self._load_drive_items(drive_id, "", recursive)

        except Exception as e:
            logger.error(f"Failed to load SharePoint library {library_name}: {e}")

        return docs

    def _load_drive_items(
        self, drive_id: str, path: str, recursive: bool
    ) -> list[LoadedDocument]:
        """Load items from a drive path."""
        docs = []

        try:
            if path:
                endpoint = f"/drives/{drive_id}/root:/{path}:/children"
            else:
                endpoint = f"/drives/{drive_id}/root/children"

            result = self._make_request(endpoint)
            if not result:
                return docs

            for item in result.get("value", []):
                if "file" in item:
                    doc = self.load_document(f"{drive_id}:{item['id']}")
                    if doc:
                        docs.append(doc)
                elif "folder" in item and recursive:
                    subpath = f"{path}/{item['name']}" if path else item["name"]
                    docs.extend(self._load_drive_items(drive_id, subpath, recursive))

        except Exception as e:
            logger.error(f"Failed to load drive items: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for files in SharePoint.

        Args:
            query: Search query
            max_results: Maximum results to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            result = self._make_request(
                f"/sites/{self._site_id}/drive/root/search(q='{query}')?$top={max_results}"
            )
            if not result:
                return docs

            for item in result.get("value", []):
                if "file" in item:
                    # Get drive ID from parentReference
                    drive_id = item.get("parentReference", {}).get("driveId", "")
                    if drive_id:
                        doc = self.load_document(f"{drive_id}:{item['id']}")
                        if doc:
                            docs.append(doc)

        except Exception as e:
            logger.error(f"SharePoint search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str = "/", recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# DISCORD LOADER
# ============================================================================

# Check for discord.py availability
try:
    import discord
    from discord.ext import commands

    DISCORD_AVAILABLE = True
except ImportError:
    DISCORD_AVAILABLE = False


class DiscordLoader(BaseLoader):
    """Document loader for Discord messages and channels.

    Load messages from Discord servers, channels, and threads.

    Features:
    - Load messages from channels
    - Load thread messages
    - Search messages
    - Support for multiple servers

    Requirements:
        pip install discord.py

    Environment variables:
        DISCORD_BOT_TOKEN: Bot token

    Example:
        loader = DiscordLoader(token="xxx")
        loader.authenticate()
        docs = loader.load_channel(123456789, days=7)
    """

    def __init__(self, token: Optional[str] = None, **kwargs):
        """Initialize Discord loader.

        Args:
            token: Discord bot token
        """
        if not DISCORD_AVAILABLE:
            raise ImportError(
                "discord.py is required for DiscordLoader. "
                "Install with: pip install discord.py"
            )

        self._token = token or os.environ.get("DISCORD_BOT_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Discord"

    def authenticate(self) -> bool:
        """Authenticate with Discord API.

        Note: Discord uses async client, so this just validates token format.
        Actual connection happens in load methods.
        """
        if not self._token:
            logger.error("No Discord token provided")
            return False

        self._authenticated = True
        logger.info("Discord token set (connection happens on load)")
        return True

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single message by ID.

        Args:
            doc_id: Message ID in format "channel_id:message_id"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        import asyncio

        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            return loop.run_until_complete(self._load_message_async(doc_id))
        finally:
            loop.close()

    async def _load_message_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async load single message."""
        import discord

        intents = discord.Intents.default()
        intents.message_content = True
        client = discord.Client(intents=intents)

        result = None

        @client.event
        async def on_ready():
            nonlocal result
            try:
                channel_id, message_id = doc_id.split(":", 1)
                channel = client.get_channel(int(channel_id))
                if channel:
                    message = await channel.fetch_message(int(message_id))
                    result = LoadedDocument(
                        content=message.content,
                        source="discord",
                        source_id=doc_id,
                        filename="",
                        created_at=message.created_at,
                        metadata={
                            "author": str(message.author),
                            "channel": str(channel),
                            "guild": str(message.guild) if message.guild else None,
                        },
                    )
            except Exception as e:
                logger.error(f"Failed to load Discord message: {e}")
            finally:
                await client.close()

        await client.start(self._token)
        return result

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load messages from a channel.

        Args:
            folder_path: Channel ID as string
            recursive: Include threads

        Returns:
            List of LoadedDocument
        """
        return self.load_channel(int(folder_path), days=7, include_threads=recursive)

    def load_channel(
        self,
        channel_id: int,
        days: int = 7,
        max_messages: int = 1000,
        include_threads: bool = True,
    ) -> list[LoadedDocument]:
        """Load messages from a Discord channel.

        Args:
            channel_id: Discord channel ID
            days: Number of days to look back
            max_messages: Maximum messages to load
            include_threads: Include thread messages

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        import asyncio

        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            return loop.run_until_complete(
                self._load_channel_async(
                    channel_id, days, max_messages, include_threads
                )
            )
        finally:
            loop.close()

    async def _load_channel_async(
        self, channel_id: int, days: int, max_messages: int, include_threads: bool
    ) -> list[LoadedDocument]:
        """Async load channel messages."""
        import discord

        intents = discord.Intents.default()
        intents.message_content = True
        client = discord.Client(intents=intents)

        docs = []
        after_date = datetime.now(timezone.utc) - timedelta(days=days)

        @client.event
        async def on_ready():
            nonlocal docs
            try:
                channel = client.get_channel(channel_id)
                if channel:
                    count = 0
                    async for message in channel.history(
                        after=after_date, limit=max_messages
                    ):
                        docs.append(
                            LoadedDocument(
                                content=f"{message.author}: {message.content}",
                                source="discord",
                                source_id=f"{channel_id}:{message.id}",
                                filename="",
                                created_at=message.created_at,
                                metadata={
                                    "author": str(message.author),
                                    "channel": str(channel),
                                    "guild": (
                                        str(message.guild) if message.guild else None
                                    ),
                                },
                            )
                        )
                        count += 1
                        if count >= max_messages:
                            break

                    # Load threads if requested
                    if include_threads and hasattr(channel, "threads"):
                        for thread in channel.threads:
                            async for msg in thread.history(
                                after=after_date, limit=100
                            ):
                                docs.append(
                                    LoadedDocument(
                                        content=f"{msg.author}: {msg.content}",
                                        source="discord",
                                        source_id=f"{thread.id}:{msg.id}",
                                        filename="",
                                        created_at=msg.created_at,
                                        metadata={
                                            "author": str(msg.author),
                                            "thread": str(thread),
                                            "channel": str(channel),
                                        },
                                    )
                                )

            except Exception as e:
                logger.error(f"Failed to load Discord channel: {e}")
            finally:
                await client.close()

        await client.start(self._token)
        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search messages (Discord API doesn't support search for bots).

        This is a stub - Discord bots cannot use search API.
        Consider loading recent messages and filtering locally.

        Args:
            query: Search query
            max_results: Maximum results

        Returns:
            Empty list (search not supported for bots)
        """
        logger.warning("Discord bot API does not support message search")
        return []

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        return await self._load_message_async(doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async channel loading."""
        return await self._load_channel_async(
            int(folder_path), days=7, max_messages=1000, include_threads=recursive
        )


# ============================================================================
# TEAMS LOADER
# ============================================================================


class TeamsLoader(BaseLoader):
    """Document loader for Microsoft Teams messages and files.

    Load messages, files, and channel content from Microsoft Teams.

    Features:
    - Load messages from channels
    - Load chat messages
    - Load shared files
    - Search across teams

    Requirements:
        pip install msal requests

    Environment variables:
        TEAMS_CLIENT_ID: Azure AD app client ID
        TEAMS_CLIENT_SECRET: Azure AD app client secret
        TEAMS_TENANT_ID: Azure AD tenant ID

    Example:
        loader = TeamsLoader(client_id="xxx", client_secret="yyy", tenant_id="zzz")
        loader.authenticate()
        docs = loader.load_channel("team_id", "channel_id", days=7)
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        tenant_id: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Teams loader.

        Args:
            client_id: Azure AD application client ID
            client_secret: Azure AD application client secret
            tenant_id: Azure AD tenant ID
        """
        if not MSAL_AVAILABLE:
            raise ImportError(
                "msal is required for TeamsLoader. " "Install with: pip install msal"
            )

        self._client_id = client_id or os.environ.get("TEAMS_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get("TEAMS_CLIENT_SECRET")
        self._tenant_id = tenant_id or os.environ.get("TEAMS_TENANT_ID")
        self._access_token: Optional[str] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Teams"

    def authenticate(self) -> bool:
        """Authenticate with Microsoft Graph API."""
        try:
            import msal

            authority = f"https://login.microsoftonline.com/{self._tenant_id}"
            app = msal.ConfidentialClientApplication(
                self._client_id,
                authority=authority,
                client_credential=self._client_secret,
            )

            scopes = ["https://graph.microsoft.com/.default"]
            result = app.acquire_token_for_client(scopes=scopes)

            if "access_token" in result:
                self._access_token = result["access_token"]
                self._authenticated = True
                logger.info("Teams authentication successful")
                return True
            else:
                logger.error(f"Teams auth failed: {result.get('error_description')}")
                return False

        except Exception as e:
            logger.error(f"Teams authentication failed: {e}")
            return False

    def _make_request(self, endpoint: str) -> Optional[dict]:
        """Make authenticated request to Graph API."""
        import requests

        headers = {"Authorization": f"Bearer {self._access_token}"}
        response = requests.get(
            f"https://graph.microsoft.com/v1.0{endpoint}",
            headers=headers,
        )

        if response.status_code == 200:
            return response.json()
        else:
            logger.error(f"Graph API error: {response.status_code}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single message by ID.

        Args:
            doc_id: Message ID in format "team_id:channel_id:message_id"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            parts = doc_id.split(":")
            if len(parts) != 3:
                raise ValueError("doc_id must be team_id:channel_id:message_id")

            team_id, channel_id, message_id = parts
            msg = self._make_request(
                f"/teams/{team_id}/channels/{channel_id}/messages/{message_id}"
            )

            if not msg:
                return None

            content = msg.get("body", {}).get("content", "")
            if msg.get("body", {}).get("contentType") == "html":
                content = self._clean_html(content)

            return LoadedDocument(
                content=content,
                source="teams",
                source_id=doc_id,
                filename="",
                created_at=(
                    datetime.fromisoformat(
                        msg["createdDateTime"].replace("Z", "+00:00")
                    )
                    if msg.get("createdDateTime")
                    else None
                ),
                metadata={
                    "from": msg.get("from", {}).get("user", {}).get("displayName"),
                    "team_id": team_id,
                    "channel_id": channel_id,
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Teams message {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load messages from a channel.

        Args:
            folder_path: "team_id:channel_id" format
            recursive: Include replies

        Returns:
            List of LoadedDocument
        """
        parts = folder_path.split(":")
        if len(parts) != 2:
            logger.error("folder_path must be team_id:channel_id")
            return []

        return self.load_channel(parts[0], parts[1], include_replies=recursive)

    def load_channel(
        self,
        team_id: str,
        channel_id: str,
        days: int = 7,
        max_messages: int = 100,
        include_replies: bool = True,
    ) -> list[LoadedDocument]:
        """Load messages from a Teams channel.

        Args:
            team_id: Team ID
            channel_id: Channel ID
            days: Days to look back
            max_messages: Maximum messages
            include_replies: Include thread replies

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            result = self._make_request(
                f"/teams/{team_id}/channels/{channel_id}/messages?$top={max_messages}"
            )

            if not result:
                return docs

            cutoff = datetime.now(timezone.utc) - timedelta(days=days)

            for msg in result.get("value", []):
                created = msg.get("createdDateTime")
                if created:
                    msg_time = datetime.fromisoformat(created.replace("Z", "+00:00"))
                    if msg_time < cutoff:
                        continue

                content = msg.get("body", {}).get("content", "")
                if msg.get("body", {}).get("contentType") == "html":
                    content = self._clean_html(content)

                sender = (
                    msg.get("from", {}).get("user", {}).get("displayName", "Unknown")
                )
                full_content = f"{sender}: {content}"

                # Get replies if requested
                if include_replies:
                    replies = self._make_request(
                        f"/teams/{team_id}/channels/{channel_id}/messages/{msg['id']}/replies"
                    )
                    if replies:
                        for reply in replies.get("value", []):
                            reply_content = reply.get("body", {}).get("content", "")
                            if reply.get("body", {}).get("contentType") == "html":
                                reply_content = self._clean_html(reply_content)
                            reply_sender = (
                                reply.get("from", {})
                                .get("user", {})
                                .get("displayName", "Unknown")
                            )
                            full_content += f"\n  {reply_sender}: {reply_content}"

                docs.append(
                    LoadedDocument(
                        content=full_content,
                        source="teams",
                        source_id=f"{team_id}:{channel_id}:{msg['id']}",
                        filename="",
                        created_at=(
                            datetime.fromisoformat(created.replace("Z", "+00:00"))
                            if created
                            else None
                        ),
                        metadata={
                            "from": sender,
                            "team_id": team_id,
                            "channel_id": channel_id,
                            "reply_count": (
                                len(replies.get("value", []))
                                if include_replies and replies
                                else 0
                            ),
                        },
                    )
                )

        except Exception as e:
            logger.error(f"Failed to load Teams channel: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for messages in Teams.

        Note: Graph API search requires specific permissions.

        Args:
            query: Search query
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        # Teams search via Graph API is complex - simplified implementation
        logger.warning("Teams search requires delegated permissions, returning empty")
        return []

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# JIRA LOADER
# ============================================================================

# Check for jira availability
try:
    from jira import JIRA

    JIRA_AVAILABLE = True
except ImportError:
    JIRA_AVAILABLE = False


class JiraLoader(BaseLoader):
    """Document loader for Atlassian JIRA tickets and comments.

    Load issues, comments, and attachments from JIRA.

    Features:
    - Load issues by JQL query
    - Load issue comments
    - Load attachments
    - Support for custom fields

    Requirements:
        pip install jira

    Environment variables:
        JIRA_URL: JIRA server URL
        JIRA_EMAIL: User email
        JIRA_API_TOKEN: API token

    Example:
        loader = JiraLoader(url="https://company.atlassian.net", email="x", token="y")
        loader.authenticate()
        docs = loader.load_issues("project = PROJ AND updated >= -7d")
    """

    def __init__(
        self,
        url: Optional[str] = None,
        email: Optional[str] = None,
        token: Optional[str] = None,
        **kwargs,
    ):
        """Initialize JIRA loader.

        Args:
            url: JIRA server URL
            email: User email
            token: API token
        """
        if not JIRA_AVAILABLE:
            raise ImportError(
                "jira is required for JiraLoader. " "Install with: pip install jira"
            )

        self._url = url or os.environ.get("JIRA_URL")
        self._email = email or os.environ.get("JIRA_EMAIL")
        self._token = token or os.environ.get("JIRA_API_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "JIRA"

    def authenticate(self) -> bool:
        """Authenticate with JIRA API."""
        try:
            self._client = JIRA(
                server=self._url,
                basic_auth=(self._email, self._token),
            )

            # Test connection
            self._client.myself()
            self._authenticated = True
            logger.info("JIRA authentication successful")
            return True

        except Exception as e:
            logger.error(f"JIRA authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single issue by key.

        Args:
            doc_id: Issue key (e.g., "PROJ-123")

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            issue = self._client.issue(doc_id, expand="renderedFields")

            # Build content
            content_parts = [
                f"# {issue.key}: {issue.fields.summary}",
                "",
                f"**Status:** {issue.fields.status.name}",
                f"**Type:** {issue.fields.issuetype.name}",
                f"**Priority:** {issue.fields.priority.name if issue.fields.priority else 'None'}",
                f"**Assignee:** {issue.fields.assignee.displayName if issue.fields.assignee else 'Unassigned'}",
                f"**Reporter:** {issue.fields.reporter.displayName if issue.fields.reporter else 'Unknown'}",
                "",
                "## Description",
                issue.fields.description or "(No description)",
            ]

            # Add comments
            comments = self._client.comments(doc_id)
            if comments:
                content_parts.append("")
                content_parts.append("## Comments")
                for comment in comments:
                    content_parts.append(
                        f"\n**{comment.author.displayName}** ({comment.created}):"
                    )
                    content_parts.append(comment.body)

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="jira",
                source_id=doc_id,
                filename=f"{doc_id}.md",
                created_at=(
                    datetime.fromisoformat(issue.fields.created.replace("Z", "+00:00"))
                    if issue.fields.created
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(issue.fields.updated.replace("Z", "+00:00"))
                    if issue.fields.updated
                    else None
                ),
                metadata={
                    "key": issue.key,
                    "status": issue.fields.status.name,
                    "type": issue.fields.issuetype.name,
                    "project": issue.fields.project.key,
                    "assignee": (
                        issue.fields.assignee.displayName
                        if issue.fields.assignee
                        else None
                    ),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load JIRA issue {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load issues from a project.

        Args:
            folder_path: Project key or JQL query
            recursive: If True, treats folder_path as JQL

        Returns:
            List of LoadedDocument
        """
        if recursive or " " in folder_path:
            # Treat as JQL
            return self.load_issues(folder_path)
        else:
            # Treat as project key
            return self.load_issues(f"project = {folder_path}")

    def load_issues(
        self, jql: str, max_results: int = 100, include_comments: bool = True
    ) -> list[LoadedDocument]:
        """Load issues matching a JQL query.

        Args:
            jql: JQL query string
            max_results: Maximum issues to return
            include_comments: Include issue comments

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            issues = self._client.search_issues(
                jql,
                maxResults=max_results,
                expand="renderedFields",
            )

            for issue in issues:
                doc = self.load_document(issue.key)
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load JIRA issues: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for issues.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        jql = f'text ~ "{query}"'
        return self.load_issues(jql, max_results=max_results)

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# ASANA LOADER
# ============================================================================

# Check for asana availability
try:
    import asana

    ASANA_AVAILABLE = True
except ImportError:
    ASANA_AVAILABLE = False


class AsanaLoader(BaseLoader):
    """Document loader for Asana tasks and projects.

    Load tasks, projects, and comments from Asana.

    Features:
    - Load tasks from projects
    - Load task comments/stories
    - Search tasks
    - Support for custom fields

    Requirements:
        pip install asana

    Environment variables:
        ASANA_ACCESS_TOKEN: Personal access token

    Example:
        loader = AsanaLoader(token="xxx")
        loader.authenticate()
        docs = loader.load_project("project_gid")
    """

    def __init__(self, token: Optional[str] = None, **kwargs):
        """Initialize Asana loader.

        Args:
            token: Personal access token
        """
        if not ASANA_AVAILABLE:
            raise ImportError(
                "asana is required for AsanaLoader. " "Install with: pip install asana"
            )

        self._token = token or os.environ.get("ASANA_ACCESS_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Asana"

    def authenticate(self) -> bool:
        """Authenticate with Asana API."""
        try:
            self._client = asana.Client.access_token(self._token)

            # Test connection
            me = self._client.users.me()
            self._authenticated = True
            logger.info(f"Asana authentication successful (user: {me['name']})")
            return True

        except Exception as e:
            logger.error(f"Asana authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single task by GID.

        Args:
            doc_id: Task GID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            task = self._client.tasks.get_task(doc_id)

            # Build content
            content_parts = [
                f"# {task['name']}",
                "",
                f"**Status:** {'Complete' if task.get('completed') else 'Incomplete'}",
                f"**Assignee:** {task.get('assignee', {}).get('name', 'Unassigned')}",
                f"**Due:** {task.get('due_on', 'No due date')}",
                "",
                "## Description",
                task.get("notes", "(No description)"),
            ]

            # Add stories (comments)
            stories = list(self._client.stories.get_stories_for_task(doc_id))
            comments = [s for s in stories if s.get("type") == "comment"]
            if comments:
                content_parts.append("")
                content_parts.append("## Comments")
                for story in comments:
                    content_parts.append(
                        f"\n**{story.get('created_by', {}).get('name', 'Unknown')}**:"
                    )
                    content_parts.append(story.get("text", ""))

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="asana",
                source_id=doc_id,
                filename=f"{task['name']}.md",
                created_at=(
                    datetime.fromisoformat(task["created_at"].replace("Z", "+00:00"))
                    if task.get("created_at")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(task["modified_at"].replace("Z", "+00:00"))
                    if task.get("modified_at")
                    else None
                ),
                metadata={
                    "name": task["name"],
                    "completed": task.get("completed", False),
                    "assignee": task.get("assignee", {}).get("name"),
                    "project": (
                        task.get("projects", [{}])[0].get("name")
                        if task.get("projects")
                        else None
                    ),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Asana task {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load tasks from a project.

        Args:
            folder_path: Project GID
            recursive: Include subtasks

        Returns:
            List of LoadedDocument
        """
        return self.load_project(folder_path, include_subtasks=recursive)

    def load_project(
        self,
        project_gid: str,
        include_subtasks: bool = True,
        include_completed: bool = True,
    ) -> list[LoadedDocument]:
        """Load all tasks from a project.

        Args:
            project_gid: Project GID
            include_subtasks: Include subtasks
            include_completed: Include completed tasks

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            tasks = list(self._client.tasks.get_tasks_for_project(project_gid))

            for task_info in tasks:
                if not include_completed and task_info.get("completed"):
                    continue

                doc = self.load_document(task_info["gid"])
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load Asana project {project_gid}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for tasks.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Get workspaces
            workspaces = list(self._client.workspaces.get_workspaces())

            for workspace in workspaces[:1]:  # Search first workspace
                results = list(
                    self._client.tasks.search_tasks_for_workspace(
                        workspace["gid"],
                        {"text": query},
                    )
                )

                for task_info in results[:max_results]:
                    doc = self.load_document(task_info["gid"])
                    if doc:
                        docs.append(doc)
                        if len(docs) >= max_results:
                            break

        except Exception as e:
            logger.error(f"Asana search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# TRELLO LOADER
# ============================================================================

# Check for trello availability
try:
    from trello import TrelloClient

    TRELLO_AVAILABLE = True
except ImportError:
    TRELLO_AVAILABLE = False


class TrelloLoader(BaseLoader):
    """Document loader for Trello boards, lists, and cards.

    Load cards and content from Trello boards.

    Features:
    - Load cards from boards
    - Load card comments
    - Load attachments
    - Search across boards

    Requirements:
        pip install py-trello

    Environment variables:
        TRELLO_API_KEY: API key
        TRELLO_API_SECRET: API secret
        TRELLO_TOKEN: OAuth token

    Example:
        loader = TrelloLoader(api_key="xxx", token="yyy")
        loader.authenticate()
        docs = loader.load_board("board_id")
    """

    def __init__(
        self,
        api_key: Optional[str] = None,
        api_secret: Optional[str] = None,
        token: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Trello loader.

        Args:
            api_key: Trello API key
            api_secret: Trello API secret
            token: OAuth token
        """
        if not TRELLO_AVAILABLE:
            raise ImportError(
                "py-trello is required for TrelloLoader. "
                "Install with: pip install py-trello"
            )

        self._api_key = api_key or os.environ.get("TRELLO_API_KEY")
        self._api_secret = api_secret or os.environ.get("TRELLO_API_SECRET")
        self._token = token or os.environ.get("TRELLO_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Trello"

    def authenticate(self) -> bool:
        """Authenticate with Trello API."""
        try:
            self._client = TrelloClient(
                api_key=self._api_key,
                api_secret=self._api_secret,
                token=self._token,
            )

            # Test connection
            self._client.list_boards()
            self._authenticated = True
            logger.info("Trello authentication successful")
            return True

        except Exception as e:
            logger.error(f"Trello authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single card by ID.

        Args:
            doc_id: Card ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            card = self._client.get_card(doc_id)

            # Build content
            content_parts = [
                f"# {card.name}",
                "",
                f"**List:** {card.get_list().name}",
                f"**Labels:** {', '.join([l.name for l in card.labels]) if card.labels else 'None'}",
                f"**Due:** {card.due_date or 'No due date'}",
                "",
                "## Description",
                card.description or "(No description)",
            ]

            # Add comments
            comments = card.get_comments()
            if comments:
                content_parts.append("")
                content_parts.append("## Comments")
                for comment in comments:
                    content_parts.append(
                        f"\n**{comment.get('memberCreator', {}).get('fullName', 'Unknown')}**:"
                    )
                    content_parts.append(comment.get("data", {}).get("text", ""))

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="trello",
                source_id=doc_id,
                filename=f"{card.name}.md",
                created_at=card.created_date,
                modified_at=(
                    datetime.fromisoformat(card.dateLastActivity.replace("Z", "+00:00"))
                    if card.dateLastActivity
                    else None
                ),
                metadata={
                    "name": card.name,
                    "list": card.get_list().name,
                    "board": card.board.name,
                    "labels": [l.name for l in card.labels] if card.labels else [],
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Trello card {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load cards from a board.

        Args:
            folder_path: Board ID
            recursive: Include archived cards

        Returns:
            List of LoadedDocument
        """
        return self.load_board(folder_path, include_archived=recursive)

    def load_board(
        self, board_id: str, include_archived: bool = False
    ) -> list[LoadedDocument]:
        """Load all cards from a board.

        Args:
            board_id: Board ID
            include_archived: Include archived cards

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            board = self._client.get_board(board_id)
            cards = board.all_cards() if include_archived else board.open_cards()

            for card in cards:
                doc = self.load_document(card.id)
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load Trello board {board_id}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for cards.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            results = self._client.search(
                query, partial_match=True, cards_limit=max_results
            )

            for card in results.cards:
                doc = self.load_document(card.id)
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Trello search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# AIRTABLE LOADER
# ============================================================================

# Check for airtable availability
try:
    from pyairtable import Api as AirtableApi

    AIRTABLE_AVAILABLE = True
except ImportError:
    AIRTABLE_AVAILABLE = False


class AirtableLoader(BaseLoader):
    """Document loader for Airtable bases and tables.

    Load records from Airtable spreadsheet databases.

    Features:
    - Load records from tables
    - Filter by formula
    - Support for linked records
    - Field formatting

    Requirements:
        pip install pyairtable

    Environment variables:
        AIRTABLE_API_KEY: API key

    Example:
        loader = AirtableLoader(api_key="xxx")
        loader.authenticate()
        docs = loader.load_table("base_id", "table_name")
    """

    def __init__(self, api_key: Optional[str] = None, **kwargs):
        """Initialize Airtable loader.

        Args:
            api_key: Airtable API key
        """
        if not AIRTABLE_AVAILABLE:
            raise ImportError(
                "pyairtable is required for AirtableLoader. "
                "Install with: pip install pyairtable"
            )

        self._api_key = api_key or os.environ.get("AIRTABLE_API_KEY")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Airtable"

    def authenticate(self) -> bool:
        """Authenticate with Airtable API."""
        try:
            self._client = AirtableApi(self._api_key)
            self._authenticated = True
            logger.info("Airtable authentication successful")
            return True

        except Exception as e:
            logger.error(f"Airtable authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single record by ID.

        Args:
            doc_id: Record ID in format "base_id:table_name:record_id"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            parts = doc_id.split(":")
            if len(parts) != 3:
                raise ValueError("doc_id must be base_id:table_name:record_id")

            base_id, table_name, record_id = parts
            table = self._client.table(base_id, table_name)
            record = table.get(record_id)

            # Format fields as content
            content_parts = [f"# Record: {record_id}", ""]
            for field, value in record.get("fields", {}).items():
                content_parts.append(f"**{field}:** {value}")

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="airtable",
                source_id=doc_id,
                filename=f"{record_id}.md",
                created_at=(
                    datetime.fromisoformat(record["createdTime"].replace("Z", "+00:00"))
                    if record.get("createdTime")
                    else None
                ),
                metadata={
                    "base_id": base_id,
                    "table": table_name,
                    "fields": list(record.get("fields", {}).keys()),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Airtable record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a table.

        Args:
            folder_path: "base_id:table_name" format
            recursive: Not used for Airtable

        Returns:
            List of LoadedDocument
        """
        parts = folder_path.split(":")
        if len(parts) != 2:
            logger.error("folder_path must be base_id:table_name")
            return []

        return self.load_table(parts[0], parts[1])

    def load_table(
        self,
        base_id: str,
        table_name: str,
        formula: Optional[str] = None,
        max_records: int = 100,
    ) -> list[LoadedDocument]:
        """Load records from an Airtable table.

        Args:
            base_id: Base ID
            table_name: Table name
            formula: Filter formula
            max_records: Maximum records to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            table = self._client.table(base_id, table_name)

            kwargs = {"max_records": max_records}
            if formula:
                kwargs["formula"] = formula

            records = table.all(**kwargs)

            for record in records:
                doc = self.load_document(f"{base_id}:{table_name}:{record['id']}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load Airtable table: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search is not directly supported - use formula filter instead."""
        logger.warning(
            "Airtable doesn't support full-text search. Use load_table with formula."
        )
        return []

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# HUBSPOT LOADER
# ============================================================================

# Check for hubspot availability
try:
    from hubspot import HubSpot

    HUBSPOT_AVAILABLE = True
except ImportError:
    HUBSPOT_AVAILABLE = False


class HubSpotLoader(BaseLoader):
    """Document loader for HubSpot CRM.

    Load contacts, deals, and tickets from HubSpot.

    Features:
    - Load contacts, companies, deals
    - Load tickets
    - Search across CRM
    - Support for custom properties

    Requirements:
        pip install hubspot-api-client

    Environment variables:
        HUBSPOT_API_KEY: API key

    Example:
        loader = HubSpotLoader(api_key="xxx")
        loader.authenticate()
        docs = loader.load_contacts(max_results=100)
    """

    def __init__(self, api_key: Optional[str] = None, **kwargs):
        """Initialize HubSpot loader.

        Args:
            api_key: HubSpot API key
        """
        if not HUBSPOT_AVAILABLE:
            raise ImportError(
                "hubspot-api-client is required for HubSpotLoader. "
                "Install with: pip install hubspot-api-client"
            )

        self._api_key = api_key or os.environ.get("HUBSPOT_API_KEY")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "HubSpot"

    def authenticate(self) -> bool:
        """Authenticate with HubSpot API."""
        try:
            self._client = HubSpot(api_key=self._api_key)

            # Test connection
            self._client.crm.contacts.basic_api.get_page(limit=1)
            self._authenticated = True
            logger.info("HubSpot authentication successful")
            return True

        except Exception as e:
            logger.error(f"HubSpot authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single CRM object.

        Args:
            doc_id: Object ID in format "type:id" (e.g., "contact:123")

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "contact":
                obj = self._client.crm.contacts.basic_api.get_by_id(obj_id)
            elif obj_type == "company":
                obj = self._client.crm.companies.basic_api.get_by_id(obj_id)
            elif obj_type == "deal":
                obj = self._client.crm.deals.basic_api.get_by_id(obj_id)
            elif obj_type == "ticket":
                obj = self._client.crm.tickets.basic_api.get_by_id(obj_id)
            else:
                raise ValueError(f"Unknown object type: {obj_type}")

            # Format properties as content
            content_parts = [f"# {obj_type.title()}: {obj_id}", ""]
            for prop, value in obj.properties.items():
                if value:
                    content_parts.append(f"**{prop}:** {value}")

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="hubspot",
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.md",
                created_at=(
                    datetime.fromisoformat(obj.created_at.replace("Z", "+00:00"))
                    if hasattr(obj, "created_at") and obj.created_at
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(obj.updated_at.replace("Z", "+00:00"))
                    if hasattr(obj, "updated_at") and obj.updated_at
                    else None
                ),
                metadata={
                    "type": obj_type,
                    "id": obj_id,
                    "properties": list(obj.properties.keys()),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load HubSpot object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load objects of a type.

        Args:
            folder_path: Object type (contacts, companies, deals, tickets)
            recursive: Not used

        Returns:
            List of LoadedDocument
        """
        if folder_path == "contacts":
            return self.load_contacts()
        elif folder_path == "companies":
            return self.load_companies()
        elif folder_path == "deals":
            return self.load_deals()
        elif folder_path == "tickets":
            return self.load_tickets()
        else:
            logger.error(f"Unknown HubSpot object type: {folder_path}")
            return []

    def load_contacts(self, max_results: int = 100) -> list[LoadedDocument]:
        """Load contacts from HubSpot."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.crm.contacts.basic_api.get_page(limit=max_results)

            for contact in response.results:
                doc = self.load_document(f"contact:{contact.id}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load HubSpot contacts: {e}")

        return docs

    def load_companies(self, max_results: int = 100) -> list[LoadedDocument]:
        """Load companies from HubSpot."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.crm.companies.basic_api.get_page(limit=max_results)

            for company in response.results:
                doc = self.load_document(f"company:{company.id}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load HubSpot companies: {e}")

        return docs

    def load_deals(self, max_results: int = 100) -> list[LoadedDocument]:
        """Load deals from HubSpot."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.crm.deals.basic_api.get_page(limit=max_results)

            for deal in response.results:
                doc = self.load_document(f"deal:{deal.id}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load HubSpot deals: {e}")

        return docs

    def load_tickets(self, max_results: int = 100) -> list[LoadedDocument]:
        """Load tickets from HubSpot."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            response = self._client.crm.tickets.basic_api.get_page(limit=max_results)

            for ticket in response.results:
                doc = self.load_document(f"ticket:{ticket.id}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load HubSpot tickets: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search across HubSpot CRM."""
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            from hubspot.crm.contacts import PublicObjectSearchRequest

            search_request = PublicObjectSearchRequest(
                query=query,
                limit=max_results,
            )

            response = self._client.crm.contacts.search_api.do_search(search_request)

            for contact in response.results:
                doc = self.load_document(f"contact:{contact.id}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"HubSpot search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# SALESFORCE LOADER
# ============================================================================

# Check for simple_salesforce availability
try:
    from simple_salesforce import Salesforce

    SALESFORCE_AVAILABLE = True
except ImportError:
    SALESFORCE_AVAILABLE = False


class SalesforceLoader(BaseLoader):
    """Document loader for Salesforce CRM.

    Load records from Salesforce objects.

    Features:
    - Load any Salesforce object
    - SOQL query support
    - Load attachments
    - Support for custom objects

    Requirements:
        pip install simple-salesforce

    Environment variables:
        SALESFORCE_USERNAME: Username
        SALESFORCE_PASSWORD: Password
        SALESFORCE_SECURITY_TOKEN: Security token
        SALESFORCE_DOMAIN: Domain (login or test)

    Example:
        loader = SalesforceLoader(username="x", password="y", security_token="z")
        loader.authenticate()
        docs = loader.load_query("SELECT Id, Name FROM Account LIMIT 10")
    """

    def __init__(
        self,
        username: Optional[str] = None,
        password: Optional[str] = None,
        security_token: Optional[str] = None,
        domain: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Salesforce loader.

        Args:
            username: Salesforce username
            password: Salesforce password
            security_token: Security token
            domain: Domain (login, test, or custom)
        """
        if not SALESFORCE_AVAILABLE:
            raise ImportError(
                "simple-salesforce is required for SalesforceLoader. "
                "Install with: pip install simple-salesforce"
            )

        self._username = username or os.environ.get("SALESFORCE_USERNAME")
        self._password = password or os.environ.get("SALESFORCE_PASSWORD")
        self._security_token = security_token or os.environ.get(
            "SALESFORCE_SECURITY_TOKEN"
        )
        self._domain = domain or os.environ.get("SALESFORCE_DOMAIN", "login")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Salesforce"

    def authenticate(self) -> bool:
        """Authenticate with Salesforce API."""
        try:
            self._client = Salesforce(
                username=self._username,
                password=self._password,
                security_token=self._security_token,
                domain=self._domain,
            )

            self._authenticated = True
            logger.info("Salesforce authentication successful")
            return True

        except Exception as e:
            logger.error(f"Salesforce authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single record.

        Args:
            doc_id: Record ID in format "ObjectName:RecordId"

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            obj_name, record_id = doc_id.split(":", 1)
            obj = getattr(self._client, obj_name)
            record = obj.get(record_id)

            # Format fields as content
            content_parts = [f"# {obj_name}: {record_id}", ""]
            for field, value in record.items():
                if value and field != "attributes":
                    content_parts.append(f"**{field}:** {value}")

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="salesforce",
                source_id=doc_id,
                filename=f"{obj_name}_{record_id}.md",
                created_at=(
                    datetime.fromisoformat(record["CreatedDate"].replace("Z", "+00:00"))
                    if record.get("CreatedDate")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(
                        record["LastModifiedDate"].replace("Z", "+00:00")
                    )
                    if record.get("LastModifiedDate")
                    else None
                ),
                metadata={
                    "object": obj_name,
                    "id": record_id,
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Salesforce record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from an object.

        Args:
            folder_path: Object name (Account, Contact, etc.)
            recursive: Not used

        Returns:
            List of LoadedDocument
        """
        return self.load_query(f"SELECT Id FROM {folder_path} LIMIT 100")

    def load_query(self, soql: str) -> list[LoadedDocument]:
        """Load records matching a SOQL query.

        Args:
            soql: SOQL query

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Extract object name from query
            obj_name = soql.split("FROM")[1].split()[0].strip()

            result = self._client.query(soql)

            for record in result.get("records", []):
                doc = self.load_document(f"{obj_name}:{record['Id']}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to execute Salesforce query: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search using SOSL.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            sosl = f"FIND {{{query}}} IN ALL FIELDS RETURNING Account(Id), Contact(Id), Lead(Id) LIMIT {max_results}"
            result = self._client.search(sosl)

            for record in result.get("searchRecords", []):
                obj_type = record["attributes"]["type"]
                doc = self.load_document(f"{obj_type}:{record['Id']}")
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Salesforce search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# ZENDESK LOADER
# ============================================================================

# Check for zenpy availability
try:
    from zenpy import Zenpy

    ZENDESK_AVAILABLE = True
except ImportError:
    ZENDESK_AVAILABLE = False


class ZendeskLoader(BaseLoader):
    """Document loader for Zendesk support tickets.

    Load tickets and comments from Zendesk.

    Features:
    - Load tickets and comments
    - Search tickets
    - Load ticket attachments
    - Support for custom fields

    Requirements:
        pip install zenpy

    Environment variables:
        ZENDESK_EMAIL: Agent email
        ZENDESK_TOKEN: API token
        ZENDESK_SUBDOMAIN: Zendesk subdomain

    Example:
        loader = ZendeskLoader(email="x", token="y", subdomain="company")
        loader.authenticate()
        docs = loader.load_tickets(status="open")
    """

    def __init__(
        self,
        email: Optional[str] = None,
        token: Optional[str] = None,
        subdomain: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Zendesk loader.

        Args:
            email: Agent email
            token: API token
            subdomain: Zendesk subdomain
        """
        if not ZENDESK_AVAILABLE:
            raise ImportError(
                "zenpy is required for ZendeskLoader. "
                "Install with: pip install zenpy"
            )

        self._email = email or os.environ.get("ZENDESK_EMAIL")
        self._token = token or os.environ.get("ZENDESK_TOKEN")
        self._subdomain = subdomain or os.environ.get("ZENDESK_SUBDOMAIN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Zendesk"

    def authenticate(self) -> bool:
        """Authenticate with Zendesk API."""
        try:
            creds = {
                "email": self._email,
                "token": self._token,
                "subdomain": self._subdomain,
            }
            self._client = Zenpy(**creds)

            # Test connection
            self._client.users.me()
            self._authenticated = True
            logger.info("Zendesk authentication successful")
            return True

        except Exception as e:
            logger.error(f"Zendesk authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single ticket by ID.

        Args:
            doc_id: Ticket ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            ticket = self._client.tickets(id=int(doc_id))

            # Build content
            content_parts = [
                f"# Ticket #{ticket.id}: {ticket.subject}",
                "",
                f"**Status:** {ticket.status}",
                f"**Priority:** {ticket.priority}",
                f"**Requester:** {ticket.requester.name if ticket.requester else 'Unknown'}",
                f"**Assignee:** {ticket.assignee.name if ticket.assignee else 'Unassigned'}",
                "",
                "## Description",
                ticket.description or "(No description)",
            ]

            # Add comments
            comments = list(self._client.tickets.comments(ticket_id=int(doc_id)))
            if len(comments) > 1:  # First comment is description
                content_parts.append("")
                content_parts.append("## Comments")
                for comment in comments[1:]:
                    content_parts.append(
                        f"\n**{comment.author.name if comment.author else 'Unknown'}**:"
                    )
                    content_parts.append(comment.body)

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="zendesk",
                source_id=doc_id,
                filename=f"ticket_{doc_id}.md",
                created_at=ticket.created_at,
                modified_at=ticket.updated_at,
                metadata={
                    "id": ticket.id,
                    "status": ticket.status,
                    "priority": ticket.priority,
                    "type": ticket.type,
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Zendesk ticket {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load tickets with a status filter.

        Args:
            folder_path: Status filter (open, pending, solved, all)
            recursive: Not used

        Returns:
            List of LoadedDocument
        """
        return self.load_tickets(status=folder_path if folder_path != "all" else None)

    def load_tickets(
        self, status: Optional[str] = None, max_results: int = 100
    ) -> list[LoadedDocument]:
        """Load tickets from Zendesk.

        Args:
            status: Filter by status (new, open, pending, hold, solved, closed)
            max_results: Maximum tickets to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            if status:
                tickets = self._client.search(type="ticket", status=status)
            else:
                tickets = self._client.tickets()

            count = 0
            for ticket in tickets:
                if count >= max_results:
                    break

                doc = self.load_document(str(ticket.id))
                if doc:
                    docs.append(doc)
                    count += 1

        except Exception as e:
            logger.error(f"Failed to load Zendesk tickets: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for tickets.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            results = self._client.search(query, type="ticket")

            count = 0
            for ticket in results:
                if count >= max_results:
                    break

                doc = self.load_document(str(ticket.id))
                if doc:
                    docs.append(doc)
                    count += 1

        except Exception as e:
            logger.error(f"Zendesk search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# INTERCOM LOADER
# ============================================================================

# Check for intercom availability
try:
    import intercom
    from intercom.client import Client as IntercomClient

    INTERCOM_AVAILABLE = True
except ImportError:
    INTERCOM_AVAILABLE = False


class IntercomLoader(BaseLoader):
    """Document loader for Intercom conversations.

    Load conversations and messages from Intercom.

    Features:
    - Load conversations
    - Load conversation messages
    - Search conversations
    - Support for contacts

    Requirements:
        pip install python-intercom

    Environment variables:
        INTERCOM_ACCESS_TOKEN: Access token

    Example:
        loader = IntercomLoader(token="xxx")
        loader.authenticate()
        docs = loader.load_conversations()
    """

    def __init__(self, token: Optional[str] = None, **kwargs):
        """Initialize Intercom loader.

        Args:
            token: Intercom access token
        """
        if not INTERCOM_AVAILABLE:
            raise ImportError(
                "python-intercom is required for IntercomLoader. "
                "Install with: pip install python-intercom"
            )

        self._token = token or os.environ.get("INTERCOM_ACCESS_TOKEN")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Intercom"

    def authenticate(self) -> bool:
        """Authenticate with Intercom API."""
        try:
            self._client = IntercomClient(personal_access_token=self._token)

            # Test connection
            self._client.admins.all()
            self._authenticated = True
            logger.info("Intercom authentication successful")
            return True

        except Exception as e:
            logger.error(f"Intercom authentication failed: {e}")
            return False

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single conversation by ID.

        Args:
            doc_id: Conversation ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            conversation = self._client.conversations.find(id=doc_id)

            # Build content from conversation parts
            content_parts = [f"# Conversation: {doc_id}", ""]

            # Add source (first message)
            if conversation.source:
                author = (
                    conversation.source.author.name
                    if hasattr(conversation.source.author, "name")
                    else "Unknown"
                )
                content_parts.append(f"**{author}:** {conversation.source.body}")

            # Add conversation parts
            for part in conversation.conversation_parts:
                author = part.author.name if hasattr(part.author, "name") else "Unknown"
                content_parts.append(f"\n**{author}:** {part.body}")

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="intercom",
                source_id=doc_id,
                filename=f"conversation_{doc_id}.md",
                created_at=(
                    datetime.fromtimestamp(conversation.created_at)
                    if conversation.created_at
                    else None
                ),
                modified_at=(
                    datetime.fromtimestamp(conversation.updated_at)
                    if conversation.updated_at
                    else None
                ),
                metadata={
                    "id": doc_id,
                    "state": conversation.state,
                    "read": conversation.read,
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Intercom conversation {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load conversations with a state filter.

        Args:
            folder_path: State filter (open, closed, all)
            recursive: Not used

        Returns:
            List of LoadedDocument
        """
        return self.load_conversations(
            state=folder_path if folder_path != "all" else None
        )

    def load_conversations(
        self, state: Optional[str] = None, max_results: int = 100
    ) -> list[LoadedDocument]:
        """Load conversations from Intercom.

        Args:
            state: Filter by state (open, closed)
            max_results: Maximum conversations to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            conversations = self._client.conversations.all()

            count = 0
            for conv in conversations:
                if count >= max_results:
                    break

                if state and conv.state != state:
                    continue

                doc = self.load_document(conv.id)
                if doc:
                    docs.append(doc)
                    count += 1

        except Exception as e:
            logger.error(f"Failed to load Intercom conversations: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for conversations.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        # Intercom API doesn't support conversation search directly
        # Load all and filter locally
        logger.warning("Intercom doesn't support conversation search, loading all")
        return self.load_conversations(max_results=max_results)

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# FRESHDESK LOADER
# ============================================================================

# Check for freshdesk availability
try:
    import requests

    FRESHDESK_AVAILABLE = True  # Uses requests directly
except ImportError:
    FRESHDESK_AVAILABLE = False


class FreshdeskLoader(BaseLoader):
    """Document loader for Freshdesk support tickets.

    Load tickets and conversations from Freshdesk.

    Features:
    - Load tickets and conversations
    - Search tickets
    - Load attachments
    - Support for custom fields

    Requirements:
        pip install requests

    Environment variables:
        FRESHDESK_DOMAIN: Freshdesk domain (company.freshdesk.com)
        FRESHDESK_API_KEY: API key

    Example:
        loader = FreshdeskLoader(domain="company", api_key="xxx")
        loader.authenticate()
        docs = loader.load_tickets(status="open")
    """

    def __init__(
        self,
        domain: Optional[str] = None,
        api_key: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Freshdesk loader.

        Args:
            domain: Freshdesk domain (without .freshdesk.com)
            api_key: API key
        """
        self._domain = domain or os.environ.get("FRESHDESK_DOMAIN")
        self._api_key = api_key or os.environ.get("FRESHDESK_API_KEY")
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "Freshdesk"

    def authenticate(self) -> bool:
        """Authenticate with Freshdesk API."""
        try:
            import requests

            response = requests.get(
                f"https://{self._domain}.freshdesk.com/api/v2/tickets",
                auth=(self._api_key, "X"),
                params={"per_page": 1},
            )

            if response.status_code == 200:
                self._authenticated = True
                logger.info("Freshdesk authentication successful")
                return True
            else:
                logger.error(f"Freshdesk auth failed: {response.status_code}")
                return False

        except Exception as e:
            logger.error(f"Freshdesk authentication failed: {e}")
            return False

    def _make_request(
        self, endpoint: str, params: Optional[dict] = None
    ) -> Optional[Any]:
        """Make authenticated request to Freshdesk API."""
        import requests

        response = requests.get(
            f"https://{self._domain}.freshdesk.com/api/v2{endpoint}",
            auth=(self._api_key, "X"),
            params=params,
        )

        if response.status_code == 200:
            return response.json()
        else:
            logger.error(f"Freshdesk API error: {response.status_code}")
            return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single ticket by ID.

        Args:
            doc_id: Ticket ID

        Returns:
            LoadedDocument or None
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        try:
            ticket = self._make_request(f"/tickets/{doc_id}")
            if not ticket:
                return None

            # Status mapping
            status_map = {2: "Open", 3: "Pending", 4: "Resolved", 5: "Closed"}
            priority_map = {1: "Low", 2: "Medium", 3: "High", 4: "Urgent"}

            # Build content
            content_parts = [
                f"# Ticket #{ticket['id']}: {ticket['subject']}",
                "",
                f"**Status:** {status_map.get(ticket['status'], 'Unknown')}",
                f"**Priority:** {priority_map.get(ticket['priority'], 'Unknown')}",
                "",
                "## Description",
                ticket.get(
                    "description_text", ticket.get("description", "(No description)")
                ),
            ]

            # Get conversations
            conversations = self._make_request(f"/tickets/{doc_id}/conversations")
            if conversations:
                content_parts.append("")
                content_parts.append("## Conversations")
                for conv in conversations:
                    content_parts.append(f"\n**Message:**")
                    content_parts.append(conv.get("body_text", conv.get("body", "")))

            content = "\n".join(content_parts)

            return LoadedDocument(
                content=content,
                source="freshdesk",
                source_id=doc_id,
                filename=f"ticket_{doc_id}.md",
                created_at=(
                    datetime.fromisoformat(ticket["created_at"].replace("Z", "+00:00"))
                    if ticket.get("created_at")
                    else None
                ),
                modified_at=(
                    datetime.fromisoformat(ticket["updated_at"].replace("Z", "+00:00"))
                    if ticket.get("updated_at")
                    else None
                ),
                metadata={
                    "id": ticket["id"],
                    "status": status_map.get(ticket["status"], "Unknown"),
                    "priority": priority_map.get(ticket["priority"], "Unknown"),
                },
            )

        except Exception as e:
            logger.error(f"Failed to load Freshdesk ticket {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load tickets with a status filter.

        Args:
            folder_path: Status filter (open, pending, resolved, closed, all)
            recursive: Not used

        Returns:
            List of LoadedDocument
        """
        status_map = {"open": 2, "pending": 3, "resolved": 4, "closed": 5}
        status = (
            status_map.get(folder_path.lower())
            if folder_path.lower() != "all"
            else None
        )
        return self.load_tickets(status=status)

    def load_tickets(
        self, status: Optional[int] = None, max_results: int = 100
    ) -> list[LoadedDocument]:
        """Load tickets from Freshdesk.

        Args:
            status: Filter by status (2=Open, 3=Pending, 4=Resolved, 5=Closed)
            max_results: Maximum tickets to return

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            params = {"per_page": min(max_results, 100)}
            if status:
                params["filter"] = f"status:{status}"

            tickets = self._make_request("/tickets", params=params)
            if not tickets:
                return docs

            for ticket in tickets[:max_results]:
                doc = self.load_document(str(ticket["id"]))
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load Freshdesk tickets: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for tickets.

        Args:
            query: Search text
            max_results: Maximum results

        Returns:
            List of LoadedDocument
        """
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

        docs = []

        try:
            # Freshdesk search endpoint
            results = self._make_request(
                "/search/tickets", params={"query": f'"{query}"'}
            )
            if not results or "results" not in results:
                return docs

            for ticket in results["results"][:max_results]:
                doc = self.load_document(str(ticket["id"]))
                if doc:
                    docs.append(doc)

        except Exception as e:
            logger.error(f"Freshdesk search failed: {e}")

        return docs

    async def load_async(self, doc_id: str) -> Optional[LoadedDocument]:
        """Async document loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load_document, doc_id)

    async def load_folder_async(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Async folder loading."""
        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self.load_folder, folder_path, recursive
        )


# ============================================================================
# CLOUD STORAGE LOADERS
# ============================================================================

# Azure Blob Storage
try:
    from azure.storage.blob import BlobServiceClient

    AZURE_BLOB_AVAILABLE = True
except ImportError:
    AZURE_BLOB_AVAILABLE = False


class AzureBlobLoader(BaseLoader):
    """Load documents from Azure Blob Storage.

    Supports loading files from Azure Blob containers with text extraction
    for common document formats.

    Environment Variables:
        AZURE_STORAGE_CONNECTION_STRING: Full connection string
        AZURE_STORAGE_ACCOUNT_NAME: Account name (with AZURE_STORAGE_ACCOUNT_KEY)
        AZURE_STORAGE_ACCOUNT_KEY: Account key

    Example:
        loader = AzureBlobLoader(
            connection_string="DefaultEndpointsProtocol=https;AccountName=...",
            container="documents"
        )
        docs = loader.load_folder("reports/2024/")
    """

    def __init__(
        self,
        container: str,
        connection_string: Optional[str] = None,
        account_name: Optional[str] = None,
        account_key: Optional[str] = None,
        max_file_size_mb: int = 50,
    ):
        if not AZURE_BLOB_AVAILABLE:
            raise ImportError(
                "Azure SDK not installed. Run: pip install azure-storage-blob"
            )

        self.container_name = container
        self.connection_string = connection_string or os.environ.get(
            "AZURE_STORAGE_CONNECTION_STRING"
        )
        self.account_name = account_name or os.environ.get("AZURE_STORAGE_ACCOUNT_NAME")
        self.account_key = account_key or os.environ.get("AZURE_STORAGE_ACCOUNT_KEY")
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self._client = None
        self._container_client = None

    @property
    def source_name(self) -> str:
        return "azure_blob"

    def authenticate(self) -> bool:
        """Authenticate with Azure Blob Storage."""
        try:
            if self.connection_string:
                self._client = BlobServiceClient.from_connection_string(
                    self.connection_string
                )
            elif self.account_name and self.account_key:
                account_url = f"https://{self.account_name}.blob.core.windows.net"
                self._client = BlobServiceClient(
                    account_url=account_url, credential=self.account_key
                )
            else:
                raise ValueError(
                    "Provide connection_string or account_name/account_key"
                )

            self._container_client = self._client.get_container_client(
                self.container_name
            )
            # Test connection
            self._container_client.exists()
            logger.info(
                f"Azure Blob authentication successful for {self.container_name}"
            )
            return True
        except Exception as e:
            logger.error(f"Azure Blob authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._container_client and not self.authenticate():
            raise RuntimeError("Failed to authenticate with Azure Blob Storage")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single blob by path."""
        self._ensure_authenticated()
        try:
            blob_client = self._container_client.get_blob_client(doc_id)
            props = blob_client.get_blob_properties()

            if props.size > self.max_file_size:
                logger.warning(f"Skipping large blob: {doc_id}")
                return None

            content_bytes = blob_client.download_blob().readall()
            mime_type, _ = mimetypes.guess_type(doc_id)
            mime_type = mime_type or "application/octet-stream"

            # Extract text based on mime type
            if mime_type == "application/pdf":
                content = self._extract_text_from_pdf(content_bytes)
            elif mime_type.startswith("text/"):
                content = content_bytes.decode("utf-8", errors="replace")
            elif mime_type == "application/json":
                content = content_bytes.decode("utf-8", errors="replace")
            else:
                logger.debug(f"Skipping unsupported type: {mime_type}")
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=Path(doc_id).name,
                mime_type=mime_type,
                created_at=props.creation_time,
                modified_at=props.last_modified,
                size_bytes=props.size,
                metadata={"container": self.container_name, "blob_path": doc_id},
            )
        except Exception as e:
            logger.error(f"Failed to load blob {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a blob prefix."""
        self._ensure_authenticated()
        docs = []
        prefix = folder_path.strip("/") + "/" if folder_path else ""

        try:
            blobs = self._container_client.list_blobs(name_starts_with=prefix)
            for blob in blobs:
                if not recursive and "/" in blob.name[len(prefix) :]:
                    continue
                doc = self.load_document(blob.name)
                if doc:
                    docs.append(doc)
        except Exception as e:
            logger.error(f"Failed to load folder {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search blobs by name pattern."""
        self._ensure_authenticated()
        docs = []
        count = 0

        try:
            blobs = self._container_client.list_blobs()
            for blob in blobs:
                if query.lower() in blob.name.lower():
                    doc = self.load_document(blob.name)
                    if doc:
                        docs.append(doc)
                        count += 1
                        if count >= max_results:
                            break
        except Exception as e:
            logger.error(f"Azure Blob search failed: {e}")

        return docs


# Google Cloud Storage
try:
    from google.cloud import storage as gcs_storage

    GCS_AVAILABLE = True
except ImportError:
    GCS_AVAILABLE = False


class GCSLoader(BaseLoader):
    """Load documents from Google Cloud Storage.

    Supports loading files from GCS buckets with text extraction.

    Environment Variables:
        GOOGLE_APPLICATION_CREDENTIALS: Path to service account JSON

    Example:
        loader = GCSLoader(
            bucket="my-documents",
            credentials_path="service-account.json"
        )
        docs = loader.load_folder("reports/")
    """

    def __init__(
        self,
        bucket: str,
        credentials_path: Optional[str] = None,
        max_file_size_mb: int = 50,
    ):
        if not GCS_AVAILABLE:
            raise ImportError(
                "Google Cloud Storage SDK not installed. Run: pip install google-cloud-storage"
            )

        self.bucket_name = bucket
        self.credentials_path = credentials_path
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self._client = None
        self._bucket = None

    @property
    def source_name(self) -> str:
        return "gcs"

    def authenticate(self) -> bool:
        """Authenticate with Google Cloud Storage."""
        try:
            if self.credentials_path:
                self._client = gcs_storage.Client.from_service_account_json(
                    self.credentials_path
                )
            else:
                self._client = gcs_storage.Client()

            self._bucket = self._client.bucket(self.bucket_name)
            # Test connection
            self._bucket.exists()
            logger.info(f"GCS authentication successful for {self.bucket_name}")
            return True
        except Exception as e:
            logger.error(f"GCS authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._bucket and not self.authenticate():
            raise RuntimeError("Failed to authenticate with GCS")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single object by path."""
        self._ensure_authenticated()
        try:
            blob = self._bucket.blob(doc_id)
            blob.reload()

            if blob.size > self.max_file_size:
                logger.warning(f"Skipping large object: {doc_id}")
                return None

            content_bytes = blob.download_as_bytes()
            mime_type = blob.content_type or "application/octet-stream"

            if mime_type == "application/pdf":
                content = self._extract_text_from_pdf(content_bytes)
            elif mime_type.startswith("text/") or mime_type == "application/json":
                content = content_bytes.decode("utf-8", errors="replace")
            else:
                logger.debug(f"Skipping unsupported type: {mime_type}")
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=Path(doc_id).name,
                mime_type=mime_type,
                created_at=blob.time_created,
                modified_at=blob.updated,
                size_bytes=blob.size,
                metadata={"bucket": self.bucket_name, "object_path": doc_id},
            )
        except Exception as e:
            logger.error(f"Failed to load object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from a prefix."""
        self._ensure_authenticated()
        docs = []
        prefix = folder_path.strip("/") + "/" if folder_path else ""
        delimiter = None if recursive else "/"

        try:
            blobs = self._client.list_blobs(
                self.bucket_name, prefix=prefix, delimiter=delimiter
            )
            for blob in blobs:
                doc = self.load_document(blob.name)
                if doc:
                    docs.append(doc)
        except Exception as e:
            logger.error(f"Failed to load folder {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search objects by name pattern."""
        self._ensure_authenticated()
        docs = []
        count = 0

        try:
            blobs = self._client.list_blobs(self.bucket_name)
            for blob in blobs:
                if query.lower() in blob.name.lower():
                    doc = self.load_document(blob.name)
                    if doc:
                        docs.append(doc)
                        count += 1
                        if count >= max_results:
                            break
        except Exception as e:
            logger.error(f"GCS search failed: {e}")

        return docs


# MinIO (S3-compatible)
class MinIOLoader(S3Loader):
    """Load documents from MinIO (self-hosted S3-compatible storage).

    MinIO is fully S3-compatible, so this extends S3Loader with
    convenient defaults for self-hosted deployments.

    Example:
        loader = MinIOLoader(
            bucket="documents",
            endpoint_url="http://minio.local:9000",
            access_key="minioadmin",
            secret_key="minioadmin"
        )
        docs = loader.load_folder("reports/")
    """

    def __init__(
        self,
        bucket: str,
        endpoint_url: str = "http://localhost:9000",
        access_key: Optional[str] = None,
        secret_key: Optional[str] = None,
        secure: bool = False,
        **kwargs,
    ):
        # MinIO uses endpoint_url, map to S3Loader params
        super().__init__(
            bucket=bucket,
            endpoint_url=endpoint_url,
            aws_access_key_id=access_key or os.environ.get("MINIO_ACCESS_KEY"),
            aws_secret_access_key=secret_key or os.environ.get("MINIO_SECRET_KEY"),
            **kwargs,
        )

    @property
    def source_name(self) -> str:
        return "minio"


# ============================================================================
# DATABASE LOADERS
# ============================================================================

# PostgreSQL
try:
    import psycopg2

    POSTGRES_AVAILABLE = True
except ImportError:
    POSTGRES_AVAILABLE = False


class PostgreSQLLoader(BaseLoader):
    """Load documents from PostgreSQL database.

    Extracts text content from tables for RAG ingestion.

    Example:
        loader = PostgreSQLLoader(
            host="localhost",
            database="knowledge_base",
            user="postgres",
            password="secret"
        )
        docs = loader.load_folder("articles")  # table name
    """

    def __init__(
        self,
        host: str = "localhost",
        port: int = 5432,
        database: str = "postgres",
        user: str = "postgres",
        password: Optional[str] = None,
        connection_string: Optional[str] = None,
        content_column: str = "content",
        id_column: str = "id",
        metadata_columns: Optional[list[str]] = None,
    ):
        if not POSTGRES_AVAILABLE:
            raise ImportError(
                "psycopg2 not installed. Run: pip install psycopg2-binary"
            )

        self.host = host
        self.port = port
        self.database = database
        self.user = user
        self.password = password or os.environ.get("POSTGRES_PASSWORD")
        self.connection_string = connection_string or os.environ.get("DATABASE_URL")
        self.content_column = content_column
        self.id_column = id_column
        self.metadata_columns = metadata_columns or []
        self._conn = None

    @property
    def source_name(self) -> str:
        return "postgresql"

    def authenticate(self) -> bool:
        """Connect to PostgreSQL."""
        try:
            if self.connection_string:
                self._conn = psycopg2.connect(self.connection_string)
            else:
                self._conn = psycopg2.connect(
                    host=self.host,
                    port=self.port,
                    database=self.database,
                    user=self.user,
                    password=self.password,
                )
            logger.info(f"PostgreSQL connection successful to {self.database}")
            return True
        except Exception as e:
            logger.error(f"PostgreSQL connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._conn and not self.authenticate():
            raise RuntimeError("Failed to connect to PostgreSQL")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single row by ID. doc_id format: table_name/row_id"""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            table = parts[0]
            row_id = parts[1] if len(parts) > 1 else doc_id

            columns = [self.id_column, self.content_column] + self.metadata_columns
            columns_str = ", ".join(columns)

            with self._conn.cursor() as cur:
                cur.execute(
                    f"SELECT {columns_str} FROM {table} WHERE {self.id_column} = %s",
                    (row_id,),
                )
                row = cur.fetchone()

            if not row:
                return None

            content = str(row[1]) if row[1] else ""
            metadata = {}
            for i, col in enumerate(self.metadata_columns):
                metadata[col] = row[i + 2]

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=f"{table}_{row_id}",
                metadata={"table": table, "row_id": row_id, **metadata},
            )
        except Exception as e:
            logger.error(f"Failed to load row {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all rows from a table. folder_path is the table name."""
        self._ensure_authenticated()
        docs = []
        table = folder_path

        try:
            columns = [self.id_column, self.content_column] + self.metadata_columns
            columns_str = ", ".join(columns)

            with self._conn.cursor() as cur:
                cur.execute(f"SELECT {columns_str} FROM {table}")
                rows = cur.fetchall()

            for row in rows:
                row_id = str(row[0])
                content = str(row[1]) if row[1] else ""
                metadata = {}
                for i, col in enumerate(self.metadata_columns):
                    metadata[col] = row[i + 2]

                docs.append(
                    LoadedDocument(
                        content=content,
                        source=self.source_name,
                        source_id=f"{table}/{row_id}",
                        filename=f"{table}_{row_id}",
                        metadata={"table": table, "row_id": row_id, **metadata},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load table {table}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for text in content column."""
        # Note: Requires table context, use load_folder with WHERE clause for real search
        logger.warning(
            "PostgreSQL search requires specifying table. Use load_with_query instead."
        )
        return []

    def load_with_query(
        self, sql: str, params: Optional[tuple] = None
    ) -> list[LoadedDocument]:
        """Load documents using custom SQL query.

        Query must return columns in order: id, content, [metadata_columns...]

        Example:
            docs = loader.load_with_query(
                "SELECT id, body, title, author FROM articles WHERE category = %s",
                ("tech",)
            )
        """
        self._ensure_authenticated()
        docs = []

        try:
            with self._conn.cursor() as cur:
                cur.execute(sql, params or ())
                rows = cur.fetchall()

            for row in rows:
                row_id = str(row[0])
                content = str(row[1]) if row[1] else ""
                metadata = {}
                for i, col in enumerate(self.metadata_columns):
                    if i + 2 < len(row):
                        metadata[col] = row[i + 2]

                docs.append(
                    LoadedDocument(
                        content=content,
                        source=self.source_name,
                        source_id=row_id,
                        filename=f"query_{row_id}",
                        metadata={"row_id": row_id, **metadata},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to execute query: {e}")

        return docs


# MySQL
try:
    import mysql.connector

    MYSQL_AVAILABLE = True
except ImportError:
    MYSQL_AVAILABLE = False


class MySQLLoader(BaseLoader):
    """Load documents from MySQL database.

    Example:
        loader = MySQLLoader(
            host="localhost",
            database="knowledge_base",
            user="root",
            password="secret"
        )
        docs = loader.load_folder("articles")  # table name
    """

    def __init__(
        self,
        host: str = "localhost",
        port: int = 3306,
        database: str = "mysql",
        user: str = "root",
        password: Optional[str] = None,
        content_column: str = "content",
        id_column: str = "id",
        metadata_columns: Optional[list[str]] = None,
    ):
        if not MYSQL_AVAILABLE:
            raise ImportError(
                "mysql-connector-python not installed. Run: pip install mysql-connector-python"
            )

        self.host = host
        self.port = port
        self.database = database
        self.user = user
        self.password = password or os.environ.get("MYSQL_PASSWORD")
        self.content_column = content_column
        self.id_column = id_column
        self.metadata_columns = metadata_columns or []
        self._conn = None

    @property
    def source_name(self) -> str:
        return "mysql"

    def authenticate(self) -> bool:
        """Connect to MySQL."""
        try:
            self._conn = mysql.connector.connect(
                host=self.host,
                port=self.port,
                database=self.database,
                user=self.user,
                password=self.password,
            )
            logger.info(f"MySQL connection successful to {self.database}")
            return True
        except Exception as e:
            logger.error(f"MySQL connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._conn and not self.authenticate():
            raise RuntimeError("Failed to connect to MySQL")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single row by ID."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            table = parts[0]
            row_id = parts[1] if len(parts) > 1 else doc_id

            columns = [self.id_column, self.content_column] + self.metadata_columns
            columns_str = ", ".join(f"`{c}`" for c in columns)

            cursor = self._conn.cursor()
            cursor.execute(
                f"SELECT {columns_str} FROM `{table}` WHERE `{self.id_column}` = %s",
                (row_id,),
            )
            row = cursor.fetchone()
            cursor.close()

            if not row:
                return None

            content = str(row[1]) if row[1] else ""
            metadata = {}
            for i, col in enumerate(self.metadata_columns):
                metadata[col] = row[i + 2]

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=f"{table}_{row_id}",
                metadata={"table": table, "row_id": row_id, **metadata},
            )
        except Exception as e:
            logger.error(f"Failed to load row {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all rows from a table."""
        self._ensure_authenticated()
        docs = []
        table = folder_path

        try:
            columns = [self.id_column, self.content_column] + self.metadata_columns
            columns_str = ", ".join(f"`{c}`" for c in columns)

            cursor = self._conn.cursor()
            cursor.execute(f"SELECT {columns_str} FROM `{table}`")
            rows = cursor.fetchall()
            cursor.close()

            for row in rows:
                row_id = str(row[0])
                content = str(row[1]) if row[1] else ""
                metadata = {}
                for i, col in enumerate(self.metadata_columns):
                    metadata[col] = row[i + 2]

                docs.append(
                    LoadedDocument(
                        content=content,
                        source=self.source_name,
                        source_id=f"{table}/{row_id}",
                        filename=f"{table}_{row_id}",
                        metadata={"table": table, "row_id": row_id, **metadata},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load table {table}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Full-text search (requires FULLTEXT index)."""
        logger.warning("MySQL search requires specifying table and FULLTEXT index.")
        return []


# Elasticsearch
try:
    from elasticsearch import Elasticsearch

    ELASTICSEARCH_AVAILABLE = True
except ImportError:
    ELASTICSEARCH_AVAILABLE = False


class ElasticsearchLoader(BaseLoader):
    """Load documents from Elasticsearch.

    Perfect for existing search indices that need RAG integration.

    Example:
        loader = ElasticsearchLoader(
            hosts=["http://localhost:9200"],
            index="articles",
            content_field="body"
        )
        docs = loader.search("machine learning")
    """

    def __init__(
        self,
        hosts: list[str] = None,
        index: str = "documents",
        content_field: str = "content",
        cloud_id: Optional[str] = None,
        api_key: Optional[str] = None,
        username: Optional[str] = None,
        password: Optional[str] = None,
    ):
        if not ELASTICSEARCH_AVAILABLE:
            raise ImportError(
                "elasticsearch not installed. Run: pip install elasticsearch"
            )

        self.hosts = hosts or ["http://localhost:9200"]
        self.index = index
        self.content_field = content_field
        self.cloud_id = cloud_id
        self.api_key = api_key or os.environ.get("ELASTICSEARCH_API_KEY")
        self.username = username
        self.password = password
        self._client = None

    @property
    def source_name(self) -> str:
        return "elasticsearch"

    def authenticate(self) -> bool:
        """Connect to Elasticsearch."""
        try:
            if self.cloud_id:
                self._client = Elasticsearch(
                    cloud_id=self.cloud_id, api_key=self.api_key
                )
            elif self.api_key:
                self._client = Elasticsearch(hosts=self.hosts, api_key=self.api_key)
            elif self.username and self.password:
                self._client = Elasticsearch(
                    hosts=self.hosts, basic_auth=(self.username, self.password)
                )
            else:
                self._client = Elasticsearch(hosts=self.hosts)

            # Test connection
            self._client.info()
            logger.info("Elasticsearch connection successful")
            return True
        except Exception as e:
            logger.error(f"Elasticsearch connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._client and not self.authenticate():
            raise RuntimeError("Failed to connect to Elasticsearch")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single document by ID."""
        self._ensure_authenticated()
        try:
            result = self._client.get(index=self.index, id=doc_id)
            source = result["_source"]
            content = source.get(self.content_field, "")

            return LoadedDocument(
                content=str(content),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{self.index}_{doc_id}",
                metadata={"index": self.index, **source},
            )
        except Exception as e:
            logger.error(f"Failed to load document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from an index. folder_path is ignored (use index)."""
        self._ensure_authenticated()
        docs = []

        try:
            # Scroll through all documents
            result = self._client.search(
                index=self.index,
                query={"match_all": {}},
                size=1000,
                scroll="2m",
            )

            scroll_id = result["_scroll_id"]
            hits = result["hits"]["hits"]

            while hits:
                for hit in hits:
                    content = hit["_source"].get(self.content_field, "")
                    docs.append(
                        LoadedDocument(
                            content=str(content),
                            source=self.source_name,
                            source_id=hit["_id"],
                            filename=f"{self.index}_{hit['_id']}",
                            metadata={"index": self.index, **hit["_source"]},
                        )
                    )

                result = self._client.scroll(scroll_id=scroll_id, scroll="2m")
                hits = result["hits"]["hits"]

            self._client.clear_scroll(scroll_id=scroll_id)
        except Exception as e:
            logger.error(f"Failed to load index {self.index}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search using Elasticsearch query."""
        self._ensure_authenticated()
        docs = []

        try:
            result = self._client.search(
                index=self.index,
                query={"match": {self.content_field: query}},
                size=max_results,
            )

            for hit in result["hits"]["hits"]:
                content = hit["_source"].get(self.content_field, "")
                docs.append(
                    LoadedDocument(
                        content=str(content),
                        source=self.source_name,
                        source_id=hit["_id"],
                        filename=f"{self.index}_{hit['_id']}",
                        metadata={
                            "index": self.index,
                            "score": hit["_score"],
                            **hit["_source"],
                        },
                    )
                )
        except Exception as e:
            logger.error(f"Elasticsearch search failed: {e}")

        return docs


# Redis
try:
    import redis

    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False


class RedisLoader(BaseLoader):
    """Load documents from Redis.

    Supports loading from Redis strings, hashes, and JSON (RedisJSON).

    Example:
        loader = RedisLoader(host="localhost", port=6379)
        docs = loader.load_folder("articles:*")  # key pattern
    """

    def __init__(
        self,
        host: str = "localhost",
        port: int = 6379,
        db: int = 0,
        password: Optional[str] = None,
        url: Optional[str] = None,
        content_field: str = "content",
    ):
        if not REDIS_AVAILABLE:
            raise ImportError("redis not installed. Run: pip install redis")

        self.host = host
        self.port = port
        self.db = db
        self.password = password or os.environ.get("REDIS_PASSWORD")
        self.url = url or os.environ.get("REDIS_URL")
        self.content_field = content_field
        self._client = None

    @property
    def source_name(self) -> str:
        return "redis"

    def authenticate(self) -> bool:
        """Connect to Redis."""
        try:
            if self.url:
                self._client = redis.from_url(self.url)
            else:
                self._client = redis.Redis(
                    host=self.host, port=self.port, db=self.db, password=self.password
                )
            # Test connection
            self._client.ping()
            logger.info("Redis connection successful")
            return True
        except Exception as e:
            logger.error(f"Redis connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._client and not self.authenticate():
            raise RuntimeError("Failed to connect to Redis")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single key."""
        self._ensure_authenticated()
        try:
            key_type = self._client.type(doc_id).decode()

            if key_type == "string":
                content = self._client.get(doc_id)
                content = content.decode() if content else ""
            elif key_type == "hash":
                data = self._client.hgetall(doc_id)
                data = {k.decode(): v.decode() for k, v in data.items()}
                content = data.get(self.content_field, json.dumps(data))
            elif key_type == "ReJSON-RL":  # RedisJSON
                content = self._client.execute_command("JSON.GET", doc_id)
                content = content if isinstance(content, str) else content.decode()
            else:
                logger.debug(f"Unsupported Redis type: {key_type}")
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=doc_id,
                metadata={"key": doc_id, "type": key_type},
            )
        except Exception as e:
            logger.error(f"Failed to load key {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all keys matching a pattern."""
        self._ensure_authenticated()
        docs = []
        pattern = folder_path if "*" in folder_path else f"{folder_path}*"

        try:
            for key in self._client.scan_iter(match=pattern):
                key_str = key.decode() if isinstance(key, bytes) else key
                doc = self.load_document(key_str)
                if doc:
                    docs.append(doc)
        except Exception as e:
            logger.error(f"Failed to scan pattern {pattern}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search using RediSearch (if available)."""
        self._ensure_authenticated()
        docs = []

        try:
            # Try RediSearch
            result = self._client.execute_command(
                "FT.SEARCH", "idx:docs", query, "LIMIT", 0, max_results
            )
            if isinstance(result, list) and len(result) > 1:
                for i in range(1, len(result), 2):
                    doc_id = (
                        result[i].decode()
                        if isinstance(result[i], bytes)
                        else result[i]
                    )
                    doc = self.load_document(doc_id)
                    if doc:
                        docs.append(doc)
        except Exception as e:
            logger.debug(f"RediSearch not available or failed: {e}")
            # Fallback to pattern matching
            docs = self.load_folder(f"*{query}*")[:max_results]

        return docs


# ============================================================================
# ENTERPRISE SYSTEM LOADERS
# ============================================================================


class SAPLoader(BaseLoader):
    """Load documents from SAP systems via RFC or OData.

    Connects to SAP ERP/S4HANA for document extraction.

    Example:
        loader = SAPLoader(
            host="sap.example.com",
            system_number="00",
            client="100",
            user="SAPUSER",
            password="secret"
        )
        docs = loader.load_folder("BKPF")  # Accounting documents
    """

    def __init__(
        self,
        host: str,
        system_number: str = "00",
        client: str = "100",
        user: Optional[str] = None,
        password: Optional[str] = None,
        odata_url: Optional[str] = None,
    ):
        self.host = host
        self.system_number = system_number
        self.client = client
        self.user = user or os.environ.get("SAP_USER")
        self.password = password or os.environ.get("SAP_PASSWORD")
        self.odata_url = odata_url
        self._connection = None
        self._session = None

    @property
    def source_name(self) -> str:
        return "sap"

    def authenticate(self) -> bool:
        """Connect to SAP system."""
        try:
            # Try pyrfc first (RFC connection)
            try:
                from pyrfc import Connection

                self._connection = Connection(
                    ashost=self.host,
                    sysnr=self.system_number,
                    client=self.client,
                    user=self.user,
                    passwd=self.password,
                )
                logger.info("SAP RFC connection successful")
                return True
            except ImportError:
                pass

            # Fallback to OData
            if self.odata_url:
                import requests

                self._session = requests.Session()
                self._session.auth = (self.user, self.password)
                # Test connection
                resp = self._session.get(f"{self.odata_url}/$metadata")
                resp.raise_for_status()
                logger.info("SAP OData connection successful")
                return True

            raise ImportError(
                "pyrfc not installed and no OData URL provided. "
                "Run: pip install pyrfc or provide odata_url"
            )
        except Exception as e:
            logger.error(f"SAP connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._connection and not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to SAP")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single SAP document."""
        self._ensure_authenticated()
        try:
            if self._connection:
                # RFC call to read document
                result = self._connection.call(
                    "RFC_READ_TABLE",
                    QUERY_TABLE=doc_id.split("/")[0],
                    DELIMITER="|",
                    OPTIONS=[{"TEXT": f"KEY = '{doc_id}'"}],
                )
                content = "\n".join([row["WA"] for row in result.get("DATA", [])])
            else:
                # OData request
                resp = self._session.get(f"{self.odata_url}/{doc_id}")
                resp.raise_for_status()
                content = json.dumps(resp.json(), indent=2)

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=doc_id,
                metadata={"system": self.host, "client": self.client},
            )
        except Exception as e:
            logger.error(f"Failed to load SAP document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load documents from a SAP table/entity set."""
        self._ensure_authenticated()
        docs = []
        table = folder_path

        try:
            if self._connection:
                result = self._connection.call(
                    "RFC_READ_TABLE",
                    QUERY_TABLE=table,
                    DELIMITER="|",
                    ROWCOUNT=1000,
                )
                for i, row in enumerate(result.get("DATA", [])):
                    docs.append(
                        LoadedDocument(
                            content=row["WA"],
                            source=self.source_name,
                            source_id=f"{table}/{i}",
                            filename=f"{table}_{i}",
                            metadata={"table": table},
                        )
                    )
            else:
                resp = self._session.get(f"{self.odata_url}/{table}")
                resp.raise_for_status()
                data = resp.json().get("d", {}).get("results", [])
                for i, item in enumerate(data):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(item, indent=2),
                            source=self.source_name,
                            source_id=f"{table}/{i}",
                            filename=f"{table}_{i}",
                            metadata={"entity_set": table, **item},
                        )
                    )
        except Exception as e:
            logger.error(f"Failed to load SAP table {table}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search SAP documents."""
        self._ensure_authenticated()
        if self._session and self.odata_url:
            try:
                resp = self._session.get(
                    f"{self.odata_url}/$search",
                    params={"$search": query, "$top": max_results},
                )
                resp.raise_for_status()
                docs = []
                for i, item in enumerate(resp.json().get("d", {}).get("results", [])):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(item, indent=2),
                            source=self.source_name,
                            source_id=f"search/{i}",
                            filename=f"search_{i}",
                            metadata=item,
                        )
                    )
                return docs
            except Exception as e:
                logger.error(f"SAP search failed: {e}")
        return []


class OracleLoader(BaseLoader):
    """Load documents from Oracle Database.

    Example:
        loader = OracleLoader(
            host="oracle.example.com",
            service_name="ORCL",
            user="system",
            password="secret"
        )
        docs = loader.load_folder("HR.EMPLOYEES")
    """

    def __init__(
        self,
        host: str = "localhost",
        port: int = 1521,
        service_name: Optional[str] = None,
        sid: Optional[str] = None,
        user: Optional[str] = None,
        password: Optional[str] = None,
        content_column: str = "content",
        id_column: str = "id",
    ):
        self.host = host
        self.port = port
        self.service_name = service_name
        self.sid = sid
        self.user = user or os.environ.get("ORACLE_USER")
        self.password = password or os.environ.get("ORACLE_PASSWORD")
        self.content_column = content_column
        self.id_column = id_column
        self._conn = None

    @property
    def source_name(self) -> str:
        return "oracle"

    def authenticate(self) -> bool:
        """Connect to Oracle Database."""
        try:
            import oracledb

            dsn = f"{self.host}:{self.port}"
            if self.service_name:
                dsn += f"/{self.service_name}"
            elif self.sid:
                dsn = oracledb.makedsn(self.host, self.port, sid=self.sid)

            self._conn = oracledb.connect(
                user=self.user, password=self.password, dsn=dsn
            )
            logger.info("Oracle connection successful")
            return True
        except ImportError:
            logger.error("oracledb not installed. Run: pip install oracledb")
            return False
        except Exception as e:
            logger.error(f"Oracle connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._conn and not self.authenticate():
            raise RuntimeError("Failed to connect to Oracle")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single row by ID."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            table = parts[0]
            row_id = parts[1] if len(parts) > 1 else doc_id

            cursor = self._conn.cursor()
            cursor.execute(
                f"SELECT {self.id_column}, {self.content_column} FROM {table} WHERE {self.id_column} = :id",
                {"id": row_id},
            )
            row = cursor.fetchone()
            cursor.close()

            if not row:
                return None

            return LoadedDocument(
                content=str(row[1]) if row[1] else "",
                source=self.source_name,
                source_id=doc_id,
                filename=f"{table}_{row_id}",
                metadata={"table": table, "row_id": row_id},
            )
        except Exception as e:
            logger.error(f"Failed to load row {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all rows from a table."""
        self._ensure_authenticated()
        docs = []
        table = folder_path

        try:
            cursor = self._conn.cursor()
            cursor.execute(
                f"SELECT {self.id_column}, {self.content_column} FROM {table} WHERE ROWNUM <= 10000"
            )
            rows = cursor.fetchall()
            cursor.close()

            for row in rows:
                row_id = str(row[0])
                content = str(row[1]) if row[1] else ""
                docs.append(
                    LoadedDocument(
                        content=content,
                        source=self.source_name,
                        source_id=f"{table}/{row_id}",
                        filename=f"{table}_{row_id}",
                        metadata={"table": table, "row_id": row_id},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load table {table}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search using Oracle Text (if available)."""
        logger.warning("Oracle search requires Oracle Text index on content column.")
        return []


class ServiceNowLoader(BaseLoader):
    """Load documents from ServiceNow.

    Connects to ServiceNow REST API for incident, knowledge, and other records.

    Example:
        loader = ServiceNowLoader(
            instance="mycompany",
            username="admin",
            password="secret"
        )
        docs = loader.load_folder("kb_knowledge")  # Knowledge articles
    """

    def __init__(
        self,
        instance: str,
        username: Optional[str] = None,
        password: Optional[str] = None,
        api_key: Optional[str] = None,
    ):
        self.instance = instance
        self.base_url = f"https://{instance}.service-now.com/api/now"
        self.username = username or os.environ.get("SERVICENOW_USERNAME")
        self.password = password or os.environ.get("SERVICENOW_PASSWORD")
        self.api_key = api_key or os.environ.get("SERVICENOW_API_KEY")
        self._session = None

    @property
    def source_name(self) -> str:
        return "servicenow"

    def authenticate(self) -> bool:
        """Connect to ServiceNow."""
        try:
            import requests

            self._session = requests.Session()
            if self.api_key:
                self._session.headers["Authorization"] = f"Bearer {self.api_key}"
            else:
                self._session.auth = (self.username, self.password)

            # Test connection
            resp = self._session.get(
                f"{self.base_url}/table/sys_user",
                params={"sysparm_limit": 1},
            )
            resp.raise_for_status()
            logger.info("ServiceNow connection successful")
            return True
        except Exception as e:
            logger.error(f"ServiceNow connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to ServiceNow")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single record by sys_id."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            table = parts[0]
            sys_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/table/{table}/{sys_id}")
            resp.raise_for_status()
            record = resp.json().get("result", {})

            # Build content from relevant fields
            content_fields = [
                "short_description",
                "description",
                "text",
                "content",
                "body",
            ]
            content = "\n".join(
                str(record.get(f, "")) for f in content_fields if record.get(f)
            )

            return LoadedDocument(
                content=content or json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{table}_{sys_id}",
                metadata={"table": table, "sys_id": sys_id, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a table."""
        self._ensure_authenticated()
        docs = []
        table = folder_path

        try:
            resp = self._session.get(
                f"{self.base_url}/table/{table}",
                params={"sysparm_limit": 1000},
            )
            resp.raise_for_status()

            for record in resp.json().get("result", []):
                sys_id = record.get("sys_id", "")
                content_fields = [
                    "short_description",
                    "description",
                    "text",
                    "content",
                ]
                content = "\n".join(
                    str(record.get(f, "")) for f in content_fields if record.get(f)
                )

                docs.append(
                    LoadedDocument(
                        content=content or json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{table}/{sys_id}",
                        filename=f"{table}_{sys_id}",
                        metadata={"table": table, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load table {table}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search ServiceNow using text search."""
        self._ensure_authenticated()
        docs = []

        try:
            # Search knowledge base
            resp = self._session.get(
                f"{self.base_url}/table/kb_knowledge",
                params={
                    "sysparm_query": f"textLIKE{query}",
                    "sysparm_limit": max_results,
                },
            )
            resp.raise_for_status()

            for record in resp.json().get("result", []):
                sys_id = record.get("sys_id", "")
                content = record.get("text", "") or record.get("short_description", "")

                docs.append(
                    LoadedDocument(
                        content=content,
                        source=self.source_name,
                        source_id=f"kb_knowledge/{sys_id}",
                        filename=f"kb_{sys_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"ServiceNow search failed: {e}")

        return docs


class WorkdayLoader(BaseLoader):
    """Load documents from Workday via REST API.

    Connects to Workday for HR data, reports, and documents.

    Example:
        loader = WorkdayLoader(
            tenant="mycompany",
            client_id="xxx",
            client_secret="yyy"
        )
        docs = loader.load_folder("workers")
    """

    def __init__(
        self,
        tenant: str,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        refresh_token: Optional[str] = None,
    ):
        self.tenant = tenant
        self.base_url = f"https://wd2-impl-services1.workday.com/ccx/api/v1/{tenant}"
        self.client_id = client_id or os.environ.get("WORKDAY_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("WORKDAY_CLIENT_SECRET")
        self.refresh_token = refresh_token or os.environ.get("WORKDAY_REFRESH_TOKEN")
        self._session = None
        self._access_token = None

    @property
    def source_name(self) -> str:
        return "workday"

    def authenticate(self) -> bool:
        """Connect to Workday."""
        try:
            import requests

            self._session = requests.Session()

            # Get access token
            token_url = (
                f"https://wd2-impl-services1.workday.com/ccx/oauth2/{self.tenant}/token"
            )
            resp = self._session.post(
                token_url,
                data={
                    "grant_type": "refresh_token",
                    "refresh_token": self.refresh_token,
                    "client_id": self.client_id,
                    "client_secret": self.client_secret,
                },
            )
            resp.raise_for_status()
            self._access_token = resp.json()["access_token"]
            self._session.headers["Authorization"] = f"Bearer {self._access_token}"

            logger.info("Workday connection successful")
            return True
        except Exception as e:
            logger.error(f"Workday connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Workday")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Workday record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            resource = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/{resource}/{record_id}")
            resp.raise_for_status()
            record = resp.json()

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{resource}_{record_id}",
                metadata={"resource": resource, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a Workday resource."""
        self._ensure_authenticated()
        docs = []
        resource = folder_path

        try:
            resp = self._session.get(
                f"{self.base_url}/{resource}", params={"limit": 100}
            )
            resp.raise_for_status()

            for record in resp.json().get("data", []):
                record_id = record.get("id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{resource}/{record_id}",
                        filename=f"{resource}_{record_id}",
                        metadata={"resource": resource, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load resource {resource}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Workday records."""
        logger.warning("Workday search requires specific report configuration.")
        return []


# ============================================================================
# CRM LOADERS
# ============================================================================


class ZohoLoader(BaseLoader):
    """Load documents from Zoho CRM.

    Example:
        loader = ZohoLoader(
            client_id="xxx",
            client_secret="yyy",
            refresh_token="zzz"
        )
        docs = loader.load_folder("Leads")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        refresh_token: Optional[str] = None,
        domain: str = "com",  # com, eu, in, com.cn, com.au
    ):
        self.client_id = client_id or os.environ.get("ZOHO_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("ZOHO_CLIENT_SECRET")
        self.refresh_token = refresh_token or os.environ.get("ZOHO_REFRESH_TOKEN")
        self.domain = domain
        self.base_url = f"https://www.zohoapis.{domain}/crm/v3"
        self._session = None
        self._access_token = None

    @property
    def source_name(self) -> str:
        return "zoho"

    def authenticate(self) -> bool:
        """Connect to Zoho CRM."""
        try:
            import requests

            self._session = requests.Session()

            # Get access token
            resp = self._session.post(
                f"https://accounts.zoho.{self.domain}/oauth/v2/token",
                data={
                    "grant_type": "refresh_token",
                    "refresh_token": self.refresh_token,
                    "client_id": self.client_id,
                    "client_secret": self.client_secret,
                },
            )
            resp.raise_for_status()
            self._access_token = resp.json()["access_token"]
            self._session.headers["Authorization"] = (
                f"Zoho-oauthtoken {self._access_token}"
            )

            logger.info("Zoho CRM connection successful")
            return True
        except Exception as e:
            logger.error(f"Zoho CRM connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Zoho CRM")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Zoho record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            module = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/{module}/{record_id}")
            resp.raise_for_status()
            record = resp.json().get("data", [{}])[0]

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{module}_{record_id}",
                metadata={"module": module, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a Zoho module."""
        self._ensure_authenticated()
        docs = []
        module = folder_path

        try:
            resp = self._session.get(
                f"{self.base_url}/{module}", params={"per_page": 200}
            )
            resp.raise_for_status()

            for record in resp.json().get("data", []):
                record_id = record.get("id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{module}/{record_id}",
                        filename=f"{module}_{record_id}",
                        metadata={"module": module, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load module {module}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Zoho CRM."""
        self._ensure_authenticated()
        docs = []

        try:
            resp = self._session.get(
                f"{self.base_url}/Leads/search",
                params={"word": query, "per_page": max_results},
            )
            resp.raise_for_status()

            for record in resp.json().get("data", []):
                record_id = record.get("id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"Leads/{record_id}",
                        filename=f"lead_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"Zoho search failed: {e}")

        return docs


class Dynamics365Loader(BaseLoader):
    """Load documents from Microsoft Dynamics 365.

    Example:
        loader = Dynamics365Loader(
            org_url="https://myorg.crm.dynamics.com",
            client_id="xxx",
            client_secret="yyy",
            tenant_id="zzz"
        )
        docs = loader.load_folder("accounts")
    """

    def __init__(
        self,
        org_url: str,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        tenant_id: Optional[str] = None,
    ):
        if not MSAL_AVAILABLE:
            raise ImportError("msal not installed. Run: pip install msal")

        self.org_url = org_url.rstrip("/")
        self.client_id = client_id or os.environ.get("DYNAMICS_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("DYNAMICS_CLIENT_SECRET")
        self.tenant_id = tenant_id or os.environ.get("DYNAMICS_TENANT_ID")
        self._session = None
        self._access_token = None

    @property
    def source_name(self) -> str:
        return "dynamics365"

    def authenticate(self) -> bool:
        """Connect to Dynamics 365."""
        try:
            import msal
            import requests

            app = msal.ConfidentialClientApplication(
                self.client_id,
                authority=f"https://login.microsoftonline.com/{self.tenant_id}",
                client_credential=self.client_secret,
            )

            result = app.acquire_token_for_client(scopes=[f"{self.org_url}/.default"])

            if "access_token" not in result:
                raise RuntimeError(
                    f"Failed to get token: {result.get('error_description')}"
                )

            self._access_token = result["access_token"]
            self._session = requests.Session()
            self._session.headers["Authorization"] = f"Bearer {self._access_token}"
            self._session.headers["OData-Version"] = "4.0"

            logger.info("Dynamics 365 connection successful")
            return True
        except Exception as e:
            logger.error(f"Dynamics 365 connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Dynamics 365")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Dynamics record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            entity = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(
                f"{self.org_url}/api/data/v9.2/{entity}({record_id})"
            )
            resp.raise_for_status()
            record = resp.json()

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{entity}_{record_id}",
                metadata={"entity": entity, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a Dynamics entity."""
        self._ensure_authenticated()
        docs = []
        entity = folder_path

        try:
            resp = self._session.get(
                f"{self.org_url}/api/data/v9.2/{entity}",
                params={"$top": 1000},
            )
            resp.raise_for_status()

            for record in resp.json().get("value", []):
                record_id = record.get(f"{entity.rstrip('s')}id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{entity}/{record_id}",
                        filename=f"{entity}_{record_id}",
                        metadata={"entity": entity, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load entity {entity}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Dynamics 365."""
        self._ensure_authenticated()
        docs = []

        try:
            resp = self._session.get(
                f"{self.org_url}/api/data/v9.2/accounts",
                params={
                    "$filter": f"contains(name,'{query}')",
                    "$top": max_results,
                },
            )
            resp.raise_for_status()

            for record in resp.json().get("value", []):
                record_id = record.get("accountid", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"accounts/{record_id}",
                        filename=f"account_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"Dynamics 365 search failed: {e}")

        return docs


# ============================================================================
# JIRA SERVICE DESK LOADER
# ============================================================================


class JiraServiceDeskLoader(BaseLoader):
    """Load documents from Jira Service Management (Service Desk).

    Extends JiraLoader with service desk specific features.

    Example:
        loader = JiraServiceDeskLoader(
            url="https://mycompany.atlassian.net",
            email="user@company.com",
            api_token="xxx"
        )
        docs = loader.load_folder("IT-SD")  # Service desk project key
    """

    def __init__(
        self,
        url: str,
        email: Optional[str] = None,
        api_token: Optional[str] = None,
    ):
        if not JIRA_AVAILABLE:
            raise ImportError("jira not installed. Run: pip install jira")

        self.url = url.rstrip("/")
        self.email = email or os.environ.get("JIRA_EMAIL")
        self.api_token = api_token or os.environ.get("JIRA_API_TOKEN")
        self._client = None

    @property
    def source_name(self) -> str:
        return "jira_service_desk"

    def authenticate(self) -> bool:
        """Connect to Jira Service Desk."""
        try:
            from jira import JIRA

            self._client = JIRA(
                server=self.url,
                basic_auth=(self.email, self.api_token),
            )
            # Test connection
            self._client.myself()
            logger.info("Jira Service Desk connection successful")
            return True
        except Exception as e:
            logger.error(f"Jira Service Desk connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._client and not self.authenticate():
            raise RuntimeError("Failed to connect to Jira Service Desk")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single request by key."""
        self._ensure_authenticated()
        try:
            issue = self._client.issue(doc_id)

            content = f"""# {issue.key}: {issue.fields.summary}

**Status:** {issue.fields.status.name}
**Reporter:** {issue.fields.reporter.displayName if issue.fields.reporter else 'Unknown'}
**Created:** {issue.fields.created}

## Description
{issue.fields.description or 'No description'}
"""

            # Add comments
            comments = self._client.comments(issue.key)
            if comments:
                content += "\n## Comments\n"
                for comment in comments:
                    content += f"\n**{comment.author.displayName}** ({comment.created}):\n{comment.body}\n"

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=f"{doc_id}.md",
                metadata={
                    "key": issue.key,
                    "summary": issue.fields.summary,
                    "status": issue.fields.status.name,
                    "priority": (
                        issue.fields.priority.name if issue.fields.priority else None
                    ),
                    "request_type": getattr(issue.fields, "customfield_10010", None),
                },
            )
        except Exception as e:
            logger.error(f"Failed to load request {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all requests from a service desk project."""
        self._ensure_authenticated()
        docs = []
        project_key = folder_path

        try:
            jql = f'project = "{project_key}" ORDER BY created DESC'
            issues = self._client.search_issues(jql, maxResults=500)

            for issue in issues:
                doc = self.load_document(issue.key)
                if doc:
                    docs.append(doc)
        except Exception as e:
            logger.error(f"Failed to load project {project_key}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search service desk requests."""
        self._ensure_authenticated()
        docs = []

        try:
            jql = f'text ~ "{query}" ORDER BY created DESC'
            issues = self._client.search_issues(jql, maxResults=max_results)

            for issue in issues:
                doc = self.load_document(issue.key)
                if doc:
                    docs.append(doc)
        except Exception as e:
            logger.error(f"Jira Service Desk search failed: {e}")

        return docs


# ============================================================================
# AUSTRALIAN BUSINESS LOADERS
# ============================================================================


class MYOBLoader(BaseLoader):
    """Load documents from MYOB AccountRight/Essentials.

    Australian accounting software. Connects via MYOB API.

    Example:
        loader = MYOBLoader(
            client_id="xxx",
            client_secret="yyy",
            company_file_id="zzz"
        )
        docs = loader.load_folder("invoices")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        access_token: Optional[str] = None,
        refresh_token: Optional[str] = None,
        company_file_id: Optional[str] = None,
        company_file_username: str = "Administrator",
        company_file_password: str = "",
    ):
        self.client_id = client_id or os.environ.get("MYOB_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("MYOB_CLIENT_SECRET")
        self.access_token = access_token or os.environ.get("MYOB_ACCESS_TOKEN")
        self.refresh_token = refresh_token or os.environ.get("MYOB_REFRESH_TOKEN")
        self.company_file_id = company_file_id or os.environ.get("MYOB_COMPANY_FILE_ID")
        self.company_file_username = company_file_username
        self.company_file_password = company_file_password
        self.base_url = "https://api.myob.com/accountright"
        self._session = None

    @property
    def source_name(self) -> str:
        return "myob"

    def authenticate(self) -> bool:
        """Connect to MYOB API."""
        try:
            import requests

            self._session = requests.Session()

            # Refresh token if needed
            if self.refresh_token and not self.access_token:
                resp = self._session.post(
                    "https://secure.myob.com/oauth2/v1/authorize",
                    data={
                        "grant_type": "refresh_token",
                        "refresh_token": self.refresh_token,
                        "client_id": self.client_id,
                        "client_secret": self.client_secret,
                    },
                )
                resp.raise_for_status()
                self.access_token = resp.json()["access_token"]

            self._session.headers["Authorization"] = f"Bearer {self.access_token}"
            self._session.headers["x-myobapi-key"] = self.client_id
            self._session.headers["x-myobapi-version"] = "v2"

            # Company file auth
            if self.company_file_username:
                import base64

                auth = base64.b64encode(
                    f"{self.company_file_username}:{self.company_file_password}".encode()
                ).decode()
                self._session.headers["x-myobapi-cftoken"] = auth

            # Test connection
            resp = self._session.get(f"{self.base_url}/{self.company_file_id}/Info")
            resp.raise_for_status()

            logger.info("MYOB connection successful")
            return True
        except Exception as e:
            logger.error(f"MYOB connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to MYOB")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single MYOB record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            endpoint = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(
                f"{self.base_url}/{self.company_file_id}/{endpoint}/{record_id}"
            )
            resp.raise_for_status()
            record = resp.json()

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{endpoint}_{record_id}",
                metadata={"endpoint": endpoint, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load MYOB record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a MYOB endpoint."""
        self._ensure_authenticated()
        docs = []
        endpoint = folder_path

        # Map common names to MYOB endpoints
        endpoint_map = {
            "invoices": "Sale/Invoice",
            "customers": "Contact/Customer",
            "suppliers": "Contact/Supplier",
            "accounts": "GeneralLedger/Account",
            "employees": "Contact/Employee",
            "items": "Inventory/Item",
        }
        myob_endpoint = endpoint_map.get(endpoint.lower(), endpoint)

        try:
            resp = self._session.get(
                f"{self.base_url}/{self.company_file_id}/{myob_endpoint}"
            )
            resp.raise_for_status()

            for record in resp.json().get("Items", []):
                record_id = record.get("UID", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{myob_endpoint}/{record_id}",
                        filename=f"{endpoint}_{record_id}",
                        metadata={"endpoint": myob_endpoint, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load MYOB endpoint {endpoint}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search MYOB records."""
        self._ensure_authenticated()
        docs = []

        try:
            # Search customers
            resp = self._session.get(
                f"{self.base_url}/{self.company_file_id}/Contact/Customer",
                params={"$filter": f"contains(Name,'{query}')"},
            )
            resp.raise_for_status()

            for record in resp.json().get("Items", [])[:max_results]:
                record_id = record.get("UID", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"Contact/Customer/{record_id}",
                        filename=f"customer_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"MYOB search failed: {e}")

        return docs


class XeroLoader(BaseLoader):
    """Load documents from Xero accounting software.

    Popular Australian/NZ accounting software.

    Example:
        loader = XeroLoader(
            client_id="xxx",
            client_secret="yyy",
            tenant_id="zzz"
        )
        docs = loader.load_folder("invoices")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        access_token: Optional[str] = None,
        refresh_token: Optional[str] = None,
        tenant_id: Optional[str] = None,
    ):
        self.client_id = client_id or os.environ.get("XERO_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("XERO_CLIENT_SECRET")
        self.access_token = access_token or os.environ.get("XERO_ACCESS_TOKEN")
        self.refresh_token = refresh_token or os.environ.get("XERO_REFRESH_TOKEN")
        self.tenant_id = tenant_id or os.environ.get("XERO_TENANT_ID")
        self.base_url = "https://api.xero.com/api.xro/2.0"
        self._session = None

    @property
    def source_name(self) -> str:
        return "xero"

    def authenticate(self) -> bool:
        """Connect to Xero API."""
        try:
            import requests

            self._session = requests.Session()

            # Refresh token if needed
            if self.refresh_token and not self.access_token:
                import base64

                auth = base64.b64encode(
                    f"{self.client_id}:{self.client_secret}".encode()
                ).decode()
                resp = self._session.post(
                    "https://identity.xero.com/connect/token",
                    headers={"Authorization": f"Basic {auth}"},
                    data={
                        "grant_type": "refresh_token",
                        "refresh_token": self.refresh_token,
                    },
                )
                resp.raise_for_status()
                tokens = resp.json()
                self.access_token = tokens["access_token"]

            self._session.headers["Authorization"] = f"Bearer {self.access_token}"
            self._session.headers["Xero-Tenant-Id"] = self.tenant_id
            self._session.headers["Accept"] = "application/json"

            # Test connection
            resp = self._session.get(f"{self.base_url}/Organisation")
            resp.raise_for_status()

            logger.info("Xero connection successful")
            return True
        except Exception as e:
            logger.error(f"Xero connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Xero")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Xero record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            endpoint = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/{endpoint}/{record_id}")
            resp.raise_for_status()
            data = resp.json()
            record = data.get(endpoint, [{}])[0] if endpoint in data else data

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{endpoint}_{record_id}",
                metadata={"endpoint": endpoint, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load Xero record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from a Xero endpoint."""
        self._ensure_authenticated()
        docs = []
        endpoint = folder_path.capitalize()

        try:
            resp = self._session.get(f"{self.base_url}/{endpoint}")
            resp.raise_for_status()
            data = resp.json()

            records = data.get(endpoint, [])
            for record in records:
                record_id = (
                    record.get("InvoiceID")
                    or record.get("ContactID")
                    or record.get("AccountID", "")
                )
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{endpoint}/{record_id}",
                        filename=f"{endpoint}_{record_id}",
                        metadata={"endpoint": endpoint, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load Xero endpoint {endpoint}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Xero records."""
        self._ensure_authenticated()
        docs = []

        try:
            # Search contacts
            resp = self._session.get(
                f"{self.base_url}/Contacts",
                params={"where": f'Name.Contains("{query}")'},
            )
            resp.raise_for_status()

            for record in resp.json().get("Contacts", [])[:max_results]:
                record_id = record.get("ContactID", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"Contacts/{record_id}",
                        filename=f"contact_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"Xero search failed: {e}")

        return docs


class EmploymentHeroLoader(BaseLoader):
    """Load documents from Employment Hero (Australian HR/Payroll).

    Example:
        loader = EmploymentHeroLoader(
            api_key="xxx",
            organization_id="yyy"
        )
        docs = loader.load_folder("employees")
    """

    def __init__(
        self,
        api_key: Optional[str] = None,
        organization_id: Optional[str] = None,
    ):
        self.api_key = api_key or os.environ.get("EMPLOYMENT_HERO_API_KEY")
        self.organization_id = organization_id or os.environ.get(
            "EMPLOYMENT_HERO_ORG_ID"
        )
        self.base_url = f"https://api.employmenthero.com/api/v1/organisations/{self.organization_id}"
        self._session = None

    @property
    def source_name(self) -> str:
        return "employment_hero"

    def authenticate(self) -> bool:
        """Connect to Employment Hero API."""
        try:
            import requests

            self._session = requests.Session()
            self._session.headers["Authorization"] = f"Bearer {self.api_key}"
            self._session.headers["Accept"] = "application/json"

            # Test connection
            resp = self._session.get(f"{self.base_url}")
            resp.raise_for_status()

            logger.info("Employment Hero connection successful")
            return True
        except Exception as e:
            logger.error(f"Employment Hero connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Employment Hero")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single employee record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            endpoint = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/{endpoint}/{record_id}")
            resp.raise_for_status()
            record = resp.json().get("data", resp.json())

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{endpoint}_{record_id}",
                metadata={"endpoint": endpoint, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from Employment Hero."""
        self._ensure_authenticated()
        docs = []
        endpoint = folder_path

        try:
            resp = self._session.get(f"{self.base_url}/{endpoint}")
            resp.raise_for_status()

            records = resp.json().get("data", [])
            for record in records:
                record_id = record.get("id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{endpoint}/{record_id}",
                        filename=f"{endpoint}_{record_id}",
                        metadata={"endpoint": endpoint, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load {endpoint}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Employment Hero records."""
        self._ensure_authenticated()
        docs = []

        try:
            resp = self._session.get(
                f"{self.base_url}/employees",
                params={"search": query, "per_page": max_results},
            )
            resp.raise_for_status()

            for record in resp.json().get("data", []):
                record_id = record.get("id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"employees/{record_id}",
                        filename=f"employee_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"Employment Hero search failed: {e}")

        return docs


class DeputyLoader(BaseLoader):
    """Load documents from Deputy (Australian workforce management).

    Example:
        loader = DeputyLoader(
            domain="mycompany",
            access_token="xxx"
        )
        docs = loader.load_folder("employees")
    """

    def __init__(
        self,
        domain: str,
        access_token: Optional[str] = None,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        refresh_token: Optional[str] = None,
    ):
        self.domain = domain
        self.access_token = access_token or os.environ.get("DEPUTY_ACCESS_TOKEN")
        self.client_id = client_id or os.environ.get("DEPUTY_CLIENT_ID")
        self.client_secret = client_secret or os.environ.get("DEPUTY_CLIENT_SECRET")
        self.refresh_token = refresh_token or os.environ.get("DEPUTY_REFRESH_TOKEN")
        self.base_url = f"https://{domain}.na.deputy.com/api/v1"
        self._session = None

    @property
    def source_name(self) -> str:
        return "deputy"

    def authenticate(self) -> bool:
        """Connect to Deputy API."""
        try:
            import requests

            self._session = requests.Session()

            # Refresh token if needed
            if self.refresh_token and not self.access_token:
                resp = self._session.post(
                    f"https://{self.domain}.na.deputy.com/oauth/access_token",
                    data={
                        "grant_type": "refresh_token",
                        "refresh_token": self.refresh_token,
                        "client_id": self.client_id,
                        "client_secret": self.client_secret,
                    },
                )
                resp.raise_for_status()
                self.access_token = resp.json()["access_token"]

            self._session.headers["Authorization"] = f"Bearer {self.access_token}"
            self._session.headers["dp-meta-option"] = "none"

            # Test connection
            resp = self._session.get(f"{self.base_url}/me")
            resp.raise_for_status()

            logger.info("Deputy connection successful")
            return True
        except Exception as e:
            logger.error(f"Deputy connection failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._session and not self.authenticate():
            raise RuntimeError("Failed to connect to Deputy")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Deputy record."""
        self._ensure_authenticated()
        try:
            parts = doc_id.split("/", 1)
            resource = parts[0]
            record_id = parts[1] if len(parts) > 1 else doc_id

            resp = self._session.get(f"{self.base_url}/resource/{resource}/{record_id}")
            resp.raise_for_status()
            record = resp.json()

            return LoadedDocument(
                content=json.dumps(record, indent=2),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{resource}_{record_id}",
                metadata={"resource": resource, **record},
            )
        except Exception as e:
            logger.error(f"Failed to load record {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load records from Deputy resource."""
        self._ensure_authenticated()
        docs = []

        # Map common names to Deputy resources
        resource_map = {
            "employees": "Employee",
            "timesheets": "Timesheet",
            "rosters": "Roster",
            "locations": "OperationalUnit",
            "leave": "Leave",
        }
        resource = resource_map.get(folder_path.lower(), folder_path)

        try:
            resp = self._session.get(f"{self.base_url}/resource/{resource}/QUERY")
            resp.raise_for_status()

            for record in resp.json():
                record_id = record.get("Id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"{resource}/{record_id}",
                        filename=f"{resource}_{record_id}",
                        metadata={"resource": resource, **record},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load {resource}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Deputy records."""
        self._ensure_authenticated()
        docs = []

        try:
            resp = self._session.post(
                f"{self.base_url}/resource/Employee/QUERY",
                json={
                    "search": {
                        "s1": {"field": "DisplayName", "data": query, "type": "lk"}
                    }
                },
            )
            resp.raise_for_status()

            for record in resp.json()[:max_results]:
                record_id = record.get("Id", "")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(record, indent=2),
                        source=self.source_name,
                        source_id=f"Employee/{record_id}",
                        filename=f"employee_{record_id}",
                        metadata=record,
                    )
                )
        except Exception as e:
            logger.error(f"Deputy search failed: {e}")

        return docs


# ============================================================================
# PAYMENT GATEWAY LOADERS
# ============================================================================


class StripeLoader(BaseLoader):
    """Document loader for Stripe payment platform.

    Load transactions, customers, invoices, and subscriptions from Stripe.

    Requirements:
        pip install stripe

    Environment variables:
        STRIPE_API_KEY: Stripe secret API key

    Example:
        loader = StripeLoader(api_key="sk_live_xxx")
        loader.authenticate()
        docs = loader.load_folder("transactions")
        docs = loader.load_folder("customers")
        docs = loader.load_folder("invoices")
        docs = loader.load_folder("subscriptions")
    """

    def __init__(self, api_key: Optional[str] = None, **kwargs):
        """Initialize Stripe loader.

        Args:
            api_key: Stripe secret API key
        """
        self._api_key = api_key or os.environ.get("STRIPE_API_KEY")
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "stripe"

    def authenticate(self) -> bool:
        """Authenticate with Stripe API."""
        try:
            import stripe

            stripe.api_key = self._api_key
            # Test with a simple API call
            stripe.Account.retrieve()
            self._client = stripe
            self._authenticated = True
            logger.info("Stripe authentication successful")
            return True
        except ImportError:
            logger.error("stripe library not installed. Run: pip install stripe")
            return False
        except Exception as e:
            logger.error(f"Stripe authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Stripe object.

        Args:
            doc_id: Object ID in format "type:id" (e.g., "customer:cus_xxx")
        """
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)
            obj_map = {
                "transaction": self._client.PaymentIntent,
                "payment_intent": self._client.PaymentIntent,
                "charge": self._client.Charge,
                "customer": self._client.Customer,
                "invoice": self._client.Invoice,
                "subscription": self._client.Subscription,
            }

            if obj_type not in obj_map:
                logger.error(f"Unknown Stripe object type: {obj_type}")
                return None

            obj = obj_map[obj_type].retrieve(obj_id)
            data = dict(obj)

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
                created_at=(
                    datetime.fromtimestamp(data.get("created", 0), tz=timezone.utc)
                    if data.get("created")
                    else None
                ),
            )
        except Exception as e:
            logger.error(f"Failed to load Stripe object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Stripe objects by type.

        Args:
            folder_path: One of 'transactions', 'customers', 'invoices', 'subscriptions'
        """
        self._ensure_authenticated()
        docs = []

        type_map = {
            "transactions": ("PaymentIntent", "payment_intent"),
            "charges": ("Charge", "charge"),
            "customers": ("Customer", "customer"),
            "invoices": ("Invoice", "invoice"),
            "subscriptions": ("Subscription", "subscription"),
        }

        folder_lower = folder_path.lower()
        if folder_lower not in type_map:
            logger.warning(
                f"Unknown Stripe folder: {folder_path}. Use: {list(type_map.keys())}"
            )
            return docs

        api_obj_name, doc_type = type_map[folder_lower]

        try:
            api_obj = getattr(self._client, api_obj_name)
            items = api_obj.list(limit=100)

            for item in items.auto_paging_iter():
                data = dict(item)
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{doc_type}:{item.id}",
                        filename=f"{doc_type}_{item.id}.json",
                        metadata={"type": doc_type, "id": item.id},
                        created_at=(
                            datetime.fromtimestamp(
                                data.get("created", 0), tz=timezone.utc
                            )
                            if data.get("created")
                            else None
                        ),
                    )
                )
                if len(docs) >= 100:
                    break
        except Exception as e:
            logger.error(f"Failed to load Stripe {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Stripe customers by email or name."""
        self._ensure_authenticated()
        docs = []

        try:
            # Search customers by email
            customers = self._client.Customer.search(
                query=f"email~'{query}' OR name~'{query}'", limit=max_results
            )
            for cust in customers.data:
                data = dict(cust)
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"customer:{cust.id}",
                        filename=f"customer_{cust.id}.json",
                        metadata={"type": "customer", "id": cust.id},
                        created_at=(
                            datetime.fromtimestamp(
                                data.get("created", 0), tz=timezone.utc
                            )
                            if data.get("created")
                            else None
                        ),
                    )
                )
        except Exception as e:
            logger.error(f"Stripe search failed: {e}")

        return docs


class PayPalLoader(BaseLoader):
    """Document loader for PayPal payment platform.

    Load transactions, customers, and invoices from PayPal.

    Requirements:
        pip install paypalrestsdk

    Environment variables:
        PAYPAL_CLIENT_ID: PayPal client ID
        PAYPAL_CLIENT_SECRET: PayPal client secret
        PAYPAL_MODE: 'sandbox' or 'live'

    Example:
        loader = PayPalLoader(client_id="xxx", client_secret="yyy")
        loader.authenticate()
        docs = loader.load_folder("transactions")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        mode: Optional[str] = None,
        **kwargs,
    ):
        """Initialize PayPal loader."""
        self._client_id = client_id or os.environ.get("PAYPAL_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get("PAYPAL_CLIENT_SECRET")
        self._mode = mode or os.environ.get("PAYPAL_MODE", "sandbox")
        self._api: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "paypal"

    def authenticate(self) -> bool:
        """Authenticate with PayPal API."""
        try:
            import paypalrestsdk

            paypalrestsdk.configure(
                {
                    "mode": self._mode,
                    "client_id": self._client_id,
                    "client_secret": self._client_secret,
                }
            )
            self._api = paypalrestsdk
            self._authenticated = True
            logger.info("PayPal authentication successful")
            return True
        except ImportError:
            logger.error("paypalrestsdk not installed. Run: pip install paypalrestsdk")
            return False
        except Exception as e:
            logger.error(f"PayPal authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single PayPal object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "payment":
                obj = self._api.Payment.find(obj_id)
            elif obj_type == "invoice":
                obj = self._api.Invoice.find(obj_id)
            else:
                logger.error(f"Unknown PayPal object type: {obj_type}")
                return None

            data = obj.to_dict()
            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load PayPal object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load PayPal objects by type."""
        self._ensure_authenticated()
        docs = []

        folder_lower = folder_path.lower()

        try:
            if folder_lower in ("transactions", "payments"):
                history = self._api.Payment.all({"count": 100})
                for payment in history.payments:
                    data = payment.to_dict()
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(data, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"payment:{payment.id}",
                            filename=f"payment_{payment.id}.json",
                            metadata={"type": "payment", "id": payment.id},
                        )
                    )
            elif folder_lower == "invoices":
                invoices = self._api.Invoice.all({"page_size": 100})
                for inv in invoices.invoices:
                    data = inv.to_dict()
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(data, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"invoice:{inv.id}",
                            filename=f"invoice_{inv.id}.json",
                            metadata={"type": "invoice", "id": inv.id},
                        )
                    )
            else:
                logger.warning(f"Unknown PayPal folder: {folder_path}")
        except Exception as e:
            logger.error(f"Failed to load PayPal {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search PayPal transactions."""
        self._ensure_authenticated()
        # PayPal API doesn't have great search - load recent and filter
        docs = self.load_folder("transactions")
        return [d for d in docs if query.lower() in d.content.lower()][:max_results]


class SquareLoader(BaseLoader):
    """Document loader for Square payment platform.

    Load transactions, customers, items, and orders from Square.

    Requirements:
        pip install squareup

    Environment variables:
        SQUARE_ACCESS_TOKEN: Square access token
        SQUARE_ENVIRONMENT: 'sandbox' or 'production'

    Example:
        loader = SquareLoader(access_token="xxx")
        loader.authenticate()
        docs = loader.load_folder("transactions")
    """

    def __init__(
        self,
        access_token: Optional[str] = None,
        environment: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Square loader."""
        self._access_token = access_token or os.environ.get("SQUARE_ACCESS_TOKEN")
        self._environment = environment or os.environ.get(
            "SQUARE_ENVIRONMENT", "production"
        )
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "square"

    def authenticate(self) -> bool:
        """Authenticate with Square API."""
        try:
            from square.client import Client

            self._client = Client(
                access_token=self._access_token,
                environment=self._environment,
            )
            # Test connection
            result = self._client.locations.list_locations()
            if result.is_error():
                raise Exception(result.errors)
            self._authenticated = True
            logger.info("Square authentication successful")
            return True
        except ImportError:
            logger.error("squareup not installed. Run: pip install squareup")
            return False
        except Exception as e:
            logger.error(f"Square authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Square object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "customer":
                result = self._client.customers.retrieve_customer(customer_id=obj_id)
                data = result.body.get("customer", {})
            elif obj_type == "order":
                result = self._client.orders.retrieve_order(order_id=obj_id)
                data = result.body.get("order", {})
            elif obj_type == "item":
                result = self._client.catalog.retrieve_catalog_object(object_id=obj_id)
                data = result.body.get("object", {})
            else:
                logger.error(f"Unknown Square object type: {obj_type}")
                return None

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load Square object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Square objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            if folder_lower == "customers":
                result = self._client.customers.list_customers()
                for cust in result.body.get("customers", []):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(cust, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"customer:{cust['id']}",
                            filename=f"customer_{cust['id']}.json",
                            metadata={"type": "customer", "id": cust["id"]},
                        )
                    )
            elif folder_lower in ("transactions", "orders"):
                # Get all locations first
                locations = self._client.locations.list_locations()
                location_ids = [
                    loc["id"] for loc in locations.body.get("locations", [])
                ]

                for loc_id in location_ids:
                    result = self._client.orders.search_orders(
                        body={"location_ids": [loc_id], "limit": 100}
                    )
                    for order in result.body.get("orders", []):
                        docs.append(
                            LoadedDocument(
                                content=json.dumps(order, indent=2, default=str),
                                source=self.source_name,
                                source_id=f"order:{order['id']}",
                                filename=f"order_{order['id']}.json",
                                metadata={"type": "order", "id": order["id"]},
                            )
                        )
            elif folder_lower == "items":
                result = self._client.catalog.list_catalog(types="ITEM")
                for obj in result.body.get("objects", []):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(obj, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"item:{obj['id']}",
                            filename=f"item_{obj['id']}.json",
                            metadata={"type": "item", "id": obj["id"]},
                        )
                    )
            else:
                logger.warning(f"Unknown Square folder: {folder_path}")
        except Exception as e:
            logger.error(f"Failed to load Square {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Square customers."""
        self._ensure_authenticated()
        docs = []

        try:
            result = self._client.customers.search_customers(
                body={
                    "query": {"filter": {"reference_id": {"fuzzy": query}}},
                    "limit": max_results,
                }
            )
            for cust in result.body.get("customers", []):
                docs.append(
                    LoadedDocument(
                        content=json.dumps(cust, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"customer:{cust['id']}",
                        filename=f"customer_{cust['id']}.json",
                        metadata={"type": "customer", "id": cust["id"]},
                    )
                )
        except Exception as e:
            logger.error(f"Square search failed: {e}")

        return docs


class AfterpayLoader(BaseLoader):
    """Document loader for Afterpay (Australian buy-now-pay-later).

    Load orders and payments from Afterpay.

    Requirements:
        pip install requests

    Environment variables:
        AFTERPAY_MERCHANT_ID: Merchant ID
        AFTERPAY_SECRET_KEY: Secret key
        AFTERPAY_ENVIRONMENT: 'sandbox' or 'production'

    Example:
        loader = AfterpayLoader(merchant_id="xxx", secret_key="yyy")
        loader.authenticate()
        docs = loader.load_folder("orders")
    """

    def __init__(
        self,
        merchant_id: Optional[str] = None,
        secret_key: Optional[str] = None,
        environment: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Afterpay loader."""
        self._merchant_id = merchant_id or os.environ.get("AFTERPAY_MERCHANT_ID")
        self._secret_key = secret_key or os.environ.get("AFTERPAY_SECRET_KEY")
        self._environment = environment or os.environ.get(
            "AFTERPAY_ENVIRONMENT", "production"
        )
        self._session: Optional[Any] = None
        self._authenticated = False

        # Set base URL based on environment
        if self._environment == "sandbox":
            self._base_url = "https://api-sandbox.afterpay.com/v2"
        else:
            self._base_url = "https://api.afterpay.com/v2"

    @property
    def source_name(self) -> str:
        return "afterpay"

    def authenticate(self) -> bool:
        """Authenticate with Afterpay API."""
        try:
            import requests

            self._session = requests.Session()
            self._session.auth = (self._merchant_id, self._secret_key)
            self._session.headers.update({"Accept": "application/json"})

            # Test connection
            resp = self._session.get(f"{self._base_url}/configuration")
            resp.raise_for_status()
            self._authenticated = True
            logger.info("Afterpay authentication successful")
            return True
        except Exception as e:
            logger.error(f"Afterpay authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Afterpay object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "order":
                resp = self._session.get(f"{self._base_url}/orders/{obj_id}")
            elif obj_type == "payment":
                resp = self._session.get(f"{self._base_url}/payments/{obj_id}")
            else:
                logger.error(f"Unknown Afterpay object type: {obj_type}")
                return None

            resp.raise_for_status()
            data = resp.json()

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load Afterpay object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Afterpay objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            if folder_lower == "orders":
                resp = self._session.get(
                    f"{self._base_url}/orders", params={"limit": 100}
                )
                resp.raise_for_status()
                for order in resp.json().get("results", []):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(order, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"order:{order['id']}",
                            filename=f"order_{order['id']}.json",
                            metadata={"type": "order", "id": order["id"]},
                        )
                    )
            elif folder_lower == "payments":
                resp = self._session.get(
                    f"{self._base_url}/payments", params={"limit": 100}
                )
                resp.raise_for_status()
                for payment in resp.json().get("results", []):
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(payment, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"payment:{payment['id']}",
                            filename=f"payment_{payment['id']}.json",
                            metadata={"type": "payment", "id": payment["id"]},
                        )
                    )
            else:
                logger.warning(f"Unknown Afterpay folder: {folder_path}")
        except Exception as e:
            logger.error(f"Failed to load Afterpay {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Afterpay orders by reference."""
        self._ensure_authenticated()
        docs = []

        try:
            resp = self._session.get(
                f"{self._base_url}/orders",
                params={"merchantReference": query, "limit": max_results},
            )
            resp.raise_for_status()
            for order in resp.json().get("results", []):
                docs.append(
                    LoadedDocument(
                        content=json.dumps(order, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"order:{order['id']}",
                        filename=f"order_{order['id']}.json",
                        metadata={"type": "order", "id": order["id"]},
                    )
                )
        except Exception as e:
            logger.error(f"Afterpay search failed: {e}")

        return docs


class BraintreeLoader(BaseLoader):
    """Document loader for Braintree (PayPal) payment platform.

    Load transactions and customers from Braintree.

    Requirements:
        pip install braintree

    Environment variables:
        BRAINTREE_MERCHANT_ID: Merchant ID
        BRAINTREE_PUBLIC_KEY: Public key
        BRAINTREE_PRIVATE_KEY: Private key
        BRAINTREE_ENVIRONMENT: 'sandbox' or 'production'

    Example:
        loader = BraintreeLoader(merchant_id="xxx", public_key="yyy", private_key="zzz")
        loader.authenticate()
        docs = loader.load_folder("transactions")
    """

    def __init__(
        self,
        merchant_id: Optional[str] = None,
        public_key: Optional[str] = None,
        private_key: Optional[str] = None,
        environment: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Braintree loader."""
        self._merchant_id = merchant_id or os.environ.get("BRAINTREE_MERCHANT_ID")
        self._public_key = public_key or os.environ.get("BRAINTREE_PUBLIC_KEY")
        self._private_key = private_key or os.environ.get("BRAINTREE_PRIVATE_KEY")
        self._environment = environment or os.environ.get(
            "BRAINTREE_ENVIRONMENT", "production"
        )
        self._gateway: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "braintree"

    def authenticate(self) -> bool:
        """Authenticate with Braintree."""
        try:
            import braintree

            env = (
                braintree.Environment.Sandbox
                if self._environment == "sandbox"
                else braintree.Environment.Production
            )
            self._gateway = braintree.BraintreeGateway(
                braintree.Configuration(
                    environment=env,
                    merchant_id=self._merchant_id,
                    public_key=self._public_key,
                    private_key=self._private_key,
                )
            )
            # Test connection
            self._gateway.client_token.generate()
            self._authenticated = True
            logger.info("Braintree authentication successful")
            return True
        except ImportError:
            logger.error("braintree not installed. Run: pip install braintree")
            return False
        except Exception as e:
            logger.error(f"Braintree authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Braintree object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "transaction":
                obj = self._gateway.transaction.find(obj_id)
            elif obj_type == "customer":
                obj = self._gateway.customer.find(obj_id)
            else:
                logger.error(f"Unknown Braintree object type: {obj_type}")
                return None

            # Convert to dict
            data = {
                attr: getattr(obj, attr, None)
                for attr in dir(obj)
                if not attr.startswith("_")
            }

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load Braintree object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Braintree objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            if folder_lower == "transactions":
                import braintree

                collection = self._gateway.transaction.search(
                    braintree.TransactionSearch.created_at.between(
                        datetime.now(tz=timezone.utc) - timedelta(days=30),
                        datetime.now(tz=timezone.utc),
                    )
                )
                for txn in collection.items:
                    data = {
                        attr: getattr(txn, attr, None)
                        for attr in dir(txn)
                        if not attr.startswith("_")
                    }
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(data, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"transaction:{txn.id}",
                            filename=f"transaction_{txn.id}.json",
                            metadata={"type": "transaction", "id": txn.id},
                        )
                    )
                    if len(docs) >= 100:
                        break
            elif folder_lower == "customers":
                import braintree

                collection = self._gateway.customer.search(
                    braintree.CustomerSearch.created_at.between(
                        datetime.now(tz=timezone.utc) - timedelta(days=365),
                        datetime.now(tz=timezone.utc),
                    )
                )
                for cust in collection.items:
                    data = {
                        attr: getattr(cust, attr, None)
                        for attr in dir(cust)
                        if not attr.startswith("_")
                    }
                    docs.append(
                        LoadedDocument(
                            content=json.dumps(data, indent=2, default=str),
                            source=self.source_name,
                            source_id=f"customer:{cust.id}",
                            filename=f"customer_{cust.id}.json",
                            metadata={"type": "customer", "id": cust.id},
                        )
                    )
                    if len(docs) >= 100:
                        break
            else:
                logger.warning(f"Unknown Braintree folder: {folder_path}")
        except Exception as e:
            logger.error(f"Failed to load Braintree {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Braintree customers by email."""
        self._ensure_authenticated()
        docs = []

        try:
            import braintree

            collection = self._gateway.customer.search(
                braintree.CustomerSearch.email.contains(query)
            )
            for cust in collection.items:
                data = {
                    attr: getattr(cust, attr, None)
                    for attr in dir(cust)
                    if not attr.startswith("_")
                }
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"customer:{cust.id}",
                        filename=f"customer_{cust.id}.json",
                        metadata={"type": "customer", "id": cust.id},
                    )
                )
                if len(docs) >= max_results:
                    break
        except Exception as e:
            logger.error(f"Braintree search failed: {e}")

        return docs


# ============================================================================
# E-COMMERCE PLATFORM LOADERS
# ============================================================================


class ShopifyLoader(BaseLoader):
    """Document loader for Shopify e-commerce platform.

    Load products, orders, customers, and collections from Shopify.

    Requirements:
        pip install ShopifyAPI

    Environment variables:
        SHOPIFY_SHOP_URL: Shop URL (e.g., 'myshop.myshopify.com')
        SHOPIFY_ACCESS_TOKEN: Access token

    Example:
        loader = ShopifyLoader(shop_url="myshop.myshopify.com", access_token="xxx")
        loader.authenticate()
        docs = loader.load_folder("products")
    """

    def __init__(
        self,
        shop_url: Optional[str] = None,
        access_token: Optional[str] = None,
        api_version: str = "2024-01",
        **kwargs,
    ):
        """Initialize Shopify loader."""
        self._shop_url = shop_url or os.environ.get("SHOPIFY_SHOP_URL")
        self._access_token = access_token or os.environ.get("SHOPIFY_ACCESS_TOKEN")
        self._api_version = api_version
        self._session: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "shopify"

    def authenticate(self) -> bool:
        """Authenticate with Shopify."""
        try:
            import shopify

            shopify.ShopifyResource.set_site(
                f"https://{self._shop_url}/admin/api/{self._api_version}"
            )
            shopify.ShopifyResource.set_headers(
                {"X-Shopify-Access-Token": self._access_token}
            )

            # Test connection
            shopify.Shop.current()
            self._session = shopify
            self._authenticated = True
            logger.info("Shopify authentication successful")
            return True
        except ImportError:
            logger.error("ShopifyAPI not installed. Run: pip install ShopifyAPI")
            return False
        except Exception as e:
            logger.error(f"Shopify authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Shopify object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "product":
                obj = self._session.Product.find(obj_id)
            elif obj_type == "order":
                obj = self._session.Order.find(obj_id)
            elif obj_type == "customer":
                obj = self._session.Customer.find(obj_id)
            elif obj_type == "collection":
                obj = self._session.CustomCollection.find(obj_id)
            else:
                logger.error(f"Unknown Shopify object type: {obj_type}")
                return None

            data = obj.to_dict()
            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load Shopify object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Shopify objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            if folder_lower == "products":
                items = self._session.Product.find(limit=100)
            elif folder_lower == "orders":
                items = self._session.Order.find(limit=100)
            elif folder_lower == "customers":
                items = self._session.Customer.find(limit=100)
            elif folder_lower == "collections":
                items = self._session.CustomCollection.find(limit=100)
            else:
                logger.warning(f"Unknown Shopify folder: {folder_path}")
                return docs

            obj_type = folder_lower.rstrip("s")  # products -> product
            for item in items:
                data = item.to_dict()
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{obj_type}:{item.id}",
                        filename=f"{obj_type}_{item.id}.json",
                        metadata={"type": obj_type, "id": str(item.id)},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load Shopify {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Shopify products."""
        self._ensure_authenticated()
        docs = []

        try:
            products = self._session.Product.find(title=query, limit=max_results)
            for prod in products:
                data = prod.to_dict()
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"product:{prod.id}",
                        filename=f"product_{prod.id}.json",
                        metadata={"type": "product", "id": str(prod.id)},
                    )
                )
        except Exception as e:
            logger.error(f"Shopify search failed: {e}")

        return docs


class WooCommerceLoader(BaseLoader):
    """Document loader for WooCommerce (WordPress e-commerce).

    Load products, orders, and customers from WooCommerce.

    Requirements:
        pip install woocommerce

    Environment variables:
        WOOCOMMERCE_URL: Store URL
        WOOCOMMERCE_CONSUMER_KEY: Consumer key
        WOOCOMMERCE_CONSUMER_SECRET: Consumer secret

    Example:
        loader = WooCommerceLoader(url="https://myshop.com", consumer_key="ck_xxx", consumer_secret="cs_yyy")
        loader.authenticate()
        docs = loader.load_folder("products")
    """

    def __init__(
        self,
        url: Optional[str] = None,
        consumer_key: Optional[str] = None,
        consumer_secret: Optional[str] = None,
        **kwargs,
    ):
        """Initialize WooCommerce loader."""
        self._url = url or os.environ.get("WOOCOMMERCE_URL")
        self._consumer_key = consumer_key or os.environ.get("WOOCOMMERCE_CONSUMER_KEY")
        self._consumer_secret = consumer_secret or os.environ.get(
            "WOOCOMMERCE_CONSUMER_SECRET"
        )
        self._api: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "woocommerce"

    def authenticate(self) -> bool:
        """Authenticate with WooCommerce."""
        try:
            from woocommerce import API

            self._api = API(
                url=self._url,
                consumer_key=self._consumer_key,
                consumer_secret=self._consumer_secret,
                version="wc/v3",
            )
            # Test connection
            self._api.get("system_status")
            self._authenticated = True
            logger.info("WooCommerce authentication successful")
            return True
        except ImportError:
            logger.error("woocommerce not installed. Run: pip install woocommerce")
            return False
        except Exception as e:
            logger.error(f"WooCommerce authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single WooCommerce object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)
            data = self._api.get(f"{obj_type}s/{obj_id}").json()

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load WooCommerce object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load WooCommerce objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        endpoint_map = {
            "products": "products",
            "orders": "orders",
            "customers": "customers",
        }

        if folder_lower not in endpoint_map:
            logger.warning(f"Unknown WooCommerce folder: {folder_path}")
            return docs

        try:
            endpoint = endpoint_map[folder_lower]
            items = self._api.get(endpoint, params={"per_page": 100}).json()
            obj_type = folder_lower.rstrip("s")

            for item in items:
                docs.append(
                    LoadedDocument(
                        content=json.dumps(item, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{obj_type}:{item['id']}",
                        filename=f"{obj_type}_{item['id']}.json",
                        metadata={"type": obj_type, "id": str(item["id"])},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load WooCommerce {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search WooCommerce products."""
        self._ensure_authenticated()
        docs = []

        try:
            products = self._api.get(
                "products", params={"search": query, "per_page": max_results}
            ).json()
            for prod in products:
                docs.append(
                    LoadedDocument(
                        content=json.dumps(prod, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"product:{prod['id']}",
                        filename=f"product_{prod['id']}.json",
                        metadata={"type": "product", "id": str(prod["id"])},
                    )
                )
        except Exception as e:
            logger.error(f"WooCommerce search failed: {e}")

        return docs


class BigCommerceLoader(BaseLoader):
    """Document loader for BigCommerce e-commerce platform.

    Load products, orders, and customers from BigCommerce.

    Requirements:
        pip install bigcommerce

    Environment variables:
        BIGCOMMERCE_STORE_HASH: Store hash
        BIGCOMMERCE_ACCESS_TOKEN: Access token

    Example:
        loader = BigCommerceLoader(store_hash="xxx", access_token="yyy")
        loader.authenticate()
        docs = loader.load_folder("products")
    """

    def __init__(
        self,
        store_hash: Optional[str] = None,
        access_token: Optional[str] = None,
        **kwargs,
    ):
        """Initialize BigCommerce loader."""
        self._store_hash = store_hash or os.environ.get("BIGCOMMERCE_STORE_HASH")
        self._access_token = access_token or os.environ.get("BIGCOMMERCE_ACCESS_TOKEN")
        self._api: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "bigcommerce"

    def authenticate(self) -> bool:
        """Authenticate with BigCommerce."""
        try:
            import bigcommerce

            self._api = bigcommerce.api.BigcommerceApi(
                store_hash=self._store_hash,
                access_token=self._access_token,
            )
            # Test connection
            self._api.Store.all()
            self._authenticated = True
            logger.info("BigCommerce authentication successful")
            return True
        except ImportError:
            logger.error("bigcommerce not installed. Run: pip install bigcommerce")
            return False
        except Exception as e:
            logger.error(f"BigCommerce authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single BigCommerce object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)

            if obj_type == "product":
                obj = self._api.Products.get(int(obj_id))
            elif obj_type == "order":
                obj = self._api.Orders.get(int(obj_id))
            elif obj_type == "customer":
                obj = self._api.Customers.get(int(obj_id))
            else:
                logger.error(f"Unknown BigCommerce object type: {obj_type}")
                return None

            data = dict(obj) if hasattr(obj, "__dict__") else obj

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load BigCommerce object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load BigCommerce objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            if folder_lower == "products":
                items = self._api.Products.all(limit=100)
            elif folder_lower == "orders":
                items = self._api.Orders.all(limit=100)
            elif folder_lower == "customers":
                items = self._api.Customers.all(limit=100)
            else:
                logger.warning(f"Unknown BigCommerce folder: {folder_path}")
                return docs

            obj_type = folder_lower.rstrip("s")
            for item in items:
                data = dict(item) if hasattr(item, "__dict__") else item
                item_id = getattr(item, "id", None) or data.get("id")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{obj_type}:{item_id}",
                        filename=f"{obj_type}_{item_id}.json",
                        metadata={"type": obj_type, "id": str(item_id)},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load BigCommerce {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search BigCommerce products."""
        self._ensure_authenticated()
        docs = []

        try:
            products = self._api.Products.all(name=query, limit=max_results)
            for prod in products:
                data = dict(prod) if hasattr(prod, "__dict__") else prod
                prod_id = getattr(prod, "id", None) or data.get("id")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"product:{prod_id}",
                        filename=f"product_{prod_id}.json",
                        metadata={"type": "product", "id": str(prod_id)},
                    )
                )
        except Exception as e:
            logger.error(f"BigCommerce search failed: {e}")

        return docs


class MagentoLoader(BaseLoader):
    """Document loader for Magento e-commerce platform.

    Load products, orders, and customers from Magento 2.

    Requirements:
        pip install requests

    Environment variables:
        MAGENTO_BASE_URL: Magento base URL
        MAGENTO_ACCESS_TOKEN: Access token (integration token)

    Example:
        loader = MagentoLoader(base_url="https://magento.myshop.com", access_token="xxx")
        loader.authenticate()
        docs = loader.load_folder("products")
    """

    def __init__(
        self,
        base_url: Optional[str] = None,
        access_token: Optional[str] = None,
        **kwargs,
    ):
        """Initialize Magento loader."""
        self._base_url = (base_url or os.environ.get("MAGENTO_BASE_URL", "")).rstrip(
            "/"
        )
        self._access_token = access_token or os.environ.get("MAGENTO_ACCESS_TOKEN")
        self._session: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "magento"

    def authenticate(self) -> bool:
        """Authenticate with Magento API."""
        try:
            import requests

            self._session = requests.Session()
            self._session.headers.update(
                {
                    "Authorization": f"Bearer {self._access_token}",
                    "Accept": "application/json",
                }
            )

            # Test connection
            resp = self._session.get(f"{self._base_url}/rest/V1/store/storeViews")
            resp.raise_for_status()
            self._authenticated = True
            logger.info("Magento authentication successful")
            return True
        except Exception as e:
            logger.error(f"Magento authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single Magento object."""
        self._ensure_authenticated()

        try:
            obj_type, obj_id = doc_id.split(":", 1)
            endpoint_map = {
                "product": f"products/{obj_id}",
                "order": f"orders/{obj_id}",
                "customer": f"customers/{obj_id}",
            }

            if obj_type not in endpoint_map:
                logger.error(f"Unknown Magento object type: {obj_type}")
                return None

            resp = self._session.get(
                f"{self._base_url}/rest/V1/{endpoint_map[obj_type]}"
            )
            resp.raise_for_status()
            data = resp.json()

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load Magento object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load Magento objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        endpoint_map = {
            "products": ("products", "product"),
            "orders": ("orders", "order"),
            "customers": ("customers/search", "customer"),
        }

        if folder_lower not in endpoint_map:
            logger.warning(f"Unknown Magento folder: {folder_path}")
            return docs

        endpoint, obj_type = endpoint_map[folder_lower]

        try:
            params = {"searchCriteria[pageSize]": 100}
            resp = self._session.get(
                f"{self._base_url}/rest/V1/{endpoint}", params=params
            )
            resp.raise_for_status()
            result = resp.json()

            items = result.get("items", result) if isinstance(result, dict) else result

            for item in items:
                item_id = item.get("entity_id") or item.get("id") or item.get("sku")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(item, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{obj_type}:{item_id}",
                        filename=f"{obj_type}_{item_id}.json",
                        metadata={"type": obj_type, "id": str(item_id)},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load Magento {folder_path}: {e}")

        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search Magento products."""
        self._ensure_authenticated()
        docs = []

        try:
            params = {
                "searchCriteria[filterGroups][0][filters][0][field]": "name",
                "searchCriteria[filterGroups][0][filters][0][value]": f"%{query}%",
                "searchCriteria[filterGroups][0][filters][0][conditionType]": "like",
                "searchCriteria[pageSize]": max_results,
            }
            resp = self._session.get(
                f"{self._base_url}/rest/V1/products", params=params
            )
            resp.raise_for_status()

            for item in resp.json().get("items", []):
                item_id = item.get("sku")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(item, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"product:{item_id}",
                        filename=f"product_{item_id}.json",
                        metadata={"type": "product", "id": str(item_id)},
                    )
                )
        except Exception as e:
            logger.error(f"Magento search failed: {e}")

        return docs


# ============================================================================
# ACCOUNTING LOADERS
# ============================================================================


class QuickBooksLoader(BaseLoader):
    """Document loader for QuickBooks accounting software.

    Load invoices, customers, bills, accounts, and reports from QuickBooks Online.

    Requirements:
        pip install python-quickbooks intuit-oauth

    Environment variables:
        QUICKBOOKS_CLIENT_ID: OAuth client ID
        QUICKBOOKS_CLIENT_SECRET: OAuth client secret
        QUICKBOOKS_REFRESH_TOKEN: OAuth refresh token
        QUICKBOOKS_REALM_ID: Company ID
        QUICKBOOKS_ENVIRONMENT: 'sandbox' or 'production'

    Example:
        loader = QuickBooksLoader(client_id="xxx", client_secret="yyy", realm_id="zzz")
        loader.authenticate()
        docs = loader.load_folder("invoices")
    """

    def __init__(
        self,
        client_id: Optional[str] = None,
        client_secret: Optional[str] = None,
        refresh_token: Optional[str] = None,
        realm_id: Optional[str] = None,
        environment: Optional[str] = None,
        **kwargs,
    ):
        """Initialize QuickBooks loader."""
        self._client_id = client_id or os.environ.get("QUICKBOOKS_CLIENT_ID")
        self._client_secret = client_secret or os.environ.get(
            "QUICKBOOKS_CLIENT_SECRET"
        )
        self._refresh_token = refresh_token or os.environ.get(
            "QUICKBOOKS_REFRESH_TOKEN"
        )
        self._realm_id = realm_id or os.environ.get("QUICKBOOKS_REALM_ID")
        self._environment = environment or os.environ.get(
            "QUICKBOOKS_ENVIRONMENT", "production"
        )
        self._client: Optional[Any] = None
        self._authenticated = False

    @property
    def source_name(self) -> str:
        return "quickbooks"

    def authenticate(self) -> bool:
        """Authenticate with QuickBooks API."""
        try:
            from intuitlib.client import AuthClient
            from quickbooks import QuickBooks

            auth_client = AuthClient(
                client_id=self._client_id,
                client_secret=self._client_secret,
                environment=self._environment,
                redirect_uri="http://localhost:8000/callback",
            )
            auth_client.refresh(refresh_token=self._refresh_token)

            self._client = QuickBooks(
                auth_client=auth_client,
                refresh_token=self._refresh_token,
                company_id=self._realm_id,
            )
            self._authenticated = True
            logger.info("QuickBooks authentication successful")
            return True
        except ImportError:
            logger.error(
                "python-quickbooks not installed. Run: pip install python-quickbooks intuit-oauth"
            )
            return False
        except Exception as e:
            logger.error(f"QuickBooks authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single QuickBooks object."""
        self._ensure_authenticated()

        try:
            from quickbooks.objects import (
                Account,
                Bill,
                Customer,
                Invoice,
            )

            obj_type, obj_id = doc_id.split(":", 1)
            type_map = {
                "invoice": Invoice,
                "customer": Customer,
                "bill": Bill,
                "account": Account,
            }

            if obj_type not in type_map:
                logger.error(f"Unknown QuickBooks object type: {obj_type}")
                return None

            obj_class = type_map[obj_type]
            obj = obj_class.get(int(obj_id), qb=self._client)
            data = obj.to_dict() if hasattr(obj, "to_dict") else vars(obj)

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=doc_id,
                filename=f"{obj_type}_{obj_id}.json",
                metadata={"type": obj_type, "id": obj_id},
            )
        except Exception as e:
            logger.error(f"Failed to load QuickBooks object {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load QuickBooks objects by type."""
        self._ensure_authenticated()
        docs = []
        folder_lower = folder_path.lower()

        try:
            from quickbooks.objects import (
                Account,
                Bill,
                Customer,
                Invoice,
            )

            type_map = {
                "invoices": (Invoice, "invoice"),
                "customers": (Customer, "customer"),
                "bills": (Bill, "bill"),
                "accounts": (Account, "account"),
            }

            if folder_lower not in type_map:
                logger.warning(
                    f"Unknown QuickBooks folder: {folder_path}. Use: {list(type_map.keys())}"
                )
                return docs

            obj_class, obj_type = type_map[folder_lower]
            items = obj_class.all(qb=self._client, max_results=100)

            for item in items:
                data = item.to_dict() if hasattr(item, "to_dict") else vars(item)
                item_id = getattr(item, "Id", None) or data.get("Id")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"{obj_type}:{item_id}",
                        filename=f"{obj_type}_{item_id}.json",
                        metadata={"type": obj_type, "id": str(item_id)},
                    )
                )
        except Exception as e:
            logger.error(f"Failed to load QuickBooks {folder_path}: {e}")

        return docs

    def load_report(self, report_name: str) -> Optional[LoadedDocument]:
        """Load a QuickBooks report.

        Args:
            report_name: Report type (e.g., 'ProfitAndLoss', 'BalanceSheet')
        """
        self._ensure_authenticated()

        try:
            from quickbooks.objects import (
                BalanceSheet,
                ProfitAndLoss,
            )

            report_map = {
                "profitandloss": ProfitAndLoss,
                "profit_and_loss": ProfitAndLoss,
                "pnl": ProfitAndLoss,
                "balancesheet": BalanceSheet,
                "balance_sheet": BalanceSheet,
            }

            report_class = report_map.get(report_name.lower())
            if not report_class:
                logger.error(f"Unknown QuickBooks report: {report_name}")
                return None

            report = report_class.get(qb=self._client)
            data = report.to_dict() if hasattr(report, "to_dict") else vars(report)

            return LoadedDocument(
                content=json.dumps(data, indent=2, default=str),
                source=self.source_name,
                source_id=f"report:{report_name}",
                filename=f"report_{report_name}.json",
                metadata={"type": "report", "report_name": report_name},
            )
        except Exception as e:
            logger.error(f"Failed to load QuickBooks report {report_name}: {e}")
            return None

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search QuickBooks customers."""
        self._ensure_authenticated()
        docs = []

        try:
            from quickbooks.objects import Customer

            customers = Customer.query(
                f"SELECT * FROM Customer WHERE DisplayName LIKE '%{query}%' MAXRESULTS {max_results}",
                qb=self._client,
            )
            for cust in customers:
                data = cust.to_dict() if hasattr(cust, "to_dict") else vars(cust)
                cust_id = getattr(cust, "Id", None) or data.get("Id")
                docs.append(
                    LoadedDocument(
                        content=json.dumps(data, indent=2, default=str),
                        source=self.source_name,
                        source_id=f"customer:{cust_id}",
                        filename=f"customer_{cust_id}.json",
                        metadata={"type": "customer", "id": str(cust_id)},
                    )
                )
        except Exception as e:
            logger.error(f"QuickBooks search failed: {e}")

        return docs


# ============================================================================
# FACTORY FUNCTIONS
# ============================================================================


def create_loader(source: str, **kwargs) -> BaseLoader:
    """Factory function to create a document loader.

    Args:
        source: One of 'google_drive', 'gmail', 'icloud', 'firestore', 's3', 'mongodb',
                'github', 'microsoft365'
        **kwargs: Arguments passed to the loader constructor

    Returns:
        BaseLoader instance

    Example:
        loader = create_loader('google_drive', credentials_path='creds.json')
        loader = create_loader('firestore', service_account_path='firebase.json')
        loader = create_loader('s3', bucket='my-docs', endpoint_url='http://minio:9000')
        loader = create_loader('mongodb', uri='mongodb://localhost:27017', database='docs')
        loader = create_loader('github', token='ghp_xxx')
        loader = create_loader('microsoft365', client_id='xxx', client_secret='yyy', tenant_id='zzz')
    """
    loaders = {
        # Google
        "google_drive": GoogleDriveLoader,
        "drive": GoogleDriveLoader,
        "gmail": GmailLoader,
        "email": GmailLoader,
        # Apple
        "icloud": iCloudLoader,
        "icloud_drive": iCloudLoader,
        # Firebase/Google Cloud
        "firestore": FirestoreLoader,
        "firebase": FirestoreLoader,
        "gcs": GCSLoader,
        "google_cloud_storage": GCSLoader,
        # AWS/Self-hosted
        "s3": S3Loader,
        "minio": MinIOLoader,
        "aws": S3Loader,
        "aws_s3": S3Loader,
        # Azure
        "azure_blob": AzureBlobLoader,
        "azure_storage": AzureBlobLoader,
        "azure": AzureBlobLoader,
        # Databases
        "mongodb": MongoDBLoader,
        "mongo": MongoDBLoader,
        "postgresql": PostgreSQLLoader,
        "postgres": PostgreSQLLoader,
        "pg": PostgreSQLLoader,
        "mysql": MySQLLoader,
        "mariadb": MySQLLoader,
        "elasticsearch": ElasticsearchLoader,
        "elastic": ElasticsearchLoader,
        "es": ElasticsearchLoader,
        "redis": RedisLoader,
        "oracle": OracleLoader,
        "oracle_db": OracleLoader,
        # GitHub
        "github": GitHubLoader,
        "gh": GitHubLoader,
        # Microsoft
        "microsoft365": Microsoft365Loader,
        "m365": Microsoft365Loader,
        "office365": Microsoft365Loader,
        "o365": Microsoft365Loader,
        "onedrive": OneDriveLoader,
        "onedrive_personal": OneDriveLoader,
        "sharepoint": SharePointLoader,
        "sharepoint_online": SharePointLoader,
        "teams": TeamsLoader,
        "ms_teams": TeamsLoader,
        "dynamics365": Dynamics365Loader,
        "dynamics": Dynamics365Loader,
        "d365": Dynamics365Loader,
        # Cloud storage
        "dropbox": DropboxLoader,
        "box": BoxLoader,
        "box_com": BoxLoader,
        # Productivity
        "notion": NotionLoader,
        "confluence": ConfluenceLoader,
        "atlassian": ConfluenceLoader,
        "jira_wiki": ConfluenceLoader,
        # Communication
        "slack": SlackLoader,
        "discord": DiscordLoader,
        # Project management
        "jira": JiraLoader,
        "atlassian_jira": JiraLoader,
        "jira_service_desk": JiraServiceDeskLoader,
        "jira_sd": JiraServiceDeskLoader,
        "jsm": JiraServiceDeskLoader,
        "asana": AsanaLoader,
        "trello": TrelloLoader,
        # Database/spreadsheet
        "airtable": AirtableLoader,
        # CRM
        "hubspot": HubSpotLoader,
        "salesforce": SalesforceLoader,
        "sfdc": SalesforceLoader,
        "zoho": ZohoLoader,
        "zoho_crm": ZohoLoader,
        # Support
        "zendesk": ZendeskLoader,
        "intercom": IntercomLoader,
        "freshdesk": FreshdeskLoader,
        # Enterprise Systems
        "sap": SAPLoader,
        "sap_erp": SAPLoader,
        "s4hana": SAPLoader,
        "servicenow": ServiceNowLoader,
        "snow": ServiceNowLoader,
        "workday": WorkdayLoader,
        # Australian Business Software
        "myob": MYOBLoader,
        "myob_accountright": MYOBLoader,
        "xero": XeroLoader,
        "employment_hero": EmploymentHeroLoader,
        "eh": EmploymentHeroLoader,
        "deputy": DeputyLoader,
        # Payment Gateways
        "stripe": StripeLoader,
        "paypal": PayPalLoader,
        "paypal_rest": PayPalLoader,
        "square": SquareLoader,
        "square_pos": SquareLoader,
        "afterpay": AfterpayLoader,
        "afterpay_au": AfterpayLoader,
        "clearpay": AfterpayLoader,  # Afterpay's UK/EU name
        "braintree": BraintreeLoader,
        "braintree_paypal": BraintreeLoader,
        # E-Commerce Platforms
        "shopify": ShopifyLoader,
        "shopify_store": ShopifyLoader,
        "woocommerce": WooCommerceLoader,
        "woo": WooCommerceLoader,
        "wordpress_ecommerce": WooCommerceLoader,
        "bigcommerce": BigCommerceLoader,
        "big_commerce": BigCommerceLoader,
        "magento": MagentoLoader,
        "magento2": MagentoLoader,
        "adobe_commerce": MagentoLoader,
        # Accounting
        "quickbooks": QuickBooksLoader,
        "quickbooks_online": QuickBooksLoader,
        "qbo": QuickBooksLoader,
        "intuit_quickbooks": QuickBooksLoader,
    }

    source_lower = source.lower().replace("-", "_").replace(" ", "_")

    if source_lower not in loaders:
        raise ValueError(f"Unknown source: {source}. Available: {list(loaders.keys())}")

    return loaders[source_lower](**kwargs)


def load_from_multiple_sources(
    sources: list[dict[str, Any]], deduplicate: bool = True
) -> list[LoadedDocument]:
    """Load documents from multiple sources.

    Args:
        sources: List of dicts with 'type', 'credentials', and optional 'folder'/'query'
        deduplicate: Remove duplicate documents based on content hash

    Example:
        docs = load_from_multiple_sources([
            {'type': 'google_drive', 'credentials_path': 'creds.json', 'folder': 'Work'},
            {'type': 'gmail', 'credentials_path': 'creds.json', 'query': 'from:boss'},
            {'type': 'icloud', 'apple_id': 'me@icloud.com', 'password': 'xxx', 'folder': 'Documents'},
        ])
    """
    all_docs = []
    seen_hashes = set()

    for source_config in sources:
        source_config = source_config.copy()  # Don't mutate original
        source_type = source_config.pop("type")
        folder = source_config.pop("folder", None)
        query = source_config.pop("query", None)

        try:
            loader = create_loader(source_type, **source_config)

            if folder:
                docs = loader.load_folder(folder)
            elif query:
                docs = loader.search(query)
            else:
                logger.warning(f"No folder or query specified for {source_type}")
                continue

            for doc in docs:
                if deduplicate:
                    import hashlib

                    content_hash = hashlib.sha256(doc.content.encode()).hexdigest()
                    if content_hash in seen_hashes:
                        continue
                    seen_hashes.add(content_hash)

                all_docs.append(doc)

        except Exception as e:
            logger.error(f"Failed to load from {source_type}: {e}")

    return all_docs
