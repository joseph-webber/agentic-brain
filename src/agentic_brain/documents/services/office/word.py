# SPDX-License-Identifier: Apache-2.0

"""Enterprise DOCX processing built on top of python-docx.

This module provides a production-oriented :class:`WordProcessor` for modern
Microsoft Word documents. The implementation intentionally uses the shared
Office models from :mod:`agentic_brain.documents.services.office.models`
instead of maintaining a Word-specific data model. That keeps the processor
compatible with the rest of the Office services layer and the unified RAG
pipeline.

Highlights
----------
- full text extraction in document order
- structured paragraph extraction with styles, heading detection, numbering,
  bookmarks, comments, and hyperlink metadata
- table extraction delegated to :class:`.tables.TableExtractor`
- image extraction delegated to :class:`.images.ImageExtractor`
- core/app/custom metadata extraction
- document structure analysis with headings, sections, outline, and WCAG status
- document creation from shared Office models
- mail-merge style templating with ``{{ placeholder }}`` expressions

The processor is best-effort by design. Some advanced DOCX features such as
tracked changes, complex content controls, and full-fidelity mail merge are not
fully exposed by python-docx, so this module supplements the public API with
careful OOXML inspection where necessary.
"""

from __future__ import annotations

import logging
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Iterator, Mapping, Optional, Sequence
from xml.etree import ElementTree as ET

from .accessibility import OfficeAccessibilityProcessor, WCAGReport
from .exceptions import DocumentValidationError, InvalidDocumentStructureError
from .images import Image as ExtractedImage
from .images import ImageExtractor
from .models import (
    Comment,
    DocumentContent,
    DocumentStyle,
    Image,
    Metadata,
    OfficeFormat,
    Paragraph,
    Table,
    TableCell,
    TextRun,
)
from .tables import Cell as ExtractedCell
from .tables import Table as ExtractedTable
from .tables import TableExtractor

try:  # pragma: no cover - optional dependency boundary
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Inches, Pt, RGBColor
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.text.run import Run as DocxRun

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency boundary
    DocxDocument = None
    DocxDocumentType = Any
    WD_STYLE_TYPE = None
    WD_ALIGN_PARAGRAPH = None
    qn = None
    CT_Tbl = None
    CT_P = None
    Inches = None
    Pt = None
    RGBColor = None
    DocxTable = Any
    DocxParagraph = Any
    DocxRun = Any
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

WORD_NS = {
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}

PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_.\-\[\]]+)\s*}}")


@dataclass(slots=True)
class HeadingNode:
    """Simplified heading structure node."""

    level: int
    text: str
    paragraph_index: int
    paragraph_id: str | None = None
    children: list["HeadingNode"] = None  # type: ignore[assignment]

    def __post_init__(self) -> None:
        if self.children is None:
            self.children = []


class WordProcessorError(Exception):
    """Base exception for DOCX processing."""


class WordNotFoundError(WordProcessorError):
    """Raised when a DOCX file does not exist."""


class WordCorruptedError(WordProcessorError):
    """Raised when a DOCX file cannot be parsed."""


class WordProcessor:
    """High-level DOCX processor using shared Office models."""

    def __init__(
        self,
        path: str | Path | None = None,
        *,
        accessibility_processor: OfficeAccessibilityProcessor | None = None,
    ) -> None:
        self._require_dependency()
        self._document: DocxDocumentType | None = None
        self._path: Path | None = None
        self._last_wcag_report: WCAGReport | None = None
        self.table_extractor = TableExtractor()
        self.image_extractor = ImageExtractor()
        self.accessibility_processor = accessibility_processor or OfficeAccessibilityProcessor()
        if path is not None:
            self.load(path)

    @property
    def last_wcag_report(self) -> WCAGReport | None:
        """Return the most recent cached WCAG report, if one exists."""

        return self._last_wcag_report

    # ------------------------------------------------------------------ #
    # Loading and parsing
    # ------------------------------------------------------------------ #
    def load(self, path: str | Path) -> "WordProcessor":
        """Load a DOCX document into memory."""

        source = self._validate_source(path)
        try:
            self._document = DocxDocument(str(source))
        except Exception as exc:  # pragma: no cover - library error boundary
            raise WordCorruptedError(f"Failed to open DOCX document: {source}") from exc
        self._path = source
        self._refresh_wcag_report(source)
        return self

    def parse(
        self,
        path: str | Path | None = None,
        *,
        include_accessibility: bool = True,
    ) -> DocumentContent:
        """Parse a DOCX file into the shared Office document model."""

        document = self._ensure_document(path)
        source = self._coerce_source(path)
        paragraphs = self.extract_paragraphs()
        structure = self.get_document_structure(include_accessibility=include_accessibility)
        content = DocumentContent(
            format=OfficeFormat.DOCX,
            paragraphs=paragraphs,
            tables=self.extract_tables(),
            images=self.extract_images(),
            comments=self.extract_comments(),
            metadata=self.get_metadata(),
            styles=self.get_styles(),
            sections=[section["title"] for section in structure["sections"]],
            document_properties={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "heading_count": len(structure["headings"]),
                "outline_depth": max((node["level"] for node in structure["headings"]), default=0),
            },
        )
        if source is not None and include_accessibility and structure.get("accessibility"):
            content.document_properties.update(structure["accessibility"])
        return content

    def extract_text(
        self,
        path: str | Path | None = None,
        *,
        include_tables: bool = True,
        include_headers: bool = False,
        include_footers: bool = False,
    ) -> str:
        """Extract full text from a DOCX document in document order."""

        document = self._ensure_document(path)
        parts: list[str] = []

        for block in self._iter_block_items(document):
            if isinstance(block, DocxParagraph):
                text = self._clean_text(block.text)
                if text:
                    parts.append(text)
                continue

            if include_tables and isinstance(block, DocxTable):
                for row in block.rows:
                    cells = [self._clean_text(cell.text) for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))

        if include_headers:
            for section in document.sections:
                header_text = "\n".join(
                    self._clean_text(paragraph.text)
                    for paragraph in section.header.paragraphs
                    if self._clean_text(paragraph.text)
                )
                if header_text:
                    parts.append(header_text)

        if include_footers:
            for section in document.sections:
                footer_text = "\n".join(
                    self._clean_text(paragraph.text)
                    for paragraph in section.footer.paragraphs
                    if self._clean_text(paragraph.text)
                )
                if footer_text:
                    parts.append(footer_text)

        return "\n\n".join(parts)

    def extract_paragraphs(
        self,
        path: str | Path | None = None,
    ) -> list[Paragraph]:
        """Extract structured paragraphs with styles and semantic metadata."""

        document = self._ensure_document(path)
        comment_map = self._load_comments_map(self._coerce_source(path))
        paragraphs: list[Paragraph] = []

        for index, paragraph in enumerate(document.paragraphs):
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = self._heading_level(style_name)
            numbering_level, numbering_id = self._numbering_info(paragraph)
            bookmark_id = self._bookmark_id(paragraph)
            style = self._paragraph_style(paragraph)
            comments = self._paragraph_comments(paragraph, comment_map)

            paragraph_model = Paragraph(
                runs=[self._run_to_model(run, document) for run in paragraph.runs],
                style=style,
                paragraph_id=self._paragraph_id(paragraph) or f"p-{index + 1}",
                numbering_level=numbering_level,
                numbering_id=numbering_id,
                is_heading=heading_level is not None,
                heading_level=heading_level,
                comments=comments,
                bookmarked=bookmark_id is not None,
                bookmark_id=bookmark_id,
            )

            if not paragraph_model.runs and paragraph.text:
                paragraph_model.runs.append(TextRun(text=paragraph.text))

            paragraphs.append(paragraph_model)

        return paragraphs

    def extract_tables(self, path: str | Path | None = None) -> list[Table]:
        """Extract tables using the shared :class:`TableExtractor` service."""

        source = self._require_source(path)
        extracted = self.table_extractor.extract_tables(source)
        return [self._convert_table(table, index) for index, table in enumerate(extracted, start=1)]

    def extract_images(self, path: str | Path | None = None) -> list[Image]:
        """Extract images using the shared :class:`ImageExtractor` service."""

        source = self._require_source(path)
        extracted = self.image_extractor.extract_images(source)
        return [self._convert_image(image) for image in extracted]

    def get_metadata(self, path: str | Path | None = None) -> Metadata:
        """Return core, app, and custom metadata for a DOCX file."""

        document = self._ensure_document(path)
        source = self._coerce_source(path)
        core = document.core_properties
        app_properties = self._load_app_properties(source)
        custom_properties = self._load_custom_properties(source)

        keywords = self._split_keywords(getattr(core, "keywords", None))
        company = app_properties.get("Company")
        category = getattr(core, "category", None) or app_properties.get("Category")

        return Metadata(
            title=getattr(core, "title", None) or None,
            subject=getattr(core, "subject", None) or None,
            author=getattr(core, "author", None) or None,
            company=str(company) if company else None,
            category=str(category) if category else None,
            keywords=keywords,
            created_at=getattr(core, "created", None),
            modified_at=getattr(core, "modified", None),
            last_printed_at=getattr(core, "last_printed", None),
            revision=str(getattr(core, "revision", "")) or None,
            custom_properties=custom_properties,
        )

    extract_metadata = get_metadata

    def extract_comments(self, path: str | Path | None = None) -> list[Comment]:
        """Extract document comments from ``word/comments.xml`` when present."""

        comment_map = self._load_comments_map(self._coerce_source(path))
        return list(comment_map.values())

    def get_styles(self, path: str | Path | None = None) -> dict[str, DocumentStyle]:
        """Return all document styles keyed by style name."""

        document = self._ensure_document(path)
        styles: dict[str, DocumentStyle] = {}

        for style in document.styles:
            style_name = getattr(style, "name", None)
            if not style_name:
                continue

            font = getattr(style, "font", None)
            paragraph_format = getattr(style, "paragraph_format", None)
            styles[style_name] = DocumentStyle(
                font_family=getattr(font, "name", None) or "Calibri",
                font_size=self._pt_value(getattr(font, "size", None), default=12.0),
                bold=bool(getattr(font, "bold", False)),
                italic=bool(getattr(font, "italic", False)),
                underline=bool(getattr(font, "underline", False)),
                strikethrough=bool(getattr(font, "strike", False)),
                text_color=self._rgb_to_hex(getattr(getattr(font, "color", None), "rgb", None)),
                alignment=self._alignment_name(getattr(paragraph_format, "alignment", None)),
                line_spacing=self._float_value(getattr(paragraph_format, "line_spacing", None), 1.15),
                spacing_before=self._pt_value(getattr(paragraph_format, "space_before", None), 0.0),
                spacing_after=self._pt_value(getattr(paragraph_format, "space_after", None), 0.0),
                indentation_left=self._pt_value(getattr(paragraph_format, "left_indent", None), 0.0),
                indentation_right=self._pt_value(getattr(paragraph_format, "right_indent", None), 0.0),
                indentation_first_line=self._pt_value(
                    getattr(paragraph_format, "first_line_indent", None),
                    0.0,
                ),
                styles={
                    "style_name": style_name,
                    "style_id": getattr(style, "style_id", None),
                    "style_type": self._style_type_name(getattr(style, "type", None)),
                    "base_style": getattr(getattr(style, "base_style", None), "name", None),
                },
            )

        return styles

    def get_document_structure(
        self,
        path: str | Path | None = None,
        *,
        include_accessibility: bool = True,
        wcag_level: str = "AA",
    ) -> dict[str, Any]:
        """Return headings, sections, outline, and optional WCAG summary."""

        document = self._ensure_document(path)
        paragraphs = self.extract_paragraphs()
        headings = [
            {
                "level": paragraph.heading_level,
                "text": paragraph.text_content(),
                "paragraph_id": paragraph.paragraph_id,
                "paragraph_index": index,
            }
            for index, paragraph in enumerate(paragraphs)
            if paragraph.is_heading and paragraph.heading_level is not None
        ]
        outline = self._build_outline(headings)
        sections = self._extract_sections(document, headings)

        structure: dict[str, Any] = {
            "headings": headings,
            "outline": outline,
            "sections": sections,
            "has_title": bool(self.get_metadata().title),
        }

        source = self._coerce_source(path)
        if include_accessibility and source is not None:
            report = self._last_wcag_report
            if report is None or report.document_path != source or report.level != wcag_level.upper():
                try:
                    report = self.accessibility_processor.check_wcag_compliance(
                        source,
                        level=wcag_level,
                    )
                    self._last_wcag_report = report
                except Exception as exc:  # pragma: no cover - best effort integration
                    logger.warning("WCAG analysis failed for %s: %s", source, exc)
                    report = None
            if report is not None:
                structure["accessibility"] = {
                    "wcag_level": report.level,
                    "wcag_score": report.score,
                    "is_compliant": report.is_compliant,
                    "issue_count": len(report.issues),
                    "recommendations": report.recommendations,
                }

        return structure

    def extract_outline(self, path: str | Path | None = None) -> list[dict[str, Any]]:
        """Compatibility helper returning just the document outline."""

        return self.get_document_structure(path).get("outline", [])

    # ------------------------------------------------------------------ #
    # Document creation and templating
    # ------------------------------------------------------------------ #
    def create_document(
        self,
        content: DocumentContent | Mapping[str, Any] | Iterable[str] | None = None,
        *,
        metadata: Metadata | Mapping[str, Any] | None = None,
        output_path: str | Path | None = None,
    ) -> DocxDocumentType:
        """Build a new DOCX document from content using shared Office models."""

        document = DocxDocument()
        self._document = document
        self._path = None

        normalized_content = self._normalize_content(content)
        normalized_metadata = self._normalize_metadata(metadata) or normalized_content.metadata

        if normalized_metadata.title is None:
            first_heading = next(
                (paragraph.text_content() for paragraph in normalized_content.paragraphs if paragraph.is_heading),
                None,
            )
            normalized_metadata.title = first_heading or "Document"
        self._apply_metadata(document, normalized_metadata)

        for paragraph in normalized_content.paragraphs:
            self.add_paragraph(paragraph)
        for table in normalized_content.tables:
            self.add_table(table)
        for image in normalized_content.images:
            self.add_image(image)

        if output_path is not None:
            self.save(output_path)

        return document

    def apply_template(
        self,
        template: str | Path,
        context: Mapping[str, Any],
        *,
        output_path: str | Path | None = None,
        strict: bool = False,
    ) -> DocxDocumentType:
        """Apply mail-merge style placeholders to a DOCX template."""

        source = self._validate_source(template)
        document = DocxDocument(str(source))
        self._document = document
        self._path = source

        for paragraph in document.paragraphs:
            self._render_paragraph(paragraph, context, strict=strict)
        for table in document.tables:
            self._render_table(table, context, strict=strict)
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                self._render_paragraph(paragraph, context, strict=strict)
            for table in section.header.tables:
                self._render_table(table, context, strict=strict)
            for paragraph in section.footer.paragraphs:
                self._render_paragraph(paragraph, context, strict=strict)
            for table in section.footer.tables:
                self._render_table(table, context, strict=strict)

        self._render_core_properties(document, context, strict=strict)

        if output_path is not None:
            self.save(output_path)

        return document

    def add_paragraph(
        self,
        paragraph: Paragraph | str,
        *,
        style_name: str | None = None,
    ) -> DocxParagraph:
        """Append a paragraph to the active document."""

        document = self._require_active_document()
        paragraph_model = self._coerce_paragraph(paragraph)
        style = style_name or paragraph_model.style.styles.get("style_name")
        if paragraph_model.is_heading and paragraph_model.heading_level:
            docx_paragraph = document.add_heading("", level=paragraph_model.heading_level)
        else:
            docx_paragraph = document.add_paragraph()

        if style and not paragraph_model.is_heading:
            try:
                docx_paragraph.style = style
            except Exception:  # pragma: no cover - style availability depends on template
                logger.debug("Word style %s unavailable in target document", style)

        self._apply_paragraph_style(docx_paragraph, paragraph_model.style)
        if paragraph_model.runs:
            for run_model in paragraph_model.runs:
                self._append_run(docx_paragraph, run_model)
        else:
            docx_paragraph.add_run(paragraph_model.text_content())
        return docx_paragraph

    def add_table(self, table: Table | Sequence[Sequence[Any]]) -> DocxTable:
        """Append a table to the active document."""

        document = self._require_active_document()
        table_model = self._coerce_table(table)
        row_count = len(table_model.rows)
        col_count = max((len(row) for row in table_model.rows), default=0)
        if row_count == 0 or col_count == 0:
            raise InvalidDocumentStructureError("Cannot create an empty Word table")

        docx_table = document.add_table(rows=row_count, cols=col_count)
        style_name = table_model.style.styles.get("style_name")
        if style_name:
            try:
                docx_table.style = style_name
            except Exception:  # pragma: no cover - style availability depends on template
                logger.debug("Word table style %s unavailable in target document", style_name)

        for row_index, row in enumerate(table_model.rows):
            for col_index, cell in enumerate(row):
                docx_cell = docx_table.cell(row_index, col_index)
                self._populate_docx_cell(docx_cell, cell)

        return docx_table

    def add_image(self, image: Image | str | Path) -> Any:
        """Append an image to the active document with accessibility metadata."""

        document = self._require_active_document()
        image_model = self._coerce_image(image)

        width = self._image_dimension(image_model.width)
        height = self._image_dimension(image_model.height)
        stream = BytesIO(image_model.data)

        if width is not None and height is not None:
            shape = document.add_picture(stream, width=width, height=height)
        elif width is not None:
            shape = document.add_picture(stream, width=width)
        elif height is not None:
            shape = document.add_picture(stream, height=height)
        else:
            shape = document.add_picture(stream)

        self._set_picture_alt_text(
            shape,
            alt_text=image_model.alternate_text or image_model.description or "",
            title=image_model.title or "",
        )

        if image_model.description:
            try:
                caption = document.add_paragraph(style="Caption")
            except Exception:  # pragma: no cover - style availability depends on template
                caption = document.add_paragraph()
            caption.add_run(image_model.description)

        return shape

    def save(self, path: str | Path) -> Path:
        """Persist the active document to disk."""

        document = self._require_active_document()
        target = Path(path).expanduser()
        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(target))
        self._path = target
        self._refresh_wcag_report(target)
        return target

    def save_to_bytes(self) -> bytes:
        """Serialize the active document to bytes."""

        document = self._require_active_document()
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def stream_paragraphs(
        self,
        path: str | Path | None = None,
    ) -> Iterator[Paragraph]:
        """Yield paragraphs one at a time for large-document processing."""

        for paragraph in self.extract_paragraphs(path):
            yield paragraph

    # ------------------------------------------------------------------ #
    # Internal: extraction helpers
    # ------------------------------------------------------------------ #
    def _require_dependency(self) -> None:
        if DOCX_AVAILABLE:
            return
        raise ImportError(
            "python-docx is required for Word processing. "
            "Install with: pip install python-docx"
        )

    def _validate_source(self, path: str | Path) -> Path:
        source = Path(path).expanduser()
        if not source.exists():
            raise WordNotFoundError(f"Word document not found: {source}")
        if source.suffix.lower() != ".docx":
            raise DocumentValidationError("Expected a .docx document", details=str(source))
        return source

    def _coerce_source(self, path: str | Path | None = None) -> Path | None:
        if path is not None:
            return self._validate_source(path)
        return self._path

    def _require_source(self, path: str | Path | None = None) -> Path:
        source = self._coerce_source(path)
        if source is None:
            raise DocumentValidationError(
                "A source path is required for this operation",
                details="Load a document first or pass a DOCX path",
            )
        return source

    def _ensure_document(self, path: str | Path | None = None) -> DocxDocumentType:
        if path is not None:
            return self.load(path)._document  # type: ignore[return-value]
        if self._document is None:
            raise DocumentValidationError("No Word document is loaded")
        return self._document

    def _require_active_document(self) -> DocxDocumentType:
        if self._document is None:
            raise DocumentValidationError(
                "No editable Word document is active",
                details="Call create_document() or apply_template() first",
            )
        return self._document

    def _iter_block_items(self, document: DocxDocumentType) -> Iterator[DocxParagraph | DocxTable]:
        """Yield paragraphs and tables in body order."""

        body = document.element.body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield DocxParagraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield DocxTable(child, document)

    def _paragraph_style(self, paragraph: DocxParagraph) -> DocumentStyle:
        paragraph_format = paragraph.paragraph_format
        base_font = paragraph.runs[0].font if paragraph.runs else None
        style_name = paragraph.style.name if paragraph.style is not None else None
        return DocumentStyle(
            font_family=getattr(base_font, "name", None) or "Calibri",
            font_size=self._pt_value(getattr(base_font, "size", None), default=12.0),
            bold=bool(getattr(base_font, "bold", False)),
            italic=bool(getattr(base_font, "italic", False)),
            underline=bool(getattr(base_font, "underline", False)),
            strikethrough=bool(getattr(base_font, "strike", False)),
            text_color=self._rgb_to_hex(getattr(getattr(base_font, "color", None), "rgb", None)),
            alignment=self._alignment_name(paragraph.alignment),
            line_spacing=self._float_value(getattr(paragraph_format, "line_spacing", None), 1.15),
            spacing_before=self._pt_value(getattr(paragraph_format, "space_before", None), 0.0),
            spacing_after=self._pt_value(getattr(paragraph_format, "space_after", None), 0.0),
            indentation_left=self._pt_value(getattr(paragraph_format, "left_indent", None), 0.0),
            indentation_right=self._pt_value(getattr(paragraph_format, "right_indent", None), 0.0),
            indentation_first_line=self._pt_value(
                getattr(paragraph_format, "first_line_indent", None),
                0.0,
            ),
            styles={"style_name": style_name or "Normal"},
        )

    def _run_to_model(self, run: DocxRun, document: DocxDocumentType) -> TextRun:
        font = run.font
        return TextRun(
            text=run.text,
            style=DocumentStyle(
                font_family=getattr(font, "name", None) or "Calibri",
                font_size=self._pt_value(getattr(font, "size", None), default=12.0),
                bold=bool(run.bold),
                italic=bool(run.italic),
                underline=bool(run.underline),
                strikethrough=bool(getattr(font, "strike", False)),
                text_color=self._rgb_to_hex(getattr(getattr(font, "color", None), "rgb", None)),
                background_color="#FFFFFF",
                styles={
                    "highlight": str(getattr(font, "highlight_color", "")) or "",
                    "subscript": bool(getattr(font, "subscript", False)),
                    "superscript": bool(getattr(font, "superscript", False)),
                },
            ),
            language=self._run_language(run),
            hyperlink=self._run_hyperlink(run, document),
        )

    def _convert_table(self, table: ExtractedTable, index: int) -> Table:
        rows: list[list[TableCell]] = []
        cell_grid: dict[str, TableCell] = {}

        for row_index, row in enumerate(table.rows):
            converted_row: list[TableCell] = []
            for col_index, cell in enumerate(row):
                cell_id = f"table-{index}-r{row_index + 1}-c{col_index + 1}"
                converted_cell = self._convert_table_cell(cell, cell_id)
                converted_row.append(converted_cell)
                cell_grid[cell_id] = converted_cell
            rows.append(converted_row)

        return Table(
            rows=rows,
            alignment="left",
            style=DocumentStyle(styles=table.style),
            has_header_row=table.header_row_count > 0,
            cell_grid=cell_grid,
            table_id=f"table-{index}",
        )

    def _convert_table_cell(self, cell: ExtractedCell, cell_id: str) -> TableCell:
        paragraph = Paragraph(
            runs=[TextRun(text=cell.text)],
            style=DocumentStyle(styles={"source_data_type": cell.data_type}),
        )
        return TableCell(
            paragraphs=[paragraph] if cell.text else [],
            rowspan=cell.rowspan,
            colspan=cell.colspan,
            style=DocumentStyle(styles=cell.style),
            shading_color=str(cell.style.get("background_color")) if cell.style.get("background_color") else None,
            cell_id=cell_id,
        )

    def _convert_image(self, image: ExtractedImage) -> Image:
        return Image(
            data=image.data,
            mime_type=image.mime_type,
            description=image.caption or image.alt_text or None,
            width=float(image.width) if image.width is not None else None,
            height=float(image.height) if image.height is not None else None,
            title=image.filename,
            alternate_text=image.alt_text or None,
            properties={
                "filename": image.filename,
                "position": image.position,
            },
        )

    def _load_comments_map(self, path: Path | None) -> dict[str, Comment]:
        if path is None:
            return {}
        try:
            with zipfile.ZipFile(path) as archive:
                if "word/comments.xml" not in archive.namelist():
                    return {}
                root = ET.fromstring(archive.read("word/comments.xml"))
        except Exception:  # pragma: no cover - corrupted optional part
            return {}

        namespace = WORD_NS["w"]
        comments: dict[str, Comment] = {}
        for element in root.findall("w:comment", WORD_NS):
            comment_id = element.get(f"{{{namespace}}}id", "")
            author = element.get(f"{{{namespace}}}author", "Unknown")
            date_value = element.get(f"{{{namespace}}}date")
            comment_date = self._parse_datetime(date_value)
            text = " ".join(
                segment.text.strip()
                for segment in element.findall(".//w:t", WORD_NS)
                if segment.text and segment.text.strip()
            )
            comments[comment_id] = Comment(
                author=author,
                text=text,
                created_at=comment_date or datetime.utcnow(),
                target_id=comment_id,
            )
        return comments

    def _paragraph_comments(
        self,
        paragraph: DocxParagraph,
        comment_map: Mapping[str, Comment],
    ) -> list[Comment]:
        comment_ids: list[str] = []
        for ref in paragraph._element.findall(".//w:commentReference", WORD_NS):
            comment_id = ref.get(qn("w:id"))
            if comment_id:
                comment_ids.append(comment_id)
        return [comment_map[comment_id] for comment_id in comment_ids if comment_id in comment_map]

    def _load_app_properties(self, path: Path | None) -> dict[str, Any]:
        if path is None:
            return {}
        try:
            with zipfile.ZipFile(path) as archive:
                if "docProps/app.xml" not in archive.namelist():
                    return {}
                root = ET.fromstring(archive.read("docProps/app.xml"))
        except Exception:  # pragma: no cover - corrupted optional part
            return {}

        properties: dict[str, Any] = {}
        for child in root:
            tag = self._local_name(child.tag)
            if child.text and child.text.strip():
                properties[tag] = child.text.strip()
        return properties

    def _load_custom_properties(self, path: Path | None) -> dict[str, Any]:
        if path is None:
            return {}
        try:
            with zipfile.ZipFile(path) as archive:
                if "docProps/custom.xml" not in archive.namelist():
                    return {}
                root = ET.fromstring(archive.read("docProps/custom.xml"))
        except Exception:  # pragma: no cover - corrupted optional part
            return {}

        properties: dict[str, Any] = {}
        for prop in root.findall(".//{*}property"):
            name = prop.get("name")
            if not name:
                continue
            value_element = next(iter(prop), None)
            if value_element is None:
                properties[name] = None
                continue
            properties[name] = self._parse_custom_property(value_element)
        return properties

    def _extract_sections(
        self,
        document: DocxDocumentType,
        headings: Sequence[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        boundaries = self._section_boundaries(document)
        sections: list[dict[str, Any]] = []

        for section_index, section in enumerate(document.sections, start=1):
            start_paragraph, end_paragraph = boundaries[min(section_index - 1, len(boundaries) - 1)]
            section_headings = [
                heading
                for heading in headings
                if start_paragraph <= self._heading_paragraph_index(heading) <= end_paragraph
            ]
            title = section_headings[0]["text"] if section_headings else f"Section {section_index}"
            sections.append(
                {
                    "index": section_index,
                    "title": title,
                    "start_paragraph": start_paragraph,
                    "end_paragraph": end_paragraph,
                    "orientation": str(section.orientation).split(".")[-1].lower(),
                    "start_type": str(section.start_type).split(".")[-1].lower(),
                    "page_width": self._length_to_inches(section.page_width),
                    "page_height": self._length_to_inches(section.page_height),
                    "left_margin": self._length_to_inches(section.left_margin),
                    "right_margin": self._length_to_inches(section.right_margin),
                    "top_margin": self._length_to_inches(section.top_margin),
                    "bottom_margin": self._length_to_inches(section.bottom_margin),
                    "header_linked": bool(getattr(section.header, "is_linked_to_previous", False)),
                    "footer_linked": bool(getattr(section.footer, "is_linked_to_previous", False)),
                }
            )

        if not sections:
            sections.append({"index": 1, "title": "Section 1"})

        return sections

    def _section_boundaries(self, document: DocxDocumentType) -> list[tuple[int, int]]:
        total_paragraphs = len(document.paragraphs)
        if total_paragraphs == 0:
            return [(0, 0)]

        boundaries: list[tuple[int, int]] = []
        start_index = 0
        paragraph_index = -1

        for child in document.element.body.iterchildren():
            tag = self._local_name(child.tag)
            if tag == "p":
                paragraph_index += 1
                if child.find(".//w:sectPr", WORD_NS) is not None:
                    boundaries.append((start_index, paragraph_index))
                    start_index = paragraph_index + 1
            elif tag == "sectPr":
                boundaries.append((start_index, max(paragraph_index, start_index)))
                start_index = paragraph_index + 1

        if not boundaries or boundaries[-1][1] < total_paragraphs - 1:
            boundaries.append((start_index, total_paragraphs - 1))

        return boundaries

    def _build_outline(self, headings: Sequence[dict[str, Any]]) -> list[dict[str, Any]]:
        stack: list[HeadingNode] = []
        roots: list[HeadingNode] = []

        for index, heading in enumerate(headings):
            level = int(heading["level"])
            node = HeadingNode(
                level=level,
                text=str(heading["text"]),
                paragraph_index=index,
                paragraph_id=heading.get("paragraph_id"),
            )
            while stack and stack[-1].level >= level:
                stack.pop()
            if stack:
                stack[-1].children.append(node)
            else:
                roots.append(node)
            stack.append(node)

        return [self._heading_node_to_dict(node) for node in roots]

    def _heading_node_to_dict(self, node: HeadingNode) -> dict[str, Any]:
        return {
            "level": node.level,
            "text": node.text,
            "paragraph_id": node.paragraph_id,
            "children": [self._heading_node_to_dict(child) for child in node.children],
        }

    # ------------------------------------------------------------------ #
    # Internal: document writing helpers
    # ------------------------------------------------------------------ #
    def _normalize_content(
        self,
        content: DocumentContent | Mapping[str, Any] | Iterable[str] | None,
    ) -> DocumentContent:
        if content is None:
            return DocumentContent(format=OfficeFormat.DOCX)
        if isinstance(content, DocumentContent):
            return content
        if isinstance(content, Mapping):
            paragraphs = [self._coerce_paragraph(value) for value in content.get("paragraphs", [])]
            tables = [self._coerce_table(value) for value in content.get("tables", [])]
            images = [self._coerce_image(value) for value in content.get("images", [])]
            metadata = self._normalize_metadata(content.get("metadata"))
            return DocumentContent(
                format=OfficeFormat.DOCX,
                paragraphs=paragraphs,
                tables=tables,
                images=images,
                metadata=metadata or Metadata(),
            )
        return DocumentContent(
            format=OfficeFormat.DOCX,
            paragraphs=[Paragraph(runs=[TextRun(text=str(value))]) for value in content],
        )

    def _normalize_metadata(self, value: Metadata | Mapping[str, Any] | None) -> Metadata | None:
        if value is None:
            return None
        if isinstance(value, Metadata):
            return value
        if isinstance(value, Mapping):
            keywords = value.get("keywords", [])
            if isinstance(keywords, str):
                keywords = self._split_keywords(keywords)
            return Metadata(
                title=self._none_if_blank(value.get("title")),
                subject=self._none_if_blank(value.get("subject")),
                author=self._none_if_blank(value.get("author")),
                company=self._none_if_blank(value.get("company")),
                category=self._none_if_blank(value.get("category")),
                keywords=list(keywords),
                created_at=value.get("created_at"),
                modified_at=value.get("modified_at"),
                last_printed_at=value.get("last_printed_at"),
                revision=self._none_if_blank(value.get("revision")),
                custom_properties=dict(value.get("custom_properties", {})),
            )
        raise InvalidDocumentStructureError("Unsupported metadata payload")

    def _coerce_paragraph(self, value: Paragraph | str | Mapping[str, Any]) -> Paragraph:
        if isinstance(value, Paragraph):
            return value
        if isinstance(value, str):
            return Paragraph(runs=[TextRun(text=value)])
        if not isinstance(value, Mapping):
            raise InvalidDocumentStructureError("Unsupported paragraph payload")

        runs = value.get("runs")
        if runs:
            run_models = [
                run
                if isinstance(run, TextRun)
                else TextRun(
                    text=str(run.get("text", "")),
                    style=DocumentStyle(
                        font_family=run.get("font_family", "Calibri"),
                        font_size=float(run.get("font_size", 12.0)),
                        bold=bool(run.get("bold", False)),
                        italic=bool(run.get("italic", False)),
                        underline=bool(run.get("underline", False)),
                    ),
                    hyperlink=run.get("hyperlink"),
                    language=run.get("language"),
                )
                for run in runs
            ]
        else:
            run_models = [TextRun(text=str(value.get("text", "")))]

        style_value = value.get("style")
        style = style_value if isinstance(style_value, DocumentStyle) else DocumentStyle(
            styles={"style_name": str(style_value)} if style_value else {}
        )

        heading_level = value.get("heading_level")
        return Paragraph(
            runs=run_models,
            style=style,
            paragraph_id=value.get("paragraph_id"),
            numbering_level=value.get("numbering_level"),
            numbering_id=value.get("numbering_id"),
            is_heading=bool(value.get("is_heading", heading_level is not None)),
            heading_level=int(heading_level) if heading_level is not None else None,
            bookmarked=bool(value.get("bookmarked", False)),
            bookmark_id=value.get("bookmark_id"),
        )

    def _coerce_table(self, value: Table | Sequence[Sequence[Any]]) -> Table:
        if isinstance(value, Table):
            return value
        rows: list[list[TableCell]] = []
        for row_index, row in enumerate(value):
            converted_row: list[TableCell] = []
            for col_index, cell_value in enumerate(row):
                if isinstance(cell_value, TableCell):
                    converted_row.append(cell_value)
                    continue
                text = str(cell_value) if cell_value is not None else ""
                converted_row.append(
                    TableCell(
                        paragraphs=[Paragraph(runs=[TextRun(text=text)])] if text else [],
                        cell_id=f"generated-r{row_index + 1}-c{col_index + 1}",
                    )
                )
            rows.append(converted_row)
        return Table(rows=rows)

    def _coerce_image(self, value: Image | str | Path | Mapping[str, Any]) -> Image:
        if isinstance(value, Image):
            return value
        if isinstance(value, (str, Path)):
            path = Path(value).expanduser()
            return Image(
                data=path.read_bytes(),
                mime_type=self._guess_mime_type(path),
                title=path.name,
            )
        if not isinstance(value, Mapping):
            raise InvalidDocumentStructureError("Unsupported image payload")

        if "data" in value:
            data = value["data"]
        elif "path" in value:
            path = Path(value["path"]).expanduser()
            data = path.read_bytes()
        else:
            raise InvalidDocumentStructureError("Image payload requires 'data' or 'path'")

        return Image(
            data=data,
            mime_type=str(value.get("mime_type", "image/png")),
            description=self._none_if_blank(value.get("description")),
            width=float(value["width"]) if value.get("width") is not None else None,
            height=float(value["height"]) if value.get("height") is not None else None,
            title=self._none_if_blank(value.get("title")),
            alternate_text=self._none_if_blank(value.get("alternate_text")),
            properties=dict(value.get("properties", {})),
        )

    def _apply_metadata(self, document: DocxDocumentType, metadata: Metadata) -> None:
        core = document.core_properties
        core.title = metadata.title or ""
        core.subject = metadata.subject or ""
        core.author = metadata.author or ""
        core.category = metadata.category or ""
        core.keywords = ", ".join(metadata.keywords)
        if metadata.created_at is not None:
            core.created = metadata.created_at
        if metadata.modified_at is not None:
            core.modified = metadata.modified_at

    def _apply_paragraph_style(self, paragraph: DocxParagraph, style: DocumentStyle) -> None:
        if style.alignment:
            paragraph.alignment = self._docx_alignment(style.alignment)
        fmt = paragraph.paragraph_format
        if style.line_spacing:
            fmt.line_spacing = style.line_spacing
        if style.spacing_before:
            fmt.space_before = Pt(style.spacing_before)
        if style.spacing_after:
            fmt.space_after = Pt(style.spacing_after)
        if style.indentation_left:
            fmt.left_indent = Pt(style.indentation_left)
        if style.indentation_right:
            fmt.right_indent = Pt(style.indentation_right)
        if style.indentation_first_line:
            fmt.first_line_indent = Pt(style.indentation_first_line)

    def _append_run(self, paragraph: DocxParagraph, run_model: TextRun) -> None:
        run = paragraph.add_run(run_model.text)
        style = run_model.style
        run.bold = style.bold
        run.italic = style.italic
        run.underline = style.underline
        if style.font_family:
            run.font.name = style.font_family
        if style.font_size:
            run.font.size = Pt(style.font_size)
        rgb = self._hex_to_rgb(style.text_color)
        if rgb is not None:
            run.font.color.rgb = rgb

    def _populate_docx_cell(self, docx_cell: Any, cell: TableCell) -> None:
        docx_cell.text = ""
        if not cell.paragraphs:
            return

        first_paragraph = docx_cell.paragraphs[0]
        template_paragraph = cell.paragraphs[0]
        self._apply_paragraph_style(first_paragraph, template_paragraph.style)
        if template_paragraph.runs:
            for run in template_paragraph.runs:
                self._append_run(first_paragraph, run)
        else:
            first_paragraph.add_run(template_paragraph.text_content())

        for extra in cell.paragraphs[1:]:
            next_paragraph = docx_cell.add_paragraph()
            self._apply_paragraph_style(next_paragraph, extra.style)
            if extra.runs:
                for run in extra.runs:
                    self._append_run(next_paragraph, run)
            else:
                next_paragraph.add_run(extra.text_content())

    def _image_dimension(self, value: float | None) -> Any | None:
        if value is None:
            return None
        inches = value / 96.0 if value > 40 else value
        return Inches(inches)

    def _set_picture_alt_text(self, shape: Any, *, alt_text: str, title: str) -> None:
        try:
            docpr = shape._inline.docPr
            if title:
                docpr.set("title", title)
            if alt_text:
                docpr.set("descr", alt_text)
        except Exception:  # pragma: no cover - python-docx internal compatibility
            logger.debug("Unable to set DOCX picture alt text")

    # ------------------------------------------------------------------ #
    # Internal: templating helpers
    # ------------------------------------------------------------------ #
    def _render_table(
        self,
        table: DocxTable,
        context: Mapping[str, Any],
        *,
        strict: bool,
    ) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._render_paragraph(paragraph, context, strict=strict)

    def _render_paragraph(
        self,
        paragraph: DocxParagraph,
        context: Mapping[str, Any],
        *,
        strict: bool,
    ) -> None:
        original_text = paragraph.text
        rendered_text = self._render_text(original_text, context, strict=strict)

        for run in paragraph.runs:
            run.text = self._render_text(run.text, context, strict=strict)

        if paragraph.text != rendered_text:
            self._rewrite_paragraph(paragraph, rendered_text)

    def _rewrite_paragraph(self, paragraph: DocxParagraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        paragraph.add_run(text)

    def _render_core_properties(
        self,
        document: DocxDocumentType,
        context: Mapping[str, Any],
        *,
        strict: bool,
    ) -> None:
        core = document.core_properties
        for field_name in ("title", "subject", "author", "category", "comments", "keywords"):
            value = getattr(core, field_name, None)
            if isinstance(value, str) and value:
                setattr(core, field_name, self._render_text(value, context, strict=strict))

    def _render_text(
        self,
        value: str,
        context: Mapping[str, Any],
        *,
        strict: bool,
    ) -> str:
        def replace(match: re.Match[str]) -> str:
            expression = match.group(1)
            resolved = self._resolve_expression(expression, context)
            if resolved is None:
                if strict:
                    raise DocumentValidationError(
                        "Template placeholder could not be resolved",
                        details=expression,
                    )
                return ""
            return self._stringify_template_value(resolved)

        return PLACEHOLDER_RE.sub(replace, value)

    def _resolve_expression(self, expression: str, context: Mapping[str, Any]) -> Any:
        current: Any = context
        parts = [part for part in re.split(r"\.(?![^\[]*\])", expression) if part]

        for part in parts:
            normalized = part
            if "[" in part and part.endswith("]"):
                base, index_text = part[:-1].split("[", 1)
                if base:
                    current = self._resolve_value(current, base)
                if current is None:
                    return None
                try:
                    current = current[int(index_text)]
                except Exception:
                    return None
                continue
            current = self._resolve_value(current, normalized)
            if current is None:
                return None
        return current

    def _resolve_value(self, current: Any, key: str) -> Any:
        if isinstance(current, Mapping):
            return current.get(key)
        if hasattr(current, key):
            return getattr(current, key)
        if isinstance(current, Sequence) and not isinstance(current, (str, bytes, bytearray)):
            try:
                return current[int(key)]
            except Exception:
                return None
        return None

    def _stringify_template_value(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.isoformat(sep=" ", timespec="seconds")
        if isinstance(value, Mapping):
            return ", ".join(f"{k}={v}" for k, v in value.items())
        if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
            return "\n".join(self._stringify_template_value(item) for item in value)
        return str(value)

    # ------------------------------------------------------------------ #
    # Internal: XML and value helpers
    # ------------------------------------------------------------------ #
    def _paragraph_id(self, paragraph: DocxParagraph) -> str | None:
        para_id = paragraph._element.get(f"{{{WORD_NS['w14']}}}paraId")
        return para_id or None

    def _heading_level(self, style_name: str | None) -> int | None:
        if not style_name:
            return None
        match = re.match(r"heading\s+(\d+)$", style_name.strip().lower())
        if not match:
            return None
        return int(match.group(1))

    def _numbering_info(self, paragraph: DocxParagraph) -> tuple[int | None, str | None]:
        num_pr = paragraph._element.find(".//w:numPr", WORD_NS)
        if num_pr is None:
            return None, None

        level = num_pr.find("w:ilvl", WORD_NS)
        numbering = num_pr.find("w:numId", WORD_NS)
        level_value = int(level.get(qn("w:val"))) if level is not None else None
        numbering_id = numbering.get(qn("w:val")) if numbering is not None else None
        return level_value, numbering_id

    def _bookmark_id(self, paragraph: DocxParagraph) -> str | None:
        bookmark = paragraph._element.find(".//w:bookmarkStart", WORD_NS)
        if bookmark is None:
            return None
        return bookmark.get(qn("w:name")) or bookmark.get(qn("w:id"))

    def _run_language(self, run: DocxRun) -> str | None:
        lang = run._element.find(".//w:lang", WORD_NS)
        if lang is None:
            return None
        return lang.get(qn("w:val")) or lang.get(qn("w:eastAsia")) or lang.get(qn("w:bidi"))

    def _run_hyperlink(self, run: DocxRun, document: DocxDocumentType) -> str | None:
        parent = run._element.getparent()
        while parent is not None and self._local_name(parent.tag) != "hyperlink":
            parent = parent.getparent()
        if parent is None:
            return None
        rel_id = parent.get(qn("r:id"))
        if not rel_id:
            return None
        relation = document.part.rels.get(rel_id)
        return getattr(relation, "target_ref", None)

    def _parse_custom_property(self, value_element: ET.Element) -> Any:
        tag = self._local_name(value_element.tag)
        raw = value_element.text or ""
        if tag in {"i1", "i2", "i4", "i8", "int", "uint"}:
            try:
                return int(raw)
            except ValueError:
                return raw
        if tag in {"r4", "r8", "decimal"}:
            try:
                return float(raw)
            except ValueError:
                return raw
        if tag == "bool":
            return raw.lower() == "true"
        if tag in {"filetime", "date"}:
            return self._parse_datetime(raw) or raw
        return raw

    def _refresh_wcag_report(self, path: Path | None) -> None:
        if path is None:
            self._last_wcag_report = None
            return
        try:
            self._last_wcag_report = self.accessibility_processor.check_wcag_compliance(path)
        except Exception as exc:  # pragma: no cover - best effort integration
            logger.debug("Unable to generate WCAG report for %s: %s", path, exc)
            self._last_wcag_report = None

    def _docx_alignment(self, value: str | None) -> Any | None:
        if value is None:
            return None
        normalized = value.lower()
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get(normalized)

    def _alignment_name(self, value: Any) -> str:
        mapping = {
            getattr(WD_ALIGN_PARAGRAPH, "LEFT", object()): "left",
            getattr(WD_ALIGN_PARAGRAPH, "CENTER", object()): "center",
            getattr(WD_ALIGN_PARAGRAPH, "RIGHT", object()): "right",
            getattr(WD_ALIGN_PARAGRAPH, "JUSTIFY", object()): "justify",
        }
        return mapping.get(value, "left")

    def _style_type_name(self, value: Any) -> str:
        mapping = {
            getattr(WD_STYLE_TYPE, "PARAGRAPH", object()): "paragraph",
            getattr(WD_STYLE_TYPE, "CHARACTER", object()): "character",
            getattr(WD_STYLE_TYPE, "TABLE", object()): "table",
            getattr(WD_STYLE_TYPE, "LIST", object()): "list",
        }
        return mapping.get(value, "unknown")

    def _heading_paragraph_index(self, heading: Mapping[str, Any]) -> int:
        try:
            return int(heading.get("paragraph_index", 0))
        except Exception:
            return 0

    def _guess_mime_type(self, path: Path) -> str:
        suffix = path.suffix.lower()
        mapping = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".tif": "image/tiff",
            ".tiff": "image/tiff",
            ".svg": "image/svg+xml",
        }
        return mapping.get(suffix, "application/octet-stream")

    def _split_keywords(self, value: str | None) -> list[str]:
        if not value:
            return []
        if ";" in value:
            parts = value.split(";")
        else:
            parts = value.split(",")
        return [part.strip() for part in parts if part.strip()]

    def _hex_to_rgb(self, value: str | None) -> RGBColor | None:
        if not value or RGBColor is None:
            return None
        normalized = value.lstrip("#")
        if len(normalized) != 6:
            return None
        try:
            return RGBColor.from_string(normalized.upper())
        except Exception:  # pragma: no cover - library validation
            return None

    def _rgb_to_hex(self, value: Any) -> str:
        if value is None:
            return "#000000"
        text = str(value)
        if re.fullmatch(r"[0-9A-Fa-f]{6}", text):
            return f"#{text.upper()}"
        return "#000000"

    def _parse_datetime(self, value: str | None) -> datetime | None:
        if not value:
            return None
        try:
            return datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            return None

    def _length_to_inches(self, value: Any) -> float | None:
        try:
            return round(float(value.inches), 4)
        except Exception:
            return None

    def _pt_value(self, value: Any, default: float) -> float:
        try:
            if value is None:
                return default
            return float(value.pt)
        except Exception:
            try:
                return float(value)
            except Exception:
                return default

    def _float_value(self, value: Any, default: float) -> float:
        try:
            if value is None:
                return default
            return float(value)
        except Exception:
            return default

    def _clean_text(self, value: str | None) -> str:
        return " ".join((value or "").split())

    def _local_name(self, tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    def _none_if_blank(self, value: Any) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        return text or None


__all__ = [
    "WordProcessor",
    "WordProcessorError",
    "WordNotFoundError",
    "WordCorruptedError",
]
