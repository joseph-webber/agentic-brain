# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Joseph Webber <joseph.webber@me.com>
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""iCloud loader for RAG pipelines.

Supports:
- iCloud Drive documents
- Text files, PDFs, Word documents
- Folder navigation
"""

import logging
import mimetypes
import os
from pathlib import Path
from typing import Any, Optional

from .base import BaseLoader, LoadedDocument

logger = logging.getLogger(__name__)

# Check for pyicloud
try:
    from pyicloud import PyiCloudService
    from pyicloud.exceptions import PyiCloudFailedLoginException

    PYICLOUD_AVAILABLE = True
except ImportError:
    PYICLOUD_AVAILABLE = False


class iCloudLoader(BaseLoader):
    """Load documents from iCloud Drive.

    Supports:
    - Text files (.txt, .md, .json, .csv)
    - PDFs (with text extraction)
    - Word documents (.docx)
    - Folder navigation

    Note: Requires 2FA. On first use, you'll need to verify via SMS/device.

    Example:
        loader = iCloudLoader(
            apple_id="your@icloud.com",
            password="your-app-specific-password"
        )

        docs = loader.load_folder("Documents/Work")
        results = loader.search("report")
    """

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".json",
        ".csv",
        ".pdf",
        ".docx",
        ".rtf",
        ".html",
    }

    def __init__(
        self,
        apple_id: Optional[str] = None,
        password: Optional[str] = None,
        cookie_directory: str = ".icloud",
        max_file_size_mb: int = 50,
    ):
        if not PYICLOUD_AVAILABLE:
            raise ImportError("pyicloud not installed. Run: pip install pyicloud")

        self.apple_id = apple_id or os.environ.get("ICLOUD_APPLE_ID")
        self.password = password or os.environ.get("ICLOUD_PASSWORD")
        self.cookie_directory = cookie_directory
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self._api = None

    @property
    def source_name(self) -> str:
        return "icloud"

    def authenticate(self) -> bool:
        """Authenticate with iCloud."""
        try:
            if not self.apple_id or not self.password:
                raise ValueError("apple_id and password required")

            self._api = PyiCloudService(
                self.apple_id, self.password, cookie_directory=self.cookie_directory
            )

            if self._api.requires_2fa:
                logger.warning("Two-factor authentication required")
                print("Two-factor authentication required.")
                code = input("Enter the code sent to your device: ")
                if not self._api.validate_2fa_code(code):
                    logger.error("Invalid 2FA code")
                    return False

                if not self._api.is_trusted_session:
                    print("Session not trusted. Requesting trust...")
                    self._api.trust_session()

            logger.info("iCloud authentication successful")
            return True

        except PyiCloudFailedLoginException as e:
            logger.error(f"iCloud login failed: {e}")
            return False
        except Exception as e:
            logger.error(f"iCloud authentication failed: {e}")
            return False

    def _ensure_authenticated(self):
        if not self._api and not self.authenticate():
            raise RuntimeError("Failed to authenticate with iCloud")

    def _get_folder(self, path: str):
        """Navigate to a folder by path."""
        if not path or path == "/":
            return self._api.drive.root

        current = self._api.drive.root
        parts = [p for p in path.split("/") if p]

        for part in parts:
            found = False
            for item in current.dir():
                if item == part:
                    current = current[item]
                    found = True
                    break

            if not found:
                raise FileNotFoundError(f"Folder not found: {part}")

        return current

    def _download_file(self, item) -> Optional[bytes]:
        """Download file content."""
        try:
            response = item.open(stream=True)
            return response.raw.read()
        except Exception as e:
            logger.error(f"Failed to download {item.name}: {e}")
            return None

    def _extract_content(
        self, data: bytes, filename: str, mime_type: str
    ) -> Optional[str]:
        """Extract text content from file data."""
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._extract_text_from_pdf(data)
        elif ext in (".txt", ".md", ".csv", ".json"):
            return data.decode("utf-8", errors="replace")
        elif ext == ".html":
            return self._clean_html(data.decode("utf-8", errors="replace"))
        elif ext == ".docx":
            try:
                from io import BytesIO

                import docx

                doc = docx.Document(BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                logger.warning("python-docx not installed")
                return None

        return None

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a document by path."""
        self._ensure_authenticated()

        try:
            parts = doc_id.rsplit("/", 1)
            if len(parts) == 2:
                folder_path, filename = parts
                folder = self._get_folder(folder_path)
                item = folder[filename]
            else:
                item = self._api.drive.root[doc_id]

            if item.type != "file":
                return None

            ext = Path(item.name).suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                logger.debug(f"Unsupported extension: {item.name}")
                return None

            size = item.size or 0
            if size > self.max_file_size:
                logger.warning(f"File too large: {item.name} ({size} bytes)")
                return None

            data = self._download_file(item)
            if not data:
                return None

            mime_type = mimetypes.guess_type(item.name)[0] or "application/octet-stream"
            content = self._extract_content(data, item.name, mime_type)

            if not content:
                return None

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=doc_id,
                filename=item.name,
                mime_type=mime_type,
                modified_at=(
                    item.date_modified if hasattr(item, "date_modified") else None
                ),
                size_bytes=size,
                metadata={"path": doc_id, "name": item.name, "type": item.type},
            )

        except Exception as e:
            logger.error(f"Failed to load iCloud document {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all documents from an iCloud Drive folder."""
        self._ensure_authenticated()

        try:
            folder = self._get_folder(folder_path)
            return self._load_folder_items(folder, folder_path, recursive)
        except FileNotFoundError as e:
            logger.error(str(e))
            return []

    def _load_folder_items(
        self, folder, path: str, recursive: bool
    ) -> list[LoadedDocument]:
        """Recursively load items from folder."""
        documents = []

        for name in folder.dir():
            item = folder[name]
            item_path = f"{path}/{name}" if path else name

            if item.type == "folder" and recursive:
                documents.extend(self._load_folder_items(item, item_path, recursive))
            elif item.type == "file":
                doc = self.load_document(item_path)
                if doc:
                    documents.append(doc)

        logger.info(f"Loaded {len(documents)} documents from {path}")
        return documents

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search for documents (simple name-based search)."""
        self._ensure_authenticated()

        query_lower = query.lower()
        documents = []

        def search_folder(folder, path: str):
            for name in folder.dir():
                if len(documents) >= max_results:
                    return

                item = folder[name]
                item_path = f"{path}/{name}" if path else name

                if item.type == "folder":
                    search_folder(item, item_path)
                elif query_lower in name.lower():
                    doc = self.load_document(item_path)
                    if doc:
                        documents.append(doc)

        search_folder(self._api.drive.root, "")
        logger.info(f"Found {len(documents)} documents matching '{query}'")
        return documents

    def list_folders(self, path: str = "") -> list[dict[str, str]]:
        """List subfolders in a path."""
        self._ensure_authenticated()

        try:
            folder = self._get_folder(path)
            folders = []

            for name in folder.dir():
                item = folder[name]
                if item.type == "folder":
                    folders.append(
                        {"name": name, "path": f"{path}/{name}" if path else name}
                    )

            return folders
        except FileNotFoundError:
            return []


__all__ = ["iCloudLoader"]
