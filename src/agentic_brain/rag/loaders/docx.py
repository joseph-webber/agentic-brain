# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Joseph Webber <joseph.webber@me.com>
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Microsoft Word document loaders for RAG pipelines.

Enhanced with Docling-grade features:
- Heading hierarchy preserved as Markdown heading levels
- Table extraction with proper Markdown table formatting
- List detection (bullet and numbered)
- Structured metadata for sections and tables
"""

import logging
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from .base import BaseLoader, LoadedDocument

logger = logging.getLogger(__name__)

# Check for python-docx
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Heading style name → Markdown level
_HEADING_MAP = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Heading 5": "#####",
    "Heading 6": "######",
    "Title": "#",
    "Subtitle": "##",
}


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to a Markdown table string."""
    rows_data: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows_data.append(cells)

    if not rows_data:
        return ""

    # Determine column count from the widest row
    col_count = max(len(r) for r in rows_data)
    # Pad short rows
    for r in rows_data:
        while len(r) < col_count:
            r.append("")

    lines: list[str] = []
    # Header row
    lines.append("| " + " | ".join(rows_data[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
    # Data rows
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _table_to_structured(table) -> dict[str, Any]:
    """Extract table as a structured dict with headers and rows."""
    rows_data: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows_data.append(cells)

    if not rows_data:
        return {"headers": [], "rows": []}

    return {
        "headers": rows_data[0],
        "rows": rows_data[1:],
    }


class DocxLoader(BaseLoader):
    """Load and extract text from DOCX documents.

    Enhanced features (inspired by Docling, zero heavy dependencies):
    - Heading hierarchy preserved with Markdown heading levels
    - Table extraction with Markdown table formatting
    - Structured metadata including tables and sections
    - to_markdown() support via LoadedDocument

    Example:
        loader = DocxLoader()
        doc = loader.load_document("report.docx")
        print(doc.to_markdown())   # full Markdown with headings & tables
        print(doc.metadata["tables"])  # structured table data
    """

    def __init__(
        self,
        base_path: str = ".",
        include_headers_footers: bool = False,
        include_comments: bool = False,
        extract_tables: bool = True,
    ):
        self.base_path = Path(base_path)
        self.include_headers_footers = include_headers_footers
        self.include_comments = include_comments
        self.extract_tables = extract_tables

    @property
    def source_name(self) -> str:
        return "docx"

    def authenticate(self) -> bool:
        """No authentication needed for local files."""
        return True

    def _extract_text(self, docx_bytes: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from DOCX bytes with structural awareness.

        Returns:
            Tuple of (content_string, extra_metadata).
        """
        extra_meta: dict[str, Any] = {}

        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed")
            return "[DOCX content - install python-docx]", extra_meta

        try:
            document = docx.Document(BytesIO(docx_bytes))
            parts: list[str] = []
            sections: list[dict[str, Any]] = []
            current_section: Optional[str] = None

            # Iterate over document body elements to preserve order
            for element in document.element.body:
                tag = element.tag.split("}")[-1]  # strip namespace

                if tag == "p":
                    # It's a paragraph
                    para = None
                    for p in document.paragraphs:
                        if p._element is element:
                            para = p
                            break
                    if para is None or not para.text.strip():
                        continue

                    style_name = para.style.name if para.style else ""
                    md_prefix = _HEADING_MAP.get(style_name, "")

                    if md_prefix:
                        heading_text = f"{md_prefix} {para.text.strip()}"
                        parts.append(heading_text)
                        current_section = para.text.strip()
                        sections.append(
                            {"heading": current_section, "level": len(md_prefix)}
                        )
                    elif style_name.startswith("List"):
                        parts.append(f"- {para.text.strip()}")
                    else:
                        parts.append(para.text.strip())

                elif tag == "tbl" and self.extract_tables:
                    # It's a table
                    for tbl in document.tables:
                        if tbl._element is element:
                            md_table = _table_to_markdown(tbl)
                            if md_table:
                                parts.append(md_table)
                            break

            # Extract structured tables into metadata
            if self.extract_tables:
                structured_tables: list[dict[str, Any]] = []
                for i, tbl in enumerate(document.tables):
                    s = _table_to_structured(tbl)
                    s["caption"] = f"Table {i + 1}"
                    structured_tables.append(s)
                if structured_tables:
                    extra_meta["tables"] = structured_tables
                    extra_meta["table_count"] = len(structured_tables)

            if sections:
                extra_meta["sections"] = sections

            return "\n\n".join(parts), extra_meta
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return "", extra_meta

    def load_document(self, doc_id: str) -> Optional[LoadedDocument]:
        """Load a single DOCX document."""
        try:
            path = Path(doc_id)
            if not path.is_absolute():
                path = self.base_path / path

            if not path.exists():
                logger.error(f"File not found: {path}")
                return None

            if path.suffix.lower() not in (".docx", ".doc"):
                logger.warning(f"Not a Word document: {path}")
                return None

            with open(path, "rb") as f:
                docx_bytes = f.read()

            content, extra_meta = self._extract_text(docx_bytes)

            metadata: dict[str, Any] = {"path": str(path)}
            metadata.update(extra_meta)

            return LoadedDocument(
                content=content,
                source=self.source_name,
                source_id=str(path),
                filename=path.name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=len(docx_bytes),
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load DOCX {doc_id}: {e}")
            return None

    def load_folder(
        self, folder_path: str, recursive: bool = True
    ) -> list[LoadedDocument]:
        """Load all DOCX files from a folder."""
        docs = []
        path = Path(folder_path)
        if not path.is_absolute():
            path = self.base_path / path

        if not path.exists():
            logger.error(f"Folder not found: {path}")
            return docs

        pattern = "**/*.docx" if recursive else "*.docx"

        for docx_path in path.glob(pattern):
            doc = self.load_document(str(docx_path))
            if doc:
                docs.append(doc)

        logger.info(f"Loaded {len(docs)} DOCX files from {folder_path}")
        return docs

    def search(self, query: str, max_results: int = 50) -> list[LoadedDocument]:
        """Search not supported for local files."""
        return []


# Alias for backward compatibility
WordLoader = DocxLoader


__all__ = ["DocxLoader", "WordLoader"]
