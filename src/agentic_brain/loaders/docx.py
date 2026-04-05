# SPDX-License-Identifier: Apache-2.0
# Copyright 2024-2026 Agentic Brain Contributors
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""DOCX (Word) document loader."""

import logging
from pathlib import Path

from .base import Document, SyncDocumentLoader

logger = logging.getLogger(__name__)


class DocxLoader(SyncDocumentLoader):
    """Load DOCX (Word) files."""

    def load_sync(self, source: str | Path) -> list[Document]:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            ) from None

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        try:
            doc = DocxDocument(path)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                text_parts.append("\n[TABLE]")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

            content = "\n".join(text_parts)

            metadata = {"filename": path.name, "size": path.stat().st_size}

            if doc.core_properties:
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.subject:
                    metadata["subject"] = props.subject

            return [
                Document(
                    content=content,
                    source=str(path),
                    metadata=metadata,
                )
            ]

        except Exception as e:
            logger.error(f"Error loading DOCX {path}: {e}")
            raise ValueError(f"Failed to load DOCX: {e}") from e

    def supported_extensions(self) -> list[str]:
        """Return supported extensions."""
        return [".docx"]
